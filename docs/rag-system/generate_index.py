"""
generate_index.py

Generates docs/rag-system/index.docx — a documentation suite index and system
overview for the Adaptive RAG Clinical Assistant.

Run:
    python docs/rag-system/generate_index.py

Requirements:
    pip install python-docx
"""

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Colour constants
# ---------------------------------------------------------------------------
NAVY = RGBColor(0x11, 0x22, 0x40)      # #112240  title
TEAL = RGBColor(0x0E, 0xA5, 0xC9)      # #0EA5C9  subtitle / table header
TEAL_HEX = "0EA5C9"
CALLOUT_BG_HEX = "EBF8FF"
DIAGRAM_BG_HEX = "F3F4F6"
RED_BORDER_HEX = "DC2626"
HEADER_TEXT_COLOR = RGBColor(0xFF, 0xFF, 0xFF)  # white header text

OUTPUT_PATH = Path(__file__).parent / "index.docx"


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def _shade_cell(cell, hex_color: str) -> None:
    """Fill a table cell with a solid background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _shade_para(para, hex_color: str) -> None:
    """Apply a solid paragraph background (shading) to simulate a coloured box."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _add_left_border(para, hex_color: str, width_pt: int = 18) -> None:
    """Add a thick left border to a paragraph (gap / warning box style)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(width_pt))   # eighths of a point
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)


def _add_page_numbers(doc: Document) -> None:
    """Insert 'Page X of Y' into the document footer via Word field codes."""
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _field(instr: str):
        run = para.add_run()
        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar_begin)

        run2 = para.add_run()
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = f" {instr} "
        run2._r.append(instrText)

        run3 = para.add_run()
        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")
        run3._r.append(fldChar_end)

    para.add_run("Page ")
    _field("PAGE")
    para.add_run(" of ")
    _field("NUMPAGES")


# ---------------------------------------------------------------------------
# High-level content helpers
# ---------------------------------------------------------------------------

def _heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level == 1 else TEAL


def add_teal_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    col_widths: list[float] | None = None,
) -> None:
    """Render a table with a teal header row and alternating body rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr_cells[i].text = text
        _shade_cell(hdr_cells[i], TEAL_HEX)
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = HEADER_TEXT_COLOR
        run.font.size = Pt(9.5)

    ALT_BG = "F9FAFB"
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        bg = ALT_BG if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, text in enumerate(row_data):
            row_cells[c_idx].text = text
            _shade_cell(row_cells[c_idx], bg)
            run = row_cells[c_idx].paragraphs[0].runs[0]
            run.font.size = Pt(9.5)

    if col_widths:
        for r_idx in range(len(table.rows)):
            for c_idx, width in enumerate(col_widths):
                table.rows[r_idx].cells[c_idx].width = Inches(width)

    doc.add_paragraph()


def add_callout(doc: Document, label: str, text: str) -> None:
    """Add a light-blue callout box (Courier New, shaded background)."""
    lbl = doc.add_paragraph()
    r = lbl.add_run(label)
    r.font.bold = True
    r.font.size = Pt(10)

    for line in text.splitlines():
        p = doc.add_paragraph(line if line.strip() else " ")
        _shade_para(p, CALLOUT_BG_HEX)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(8.5)

    doc.add_paragraph()


def add_ascii(doc: Document, caption: str, lines: list[str]) -> None:
    """Render an ASCII diagram in a grey code block with a caption."""
    cap = doc.add_paragraph()
    r = cap.add_run(f"Figure: {caption}")
    r.font.bold = True
    r.font.italic = True
    r.font.size = Pt(10)

    for line in lines:
        p = doc.add_paragraph(line if line else " ")
        _shade_para(p, DIAGRAM_BG_HEX)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.15)
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(8)

    doc.add_paragraph()


# mmdc may not be on PATH in all shells on Windows; resolve the full path via npm prefix.
_NPM_PREFIX = subprocess.run(
    ["node", "-e", "process.stdout.write(require('child_process').execSync('npm config get prefix').toString().trim())"],
    capture_output=True,
    text=True,
).stdout.strip()
_MMDC_CMD = str(Path(_NPM_PREFIX) / "mmdc.cmd") if _NPM_PREFIX else "mmdc"


def _try_mermaid(mmd: str, out_path: Path, mmd_source_path: Path | None = None) -> bool:
    """Attempt to render a Mermaid diagram to PNG. Returns True on success."""
    if mmd_source_path and mmd_source_path.exists():
        mmd_file = mmd_source_path
        temp_created = False
    else:
        mmd_file = out_path.with_suffix(".mmd")
        mmd_file.write_text(mmd, encoding="utf-8")
        temp_created = True

    try:
        result = subprocess.run(
            [
                _MMDC_CMD,
                "-i", str(mmd_file),
                "-o", str(out_path),
                "--scale", "2",
                "--backgroundColor", "white",
            ],
            capture_output=True,
            timeout=60,
        )
        return result.returncode == 0 and out_path.exists()
    except (FileNotFoundError, subprocess.TimeoutExpired, OSError):
        return False
    finally:
        if temp_created and mmd_file.exists():
            mmd_file.unlink()


def add_diagram(
    doc: Document,
    caption: str,
    mmd: str,
    ascii_lines: list[str],
    png_name: str,
    mmd_source_path: Path | None = None,
) -> None:
    """Attempt Mermaid PNG render; fall back to ASCII art if mmdc is unavailable."""
    png_path = OUTPUT_PATH.parent / png_name
    if _try_mermaid(mmd, png_path, mmd_source_path):
        cap = doc.add_paragraph()
        r = cap.add_run(f"Figure: {caption}")
        r.font.bold = True
        r.font.italic = True
        r.font.size = Pt(10)
        doc.add_picture(str(png_path), width=Inches(6.0))
        doc.add_paragraph()
    else:
        add_ascii(doc, caption, ascii_lines)


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def add_cover(doc: Document) -> None:
    """Add a styled cover page."""
    doc.add_paragraph()
    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Clinical RAG Assistant — Technical Documentation")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = NAVY

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_para.add_run("Documentation Suite Index & System Overview")
    run.font.size = Pt(14)
    run.font.color.rgb = TEAL

    doc.add_paragraph()

    meta_lines = [
        ("System:", "Adaptive RAG Clinical Assistant"),
        ("Version:", "1.0"),
        ("Date:", "2026-02-23"),
        ("Description:", "An adaptive, hallucination-mitigated Retrieval-Augmented Generation system for clinical trial document intelligence."),
    ]
    for label, value in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label + "  ")
        r1.font.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(value)
        r2.font.size = Pt(11)

    doc.add_page_break()


def add_toc(doc: Document) -> None:
    """Add a Table of Contents placeholder."""
    h = doc.add_heading("Table of Contents", level=1)
    h.runs[0].font.color.rgb = NAVY

    note = doc.add_paragraph(
        "Right-click this area in Microsoft Word and select 'Update Field' "
        "to generate the Table of Contents automatically."
    )
    note.runs[0].font.italic = True
    note.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    note.runs[0].font.size = Pt(10)

    doc.add_page_break()


def build_section_1_overview(doc: Document) -> None:
    """Section 1 — System Overview with architecture diagram."""
    _heading(doc, "1. System Overview")

    doc.add_paragraph(
        "The Adaptive RAG Clinical Assistant is a two-phase Retrieval-Augmented Generation "
        "system purpose-built for clinical trial document intelligence. "
        "In the offline indexing phase, source documents (PDF, DOCX, XML) are ingested, "
        "chunked with a 800-character window and 150-character overlap, embedded using a "
        "medical-domain sentence transformer (S-PubMedBert-MS-MARCO, 768-dim), and stored "
        "in two complementary indexes: a ChromaDB dense vector store (HNSW, cosine distance) "
        "and a BM25 sparse lexical index. "
        "In the online query phase, a HybridRetriever fuses results from both indexes using "
        "Reciprocal Rank Fusion (RRF, k=60), a PromptBuilder assembles a persona-aware prompt "
        "with grounding instructions and source citation tags, an LLM generates the response "
        "(cloud: gpt-4o-mini; local: llama3.1:8b), and a FaithfulnessChecker scores the "
        "response against retrieved context to detect potential hallucination."
    )

    doc.add_paragraph(
        "The system adapts its response style to five distinct user personas defined in "
        "src/personas.py: NOVICE (simplified language, definitions, 300-token limit), "
        "INTERMEDIATE (practical guidance, tables enabled, 500-token limit), "
        "EXPERT (full technical depth, 1000-token limit), "
        "REGULATORY (audit-ready format, compliance focus, 800-token limit), and "
        "EXECUTIVE (concise summary with key metrics, 250-token limit). "
        "For each persona, the system applies different response length limits "
        "(src/config.py RESPONSE_LENGTH_LIMITS), detail levels, format rules (tables, "
        "definitions, key takeaways), and prompt instruction templates "
        "(src/prompts.py:47-107), producing measurably differentiated outputs compared "
        "to a generic non-adaptive baseline."
    )

    doc.add_paragraph(
        "Clinical trial document systems operate in a regulated, safety-critical context "
        "where hallucinated citations, fabricated protocol steps, or incorrect eligibility "
        "criteria can directly harm patients or delay regulatory approval. "
        "This documentation suite provides the technical evidence required for clinical "
        "governance review: it traces every system parameter to its source file, quantifies "
        "retrieval and generation performance against SLA targets, documents the three-layer "
        "hallucination mitigation strategy, and catalogues known gaps and their mitigations. "
        "All claims in this suite are grounded in the actual codebase — no speculative or "
        "aspirational descriptions are included."
    )

    # Architecture diagram
    mmd = """graph TD
  subgraph OFFLINE ["Offline — Indexing Phase"]
    A["Clinical Documents\\nPDF · DOCX · XML"] --> B["DocumentLoader\\nsrc/ingestion.py"]
    B --> C["RecursiveCharacterTextSplitter\\nchunk_size=800 · overlap=150"]
    C --> D["EmbeddingService\\nS-PubMedBert-MS-MARCO · 768-dim"]
    D --> E[("ChromaDB\\nHNSW · cosine")]
    C --> F[("BM25Index")]
  end
  subgraph ONLINE ["Online — Query Phase"]
    G["User Query\\n+ Persona (5 levels)"] --> H["HybridRetriever\\nDense + BM25 + RRF k=60"]
    E --> H
    F --> H
    H --> I["PromptBuilder\\nSYSTEM_PROMPT + _GROUNDING_INSTRUCTION\\n+ [Source N] tags"]
    I --> J["LLMGenerator\\ngpt-4o-mini (cloud) · llama3.1:8b (local)"]
    J --> K["FaithfulnessChecker\\ncosine similarity · warn@0.45 · flag@0.35"]
    K --> L["Verified Response\\n+ faithfulness_score"]
  end
  subgraph EVAL ["Evaluation & Benchmarking"]
    M["Adaptive Evals\\nclassification · readability\\nformat · adaptive_vs_generic"]
    N["Benchmark Evals\\nmodel_comparison · hybrid_vs_semantic\\npersona_evaluation · latency_measurement"]
  end
  L --> M
  L --> N
  style OFFLINE fill:#0B1829,color:#7DD3FC
  style ONLINE  fill:#112240,color:#7DD3FC
  style EVAL    fill:#1A3356,color:#7DD3FC
  style L       fill:#10B981,color:#fff"""

    ascii_lines = [
        "  ┌─────────────────────────────────────────────────────────────────┐",
        "  │  OFFLINE — Indexing Phase                                       │",
        "  │                                                                  │",
        "  │  Clinical Documents (PDF · DOCX · XML)                          │",
        "  │      │                                                           │",
        "  │      ▼                                                           │",
        "  │  DocumentLoader  (src/ingestion.py)                             │",
        "  │      │                                                           │",
        "  │      ▼                                                           │",
        "  │  RecursiveCharacterTextSplitter  chunk_size=800 · overlap=150   │",
        "  │      │                                                           │",
        "  │      ├──────────────────────────────────┐                       │",
        "  │      ▼                                  ▼                       │",
        "  │  EmbeddingService                   BM25Index                   │",
        "  │  S-PubMedBert-MS-MARCO 768-dim      (sparse lexical)            │",
        "  │      │                                                           │",
        "  │      ▼                                                           │",
        "  │  ChromaDB (HNSW · cosine)                                       │",
        "  └─────────────────────────────────────────────────────────────────┘",
        "",
        "  ┌─────────────────────────────────────────────────────────────────┐",
        "  │  ONLINE — Query Phase                                           │",
        "  │                                                                  │",
        "  │  User Query + Persona (5 levels)                                │",
        "  │      │                                                           │",
        "  │      ▼                                                           │",
        "  │  HybridRetriever  Dense + BM25 + RRF k=60                       │",
        "  │      │                                                           │",
        "  │      ▼                                                           │",
        "  │  PromptBuilder  SYSTEM_PROMPT + _GROUNDING_INSTRUCTION          │",
        "  │                 + [Source N] citation tags                      │",
        "  │      │                                                           │",
        "  │      ▼                                                           │",
        "  │  LLMGenerator  gpt-4o-mini (cloud) · llama3.1:8b (local)       │",
        "  │      │                                                           │",
        "  │      ▼                                                           │",
        "  │  FaithfulnessChecker  warn@0.45 · sentence-flag@0.35           │",
        "  │      │                                                           │",
        "  │      ▼                                                           │",
        "  │  ✓ Verified Response + faithfulness_score                       │",
        "  └─────────────────────────────────────────────────────────────────┘",
        "",
        "  ┌─────────────────────────────────────────────────────────────────┐",
        "  │  EVALUATION & BENCHMARKING                                      │",
        "  │                                                                  │",
        "  │  Adaptive:   classification · readability · format ·            │",
        "  │              adaptive_vs_generic                                 │",
        "  │  Benchmark:  model_comparison · hybrid_vs_semantic ·            │",
        "  │              persona_evaluation · latency_measurement            │",
        "  └─────────────────────────────────────────────────────────────────┘",
    ]

    add_diagram(
        doc,
        "Figure 1 — Clinical RAG Assistant: Complete System Architecture",
        mmd,
        ascii_lines,
        "fig0_system_architecture.png",
        mmd_source_path=Path(__file__).parent / "figures" / "fig0_system_architecture.mmd",
    )


def build_section_2_document_map(doc: Document) -> None:
    """Section 2 — Document Map table."""
    _heading(doc, "2. Document Map")

    doc.add_paragraph(
        "The documentation suite consists of five documents. The four specialist documents "
        "below are complementary — each covers a distinct system concern at a level of "
        "detail appropriate for its primary audience. This index document provides the "
        "entry point and navigation hub for all five."
    )

    add_teal_table(
        doc,
        ["Document", "File", "Primary Audience", "Purpose"],
        [
            [
                "RAG Pipeline Architecture",
                "rag_pipeline_updated.docx",
                "ML Engineers, Developers",
                "End-to-end pipeline from ingestion to generation",
            ],
            [
                "Adaptive Evaluation Suite",
                "adaptive_evaluation_updated.docx",
                "Researchers, QA Engineers",
                "Four adaptive evals measuring system intelligence",
            ],
            [
                "Benchmark Evaluations",
                "benchmark_evaluations_updated.docx",
                "ML Engineers, Clinical Informatics",
                "Retrieval quality, embedding, latency benchmarks",
            ],
            [
                "Hallucination Mitigation",
                "hallucination_mitigation_updated.docx",
                "Clinical Governance, Safety Officers",
                "Three-layer defense against fabricated outputs",
            ],
        ],
        col_widths=[1.9, 2.0, 1.9, 2.0],
    )


def build_section_3_summaries(doc: Document) -> None:
    """Section 3 — Document Summaries in callout boxes."""
    _heading(doc, "3. Document Summaries")

    doc.add_paragraph(
        "Each subsection below provides a summary of one specialist document: "
        "the components it covers, its headline claim, the key metrics or parameters "
        "it documents, and the recommended reading context."
    )

    _heading(doc, "3.1 RAG Pipeline Architecture", level=2)
    add_callout(
        doc,
        "rag_pipeline_updated.docx",
        """Components covered:
  src/ingestion.py      — DocumentIngester, RecursiveCharacterTextSplitter
  src/config.py         — all pipeline parameters (chunking, HNSW, retrieval, LLM)
  src/retrieval.py      — HybridRetriever, ReciprocalRankFusion
  src/prompts.py        — ResponseStyler, build_adaptive_prompt(), SYSTEM_PROMPT,
                          _GROUNDING_INSTRUCTION, [Source N] citation tags

Pipeline stages (7):
  1. Document Ingestion (pdfplumber / OCR fallback)
  2. Chunking (RecursiveCharacterTextSplitter, chunk_size=800, overlap=150)
  3. Embedding  (S-PubMedBert-MS-MARCO, 768-dim, three-tier fallback chain)
  4. Index Construction (ChromaDB HNSW cosine + BM25 sparse)
  5. Hybrid Retrieval (Dense + BM25 + RRF k=60, top-5 final)
  6. Prompt Augmentation (persona-aware + grounding + [Source N] tags)
  7. Response Generation (OpenAIProvider or OllamaProvider via LLMFactory)

Key technologies: ChromaDB HNSW, BM25 (rank_bm25), S-PubMedBert-MS-MARCO,
                  gpt-4o-mini (cloud), llama3.1:8b (local)

Read when: modifying any pipeline stage, onboarding as an ML engineer, or
           performing architecture review.""",
    )

    _heading(doc, "3.2 Adaptive Evaluation Suite", level=2)
    add_callout(
        doc,
        "adaptive_evaluation_updated.docx",
        """Components covered:
  eval/classification_accuracy.py — query intent classification accuracy
  eval/readability_analysis.py    — Flesch-Kincaid + Coleman-Liau per persona
  eval/format_compliance.py       — persona format rule adherence
  eval/adaptive_vs_generic.py     — composite Adaptive Advantage Score

Core claim:
  Persona adaptation produces measurably better outputs than a generic baseline.
  Headline metric: Adaptive Advantage Score (composite of compliance_delta +
  readability_fit + length_delta).

Read when: validating that persona adaptation works correctly, reporting evaluation
           results, or modifying the query classifier or persona response rules.""",
    )

    _heading(doc, "3.3 Benchmark Evaluations", level=2)
    add_callout(
        doc,
        "benchmark_evaluations_updated.docx",
        """Components covered:
  eval/model_comparison.py     — multi-model response quality benchmarks
  eval/hybrid_comparison.py    — hybrid vs. semantic-only retrieval comparison
  eval/persona_evaluation.py   — per-persona faithfulness and format benchmarks
  eval/latency_measurement.py  — retrieval and generation latency against SLAs

Measures raw system capability independent of adaptive logic.
Key metrics:
  - Retrieval diversity score: hybrid (+0.243 vs. semantic-only)
  - Latency SLAs: retrieval < 500 ms · generation < 8000 ms  (src/config.py)
  - Persona faithfulness gap: cross-persona consistency of grounded responses

Read when: selecting embedding models, evaluating retrieval strategies, measuring
           latency compliance, or comparing model providers.""",
    )

    _heading(doc, "3.4 Hallucination Mitigation", level=2)
    add_callout(
        doc,
        "hallucination_mitigation_updated.docx",
        """Components covered:
  src/prompts.py      — SYSTEM_PROMPT (forbids fabrication)
  src/prompts.py      — _GROUNDING_INSTRUCTION + [Source N] citation tags
                        (enforce context-only answers)
  src/faithfulness.py — FaithfulnessChecker (cosine similarity scoring)
  src/config.py       — FAITHFULNESS_BLOCK_THRESHOLD = 0.25

Three-layer defense:
  Layer 1 — SYSTEM_PROMPT:           instructs the model to never fabricate
  Layer 2 — _GROUNDING_INSTRUCTION:  user-role grounding + source citation format
  Layer 3 — FaithfulnessChecker:     post-generation cosine similarity scoring
                                     warn  threshold: 0.45  (src/faithfulness.py:17)
                                     sentence flag:   0.35  (src/faithfulness.py:31)
                                     block threshold: 0.25  (src/config.py:203)

Important: The block threshold (0.25) is defined in config but the generation
path does not enforce it as a hard gate — this is a documented known gap.

Read when: conducting clinical governance review, modifying safety controls,
           or assessing regulatory compliance.""",
    )


def build_section_4_parameters(doc: Document) -> None:
    """Section 4 — Quick-Reference Parameter Table."""
    _heading(doc, "4. Quick-Reference Parameter Table")

    doc.add_paragraph(
        "All key system parameters, their configured values, the pipeline stage they "
        "affect, and the source file where they are defined. "
        "Line numbers reference the codebase at version 1.0 (2026-02-23)."
    )

    add_teal_table(
        doc,
        ["Parameter", "Value", "Stage", "Description", "Config File"],
        [
            ["chunk_size", "800", "Chunking", "Max chars per chunk", "src/config.py:15"],
            ["chunk_overlap", "150", "Chunking", "Shared chars between chunks", "src/config.py:16"],
            ["embedding_model", "S-PubMedBert-MS-MARCO", "Embedding", "Recommended medical model (768-dim, 512 seq)", "src/config.py:78-85"],
            ["embedding_dimensions", "768", "Embedding", "Vector dimensionality", "src/config.py"],
            ["top_k", "5", "Retrieval", "Chunks returned per retriever + final", "src/config.py:23"],
            ["rrf_k", "60", "Retrieval", "RRF smoothing constant", "src/config.py:24"],
            ["semantic_weight", "0.6", "Retrieval", "Dense retrieval weight in fusion", "src/config.py:25"],
            ["bm25_weight", "0.4", "Retrieval", "BM25 weight in fusion", "src/config.py:26"],
            ["hnsw_M", "16", "Indexing", "HNSW bidirectional links", "src/config.py:46"],
            ["hnsw_construction_ef", "200", "Indexing", "HNSW build-time search depth", "src/config.py:47"],
            ["hnsw_search_ef", "100", "Indexing", "HNSW query-time search depth", "src/config.py:48"],
            ["faithfulness_warning", "0.45", "Mitigation", "Overall score warning threshold", "src/faithfulness.py:17"],
            ["faithfulness_sentence", "0.35", "Mitigation", "Per-sentence flag threshold", "src/faithfulness.py:31"],
            ["faithfulness_block", "0.25", "Mitigation", "Hard block threshold (config)", "src/config.py:203"],
            ["llm_temperature", "0 (config) / 0.1 (method)", "Generation", "Sampling temperature", "src/config.py:182, src/llm/base.py"],
            ["llm_max_tokens", "2048", "Generation", "Response length cap", "src/llm/base.py"],
            ["llm_model_cloud", "gpt-4o-mini", "Generation", "Default cloud model", "src/config.py:180"],
            ["llm_model_local", "llama3.1:8b", "Generation", "Default local model", "src/config.py:181"],
            ["retrieval_sla_ms", "500", "Latency", "Retrieval SLA target", "src/config.py:212"],
            ["generation_sla_ms", "8000", "Latency", "Generation SLA target", "src/config.py:213"],
        ],
        col_widths=[1.5, 1.4, 0.9, 2.1, 1.9],
    )


def build_section_5_gaps(doc: Document) -> None:
    """Section 5 — Known System Gaps with red left-border paragraphs."""
    _heading(doc, "5. Known System Gaps")

    doc.add_paragraph(
        "The following gaps are confirmed from codebase inspection. Each entry states "
        "the risk and identifies which specialist document contains the full mitigation "
        "analysis. These gaps do not represent defects in the current scope but rather "
        "boundaries of the system's current capability that are important for clinical "
        "governance reviewers to understand."
    )

    gaps = [
        (
            "NLI entailment checking",
            "Cosine similarity cannot detect logically contradictory but lexically similar text. "
            "A response that rephrases context with inverted meaning may score highly faithful "
            "while being factually wrong.",
            "hallucination_mitigation_updated.docx",
        ),
        (
            "Automatic response blocking",
            "The faithfulness_block threshold (0.25) is defined in src/config.py:203 but is "
            "not enforced as a hard gate in the generation path. Low-faithfulness responses "
            "are flagged but not suppressed.",
            "hallucination_mitigation_updated.docx",
        ),
        (
            "Self-consistency sampling",
            "A single LLM pass cannot detect stochastic hallucination variance. Running multiple "
            "samples and comparing outputs would surface cases where the model gives contradictory "
            "answers to the same question.",
            "hallucination_mitigation_updated.docx",
        ),
        (
            "No ground-truth retrieval metrics",
            "Precision, recall, and MRR cannot be computed without manually labelled relevance "
            "judgements for each query. Current retrieval evaluation relies on diversity score "
            "and latency only.",
            "benchmark_evaluations_updated.docx",
        ),
        (
            "Regex-only query classification",
            "Intent detection in src/query_classifier.py uses regex pattern matching. It will "
            "fail on novel phrasings, sarcasm, negation, or multi-intent queries not covered "
            "by the 45-query classification dataset.",
            "adaptive_evaluation_updated.docx",
        ),
        (
            "Faithfulness thresholds uncalibrated",
            "warn@0.45 and block@0.25 are manually set with no user-study validation. There is "
            "no evidence that these thresholds correspond to clinically meaningful faithfulness "
            "levels.",
            "hallucination_mitigation_updated.docx",
        ),
        (
            "No human preference data",
            "All evaluations are fully automated. There is no patient-safety signal, clinician "
            "satisfaction survey, or expert review of generated responses to validate that "
            "automated metrics correlate with real-world quality.",
            "benchmark_evaluations_updated.docx",
        ),
        (
            "Classification dataset limited to 45 queries",
            "The query classification accuracy evaluation uses 45 hand-crafted test queries. "
            "Edge cases, domain-specific phrasings, and adversarial inputs are likely "
            "underrepresented, making accuracy estimates optimistic.",
            "adaptive_evaluation_updated.docx",
        ),
    ]

    for gap_name, risk_text, doc_ref in gaps:
        p = doc.add_paragraph()
        _add_left_border(p, RED_BORDER_HEX, 18)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)

        r_title = p.add_run(f"⚠️  {gap_name}  — ")
        r_title.font.bold = True
        r_title.font.size = Pt(10)

        r_body = p.add_run(f"{risk_text}  ")
        r_body.font.size = Pt(10)

        r_ref = p.add_run(f"[documented in: {doc_ref}]")
        r_ref.font.size = Pt(9.5)
        r_ref.font.italic = True
        r_ref.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()


def build_section_6_usage(doc: Document) -> None:
    """Section 6 — How to Use This Documentation Suite."""
    _heading(doc, "6. How to Use This Documentation Suite")

    doc.add_paragraph(
        "Different readers have different entry points depending on their role. "
        "An ML engineer or developer joining the project should read the RAG Pipeline "
        "Architecture document first — it explains the full seven-stage pipeline, maps "
        "canonical brief names to actual code identifiers, and provides the complete "
        "configuration reference. After that, the Benchmark Evaluations document shows "
        "measured performance against SLA targets, which is essential context before "
        "making any changes to retrieval or embedding configuration."
    )

    doc.add_paragraph(
        "A clinical governance reviewer or safety officer should start with the "
        "Hallucination Mitigation document, which traces the three-layer defense "
        "(SYSTEM_PROMPT, grounding instruction with citation tags, and "
        "FaithfulnessChecker) to specific lines of code, and is candid about the gaps "
        "in the current implementation — particularly the uncalibrated thresholds and "
        "the absence of automatic blocking. The Known System Gaps section of this "
        "index document (Section 5) provides a consolidated view across all four "
        "documents. A researcher validating that persona adaptation works correctly "
        "should read the Adaptive Evaluation Suite document, which defines the "
        "Adaptive Advantage Score and reports per-eval results."
    )

    doc.add_paragraph(
        "The four specialist documents cross-reference each other in several important "
        "ways. The Hallucination Mitigation document references the SYSTEM_PROMPT, "
        "which is also fully documented (with verbatim text) in the RAG Pipeline "
        "Architecture document — readers should consult both if they are modifying the "
        "system prompt. The Benchmark Evaluations share embedding model and HNSW "
        "configuration with the RAG Pipeline document; any change to these parameters "
        "requires updating both documents. The Adaptive Evaluation Suite references "
        "the five persona definitions, which are also tabulated in the RAG Pipeline "
        "document's persona token budget section."
    )

    doc.add_paragraph(
        "When the system changes, apply updates to the relevant documents as follows: "
        "changing the embedding model requires updating RAG Pipeline Architecture "
        "(Section 4, embedding parameters) and Benchmark Evaluations (model comparison "
        "results); adding a new evaluation requires updating Adaptive Evaluation Suite "
        "or Benchmark Evaluations and this index (Section 3 and the Document Map); "
        "modifying the system prompt or grounding instruction requires updating both "
        "Hallucination Mitigation and RAG Pipeline Architecture; adding a new persona "
        "requires updating Adaptive Evaluation Suite (persona format rules), "
        "RAG Pipeline Architecture (persona token budget table), and this index "
        "(Section 1 system overview and Section 4 parameter table). "
        "This index document should always be regenerated last, after the specialist "
        "documents are finalised, by running: python docs/rag-system/generate_index.py"
    )


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def main() -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    _add_page_numbers(doc)

    add_cover(doc)
    add_toc(doc)

    build_section_1_overview(doc)
    build_section_2_document_map(doc)
    build_section_3_summaries(doc)
    build_section_4_parameters(doc)
    build_section_5_gaps(doc)
    build_section_6_usage(doc)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
