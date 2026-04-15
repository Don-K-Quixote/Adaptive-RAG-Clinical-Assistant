"""
generate_rag_pipeline.py

Generates docs/rag-system/rag_pipeline.docx — a full technical reference for the
Adaptive RAG Clinical Assistant pipeline.

Run:
    python docs/rag-system/generate_rag_pipeline.py

Requirements:
    pip install python-docx
"""

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Colour constants
# ---------------------------------------------------------------------------
NAVY = RGBColor(0x11, 0x22, 0x40)      # #112240  title
TEAL = RGBColor(0x0E, 0xA5, 0xC9)      # #0EA5C9  subtitle / table header
TEAL_HEX = "0EA5C9"
CALLOUT_BG_HEX = "EBF8FF"
DIAGRAM_BG_HEX = "F3F4F6"
RED_BORDER_HEX = "DC2626"
HEADER_TEXT_COLOR = RGBColor(0xFF, 0xFF, 0xFF)  # white header text

OUTPUT_PATH = Path(__file__).parent / "rag_pipeline.docx"


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def _shade_cell(cell, hex_color: str) -> None:
    """Fill a table cell with a solid background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _shade_para(para, hex_color: str) -> None:
    """Apply a solid paragraph background (shading) to simulate a coloured box."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _add_left_border(para, hex_color: str, width_pt: int = 18) -> None:
    """Add a thick left border to a paragraph (gap / warning box style)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(width_pt))   # eighths of a point
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)


def _add_page_numbers(doc: Document) -> None:
    """Insert 'Page X of Y' into the document footer via Word field codes."""
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _field(instr: str):
        run = para.add_run()
        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar_begin)

        run2 = para.add_run()
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = f" {instr} "
        run2._r.append(instrText)

        run3 = para.add_run()
        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")
        run3._r.append(fldChar_end)

    para.add_run("Page ")
    _field("PAGE")
    para.add_run(" of ")
    _field("NUMPAGES")


# ---------------------------------------------------------------------------
# High-level content helpers
# ---------------------------------------------------------------------------

def add_cover(doc: Document) -> None:
    """Add a styled cover page."""
    doc.add_paragraph()  # top margin spacer
    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("RAG Pipeline Architecture")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = NAVY

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_para.add_run("Clinical RAG Assistant — Technical Documentation")
    run.font.size = Pt(14)
    run.font.color.rgb = TEAL

    doc.add_paragraph()

    meta_lines = [
        ("System:", "Adaptive RAG Clinical Assistant"),
        ("Version:", "1.0"),
        ("Date:", "2026-02-19"),
        ("Classification:", "Internal Technical Reference"),
    ]
    for label, value in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label + "  ")
        r1.font.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(value)
        r2.font.size = Pt(11)

    doc.add_page_break()


def add_toc(doc: Document) -> None:
    """Add a Table of Contents placeholder."""
    h = doc.add_heading("Table of Contents", level=1)
    h.runs[0].font.color.rgb = NAVY

    note = doc.add_paragraph(
        "Right-click this area in Microsoft Word and select 'Update Field' "
        "to generate the Table of Contents automatically."
    )
    note.runs[0].font.italic = True
    note.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)  # grey
    note.runs[0].font.size = Pt(10)

    doc.add_page_break()


def _heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level == 1 else TEAL


def add_teal_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    col_widths: list[float] | None = None,
) -> None:
    """
    Render a table with a teal header row and alternating body rows.

    col_widths: list of Inches widths; if None, Word auto-sizes.
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr_cells[i].text = text
        _shade_cell(hdr_cells[i], TEAL_HEX)
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = HEADER_TEXT_COLOR
        run.font.size = Pt(9.5)

    # Body rows — alternate light grey
    ALT_BG = "F9FAFB"
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        bg = ALT_BG if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, text in enumerate(row_data):
            row_cells[c_idx].text = text
            _shade_cell(row_cells[c_idx], bg)
            run = row_cells[c_idx].paragraphs[0].runs[0]
            run.font.size = Pt(9.5)

    # Apply column widths if provided
    if col_widths:
        for r_idx in range(len(table.rows)):
            for c_idx, width in enumerate(col_widths):
                table.rows[r_idx].cells[c_idx].width = Inches(width)

    doc.add_paragraph()  # spacing after table


def add_callout(doc: Document, label: str, text: str) -> None:
    """
    Add a verbatim callout box (light-blue background, Courier New).

    label: displayed as bold heading above the box.
    text:  monospaced verbatim content.
    """
    lbl = doc.add_paragraph()
    r = lbl.add_run(label)
    r.font.bold = True
    r.font.size = Pt(10)

    for line in text.splitlines():
        p = doc.add_paragraph(line if line.strip() else " ")
        _shade_para(p, CALLOUT_BG_HEX)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(8.5)

    doc.add_paragraph()  # spacing


def add_gap_box(doc: Document, title: str, items: list[str]) -> None:
    """Add a red-bordered warning box listing pipeline gaps."""
    title_p = doc.add_paragraph()
    _add_left_border(title_p, RED_BORDER_HEX, 24)
    title_p.paragraph_format.left_indent = Inches(0.25)
    r = title_p.add_run(title)
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

    for item in items:
        p = doc.add_paragraph()
        _add_left_border(p, RED_BORDER_HEX, 18)
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(f"• {item}")
        r.font.size = Pt(10)

    doc.add_paragraph()


def add_ascii(doc: Document, caption: str, lines: list[str]) -> None:
    """Render an ASCII diagram in a grey code block with a caption."""
    cap = doc.add_paragraph()
    r = cap.add_run(f"Figure: {caption}")
    r.font.bold = True
    r.font.italic = True
    r.font.size = Pt(10)

    for line in lines:
        p = doc.add_paragraph(line if line else " ")
        _shade_para(p, DIAGRAM_BG_HEX)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.15)
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(8)

    doc.add_paragraph()


# mmdc may not be on PATH in all shells on Windows; resolve the full path via npm prefix.
_NPM_PREFIX = subprocess.run(
    ["node", "-e", "process.stdout.write(require('child_process').execSync('npm config get prefix').toString().trim())"],
    capture_output=True,
    text=True,
).stdout.strip()
_MMDC_CMD = str(Path(_NPM_PREFIX) / "mmdc.cmd") if _NPM_PREFIX else "mmdc"


def _try_mermaid(mmd: str, out_path: Path, mmd_source_path: Path | None = None) -> bool:
    """Attempt to render a Mermaid diagram to PNG. Returns True on success.

    If mmd_source_path is provided and exists, it is used as the .mmd input
    (single source of truth). Otherwise the inline mmd string is written to
    a temporary file that is cleaned up after rendering.
    """
    if mmd_source_path and mmd_source_path.exists():
        mmd_file = mmd_source_path
        temp_created = False
    else:
        mmd_file = out_path.with_suffix(".mmd")
        mmd_file.write_text(mmd, encoding="utf-8")
        temp_created = True

    try:
        result = subprocess.run(
            [
                _MMDC_CMD,
                "-i", str(mmd_file),
                "-o", str(out_path),
                "--scale", "2",
                "--backgroundColor", "white",
            ],
            capture_output=True,
            timeout=60,
        )
        return result.returncode == 0 and out_path.exists()
    except (FileNotFoundError, subprocess.TimeoutExpired, OSError):
        return False
    finally:
        if temp_created and mmd_file.exists():
            mmd_file.unlink()


def add_diagram(
    doc: Document,
    caption: str,
    mmd: str,
    ascii_lines: list[str],
    png_name: str,
    mmd_source_path: Path | None = None,
) -> None:
    """
    Attempt Mermaid PNG render; fall back to ASCII art if mmdc is unavailable.

    mmd_source_path: if supplied and the file exists, it is used as the source
    .mmd file instead of the inline mmd string. Keeps .mmd files as the single
    source of truth for editable diagrams.
    """
    png_path = OUTPUT_PATH.parent / png_name
    if _try_mermaid(mmd, png_path, mmd_source_path):
        cap = doc.add_paragraph()
        r = cap.add_run(f"Figure: {caption}")
        r.font.bold = True
        r.font.italic = True
        r.font.size = Pt(10)
        doc.add_picture(str(png_path), width=Inches(6.0))
        doc.add_paragraph()
    else:
        add_ascii(doc, caption, ascii_lines)


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def build_section_1_overview(doc: Document) -> None:
    _heading(doc, "1. Pipeline Overview")

    doc.add_paragraph(
        "The Adaptive RAG Clinical Assistant implements a two-phase architecture: "
        "an offline indexing phase that ingests and indexes clinical trial documents, "
        "and an online query phase that retrieves relevant context, augments it with "
        "persona-aware prompting, and generates grounded responses. "
        "The system supports five persona classes (NOVICE, INTERMEDIATE, EXPERT, "
        "REGULATORY, EXECUTIVE) and enforces a post-generation faithfulness check "
        "to reduce hallucination risk."
    )

    # Figure 1 — End-to-End Pipeline
    mmd = """flowchart LR
  subgraph OFFLINE
    A["PDF Document"] -->|"pdfplumber / OCR"| B["DocumentIngester\\nsrc/ingestion.py:88"]
    B --> C["RecursiveCharacterTextSplitter\\nchunk_size=800, overlap=150"]
    C --> D["create_embedder()\\nS-PubMedBert-MS-MARCO 768-dim"]
    D --> E[("ChromaDB\\nVector Store\\nhnsw:space=cosine M=16")]
    C --> F[("BM25Retriever\\nSparse Index")]
  end
  subgraph ONLINE
    G["User Query"] --> H["detect_user_type()\\nget_response_config()"]
    G --> I["QueryClassifier\\n9 QueryTypes"]
    G --> J["Query Embedding"]
    J --> K["ChromaDB\\nsimilarity_search k=10"]
    G --> L["BM25Retriever\\ninvoke k=10"]
    E --> K
    F --> L
    K --> M["ReciprocalRankFusion\\n1/(60+rank)"]
    L --> M
    M --> ST{"score_threshold\\n(optional gate)"}
    ST --> N["HybridRetriever\\nTop-5 Chunks"]
    N --> O["build_adaptive_prompt()\\nResponseStyler"]
    H --> O
    I --> O
    O --> P["LLMFactory\\nOpenAI / Ollama\\ngenerate() / generate_stream()\\ntop_p optional"]
    P --> Q["FaithfulnessChecker\\nscore >= 0.45"]
    Q --> R["Response -> Streamlit UI"]
  end"""

    ascii_lines = [
        "OFFLINE PHASE",
        "──────────────────────────────────────────────────────────────────",
        " PDF Document",
        "    │ pdfplumber / OCR (Surya or OpenAI Vision)",
        "    ▼",
        " DocumentIngester  (src/ingestion.py:88)",
        "    │ TEXT_NATIVE pages: pdfplumber",
        "    │ NEEDS_OCR pages:   OCR provider → eval/ocr_retrieval_quality.py",
        "    ▼",
        " RecursiveCharacterTextSplitter  (src/ingestion.py:115)",
        "    │ chunk_size=800  chunk_overlap=150",
        "    ▼",
        " Document Chunks + Metadata",
        "    │ (page, doc_name, classification, ocr_provider, chunk_id)",
        "    ├─────────────────────────┐",
        "    ▼                         ▼",
        " create_embedder()          BM25Retriever.from_documents()",
        " HuggingFaceEmbeddings      (app.py inline)",
        " S-PubMedBert-MS-MARCO",
        "    ▼",
        " Chroma.from_documents(collection_metadata=HNSW_COLLECTION_METADATA)",
        " hnsw:space=cosine  hnsw:M=16  construction_ef=200  search_ef=100",
        " (app.py, src/config.py:HNSW_COLLECTION_METADATA)",
        "",
        "ONLINE PHASE",
        "──────────────────────────────────────────────────────────────────",
        " User Query",
        "    │",
        "    ├──► detect_user_type()  (src/personas.py:111) ──► UserType",
        "    ├──► QueryClassifier.classify()  (src/query_classifier.py:162) ──► QueryType",
        "    │",
        "    ├──► Query Embedding ──► ChromaDB.similarity_search(k=10)",
        "    └──► BM25Retriever.invoke(k=10)",
        "              │",
        "              ▼",
        " ReciprocalRankFusion.fuse()  (src/retrieval.py:90)",
        " score(d) = Σ 1/(60 + rank_i(d))",
        "              │",
        "              ▼",
        " [score_threshold gate]  filter rrf_score < threshold  (optional)",
        " HybridRetriever.score_threshold  (src/retrieval.py, src/config.py)",
        "              │",
        "              ▼",
        " HybridRetriever  top-5 Documents",
        "              │",
        "    ┌─────────┴──────────┐",
        "    │   get_response_config()  (src/personas.py:167)",
        "    │   ResponseConfig (detail_level, max_length, tables, ...)",
        "    └─────────┬──────────┘",
        "              ▼",
        " build_adaptive_prompt()  (src/prompts.py:258)",
        " ResponseStyler.generate_prompt()",
        " [UserTypeInstr + QueryTypeInstr + CONTEXT + GROUNDING + QUESTION]",
        "              │",
        "              ▼",
        " LLMFactory.create()  (src/llm/factory.py:46)",
        " OpenAIProvider or OllamaProvider",
        " generate(prompt, system_prompt, temperature=0, top_p=None)",
        " generate_stream() → Iterator[str]  available on both providers",
        "              │",
        "              ▼",
        " FaithfulnessChecker.check()  (src/faithfulness.py:66)",
        " warning if score < 0.45",
        "              │",
        "              ▼",
        " Response + Source Citations → Streamlit UI",
    ]

    add_diagram(
        doc,
        "End-to-End Pipeline Flow",
        mmd,
        ascii_lines,
        "fig1_pipeline_overview.png",
        mmd_source_path=Path(__file__).parent / "fig1_pipeline_overview.mmd",
    )


def build_section_2_ingestion(doc: Document) -> None:
    _heading(doc, "2. Data Ingestion")

    doc.add_paragraph(
        "Document ingestion is handled by DocumentIngester (src/ingestion.py:88). "
        "On loading, each page is classified as TEXT_NATIVE if pdfplumber extracts "
        "at least 50 characters of text, or NEEDS_OCR otherwise. "
        "TEXT_NATIVE pages are processed directly by pdfplumber; NEEDS_OCR pages "
        "are rasterised at 200 DPI and routed to either the Surya or OpenAI Vision "
        "OCR provider."
    )

    note = doc.add_paragraph(
        "Note on canonical names: The project brief refers to this component as "
        "DocumentLoader. The actual class name in the codebase is DocumentIngester "
        "(src/ingestion.py:88)."
    )
    note.runs[0].font.italic = True
    note.runs[0].font.size = Pt(9.5)
    note.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    _heading(doc, "2.1 Document Schema", level=2)
    add_teal_table(
        doc,
        ["Field", "Type", "Description", "Source"],
        [
            ["page", "int", "1-based page number", "src/ingestion.py:199"],
            ["doc_name", "str", "Source filename", "src/ingestion.py:199"],
            ["classification", "str", '"text_native" or "needs_ocr"', "src/ingestion.py:43"],
            ["ocr_provider", "str", '"pdfplumber", "surya", or "openai_vision"', "src/ingestion.py:199"],
            ["chunk_id", "str", "MD5 hash of chunk content", "src/ingestion.py:196"],
            ["page_content", "str", "Extracted chunk text", "LangChain Document"],
        ],
        col_widths=[1.3, 0.7, 2.8, 2.0],
    )

    _heading(doc, "2.2 Ingestion Parameters", level=2)
    add_teal_table(
        doc,
        ["Parameter", "Default", "Description", "Configured In"],
        [
            ["ocr_text_threshold", "50 chars", "Min text to classify page as TEXT_NATIVE", "src/ingestion.py:37"],
            ["raster_dpi", "200", "Resolution for page rasterisation", "src/ingestion.py:40"],
            ["supported_formats", "PDF only", "File formats accepted", "src/ingestion.py"],
            ["ocr_providers", "surya, openai_vision", "Available OCR backends", "src/ocr/"],
        ],
        col_widths=[1.7, 1.2, 2.9, 1.9],
    )


def build_section_3_chunking(doc: Document) -> None:
    _heading(doc, "3. Chunking Strategy")

    doc.add_paragraph(
        "Extracted text is split using LangChain's RecursiveCharacterTextSplitter "
        "(src/ingestion.py:115) with a target chunk size of 800 characters and a "
        "150-character overlap between adjacent chunks. The splitter attempts "
        "separators in priority order — paragraph breaks (\\n\\n), line breaks (\\n), "
        'sentence endings (". "), spaces (" "), and finally splitting at any character. '
        "This hierarchy preserves clinical trial sentence structure wherever possible. "
        "Each chunk receives a deterministic chunk_id computed as the MD5 hash of its "
        "page_content, enabling deduplication during re-ingestion."
    )

    note = doc.add_paragraph(
        "Note on canonical names: The project brief calls this component "
        "RecursiveSemanticSplitter. The actual implementation uses LangChain's "
        "RecursiveCharacterTextSplitter — character and separator-based, not "
        "semantic-similarity-based."
    )
    note.runs[0].font.italic = True
    note.runs[0].font.size = Pt(9.5)
    note.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    _chunking_mmd = """flowchart TD
  S["Source Document (Pages 1-N)\\nClinical trial protocol text"]
  S -->|"RecursiveCharacterTextSplitter\\nsrc/ingestion.py:115\\nseparators: paragraph break, line break, sentence, space"| C1
  C1["Chunk 1\\n~800 chars\\npage=N, chunk_id=MD5(content)"]
  C2["Chunk 2\\n~800 chars\\npage=N, chunk_id=MD5(content)"]
  C3["Chunk 3\\n~800 chars\\npage=N, chunk_id=MD5(content)"]
  CN["... Chunk N"]
  C1 -.->|"150-char overlap"| C2
  C2 -.->|"150-char overlap"| C3
  C3 -.->|"150-char overlap"| CN"""

    _chunking_ascii = [
        "  Source Document (pages 1-N)",
        "  +----------------------------------------------------------+",
        "  | ... clinical trial protocol text ...                     |",
        "  +----------------------------------------------------------+",
        "                 |  RecursiveCharacterTextSplitter",
        "                 |  separator priority: \\n\\n -> \\n -> \". \" -> \" \" -> \"\"",
        "                 v",
        "  +------------------+",
        "  |   Chunk 1        |  ~800 chars   page=3, chunk_id=MD5(content)",
        "  |   ...content...  |",
        "  +------------------+",
        "          +-- 150 char overlap --+",
        "          +-----------------------------+",
        "          |         Chunk 2             |  ~800 chars",
        "          |         ...content...       |",
        "          +-----------------------------+",
        "                    +-- 150 char overlap --+",
        "                    +----------------------------+",
        "                    |         Chunk 3            |  ~800 chars",
        "                    |         ...content...      |",
        "                    +----------------------------+",
    ]

    add_diagram(
        doc,
        "Chunking and Overlap",
        _chunking_mmd,
        _chunking_ascii,
        "fig2_chunking.png",
        mmd_source_path=Path(__file__).parent / "fig2_chunking.mmd",
    )

    _heading(doc, "3.1 Chunking Parameters", level=2)
    add_teal_table(
        doc,
        ["Parameter", "Default", "Description", "Configured In"],
        [
            ["chunk_size", "800", "Target character count per chunk", "src/config.py:15"],
            ["chunk_overlap", "150", "Character overlap between adjacent chunks", "src/config.py:16"],
            ["separators", '[\"\\\\n\\\\n\",\"\\\\n\",\". \",\" \",\"\"]', "Priority-ordered split points", "src/ingestion.py:115"],
            ["chunk_id algorithm", "MD5", "Hash of page_content for deduplication", "src/ingestion.py:196"],
        ],
        col_widths=[1.7, 1.6, 2.7, 1.7],
    )


def build_section_4_embedding(doc: Document) -> None:
    _heading(doc, "4. Embedding")

    doc.add_paragraph(
        "Embeddings are produced by create_embedder() (src/embeddings.py:57), "
        "which returns a HuggingFaceEmbeddings instance. The default model is "
        "pritamdeka/S-PubMedBert-MS-MARCO, a medical domain model producing "
        "768-dimensional vectors. "
        "The function implements a three-tier fallback chain: (1) the configured "
        "primary model, (2) a configured fallback model, and (3) all-MiniLM-L6-v2 "
        "as a last resort. "
        "Once created, the embedder instance is cached in Streamlit session state and "
        "shared with FaithfulnessChecker to avoid loading the model twice."
    )

    note = doc.add_paragraph(
        "Note on canonical names: The project brief refers to this component as "
        "EmbeddingService. The actual implementation is the function create_embedder() "
        "in src/embeddings.py:57 — there is no EmbeddingService class."
    )
    note.runs[0].font.italic = True
    note.runs[0].font.size = Pt(9.5)
    note.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    _heading(doc, "4.1 Available Embedding Models", level=2)
    add_teal_table(
        doc,
        ["Model Key", "Type", "Dimensions", "Recommended For"],
        [
            ["S-PubMedBert-MS-MARCO (DEFAULT)", "medical", "768", "Clinical trials, IRC charters"],
            ["BioSimCSE-BioLinkBERT", "medical", "768", "Biomedical literature"],
            ["BioBERT", "medical", "768", "Biomedical NLP"],
            ["all-mpnet-base-v2", "general", "768", "Mixed content"],
            ["all-MiniLM-L6-v2", "general", "384", "Rapid prototyping / fallback"],
            ["bert-tiny-mnli", "lightweight", "128", "Testing, low-resource"],
        ],
        col_widths=[2.5, 1.1, 1.1, 2.1],
    )

    _heading(doc, "4.2 Embedding Parameters", level=2)
    add_teal_table(
        doc,
        ["Parameter", "Default", "Description", "Configured In"],
        [
            ["model_key", "S-PubMedBert-MS-MARCO", "Primary embedding model", "src/config.py:35"],
            ["dimensions", "768", "Output vector dimensionality", "src/config.py"],
            ["fallback_model", "all-MiniLM-L6-v2", "Last-resort fallback", "src/embeddings.py:100"],
            ["reuse_across_faithfulness", "True", "Embedder shared with FaithfulnessChecker", "app.py session state"],
        ],
        col_widths=[2.0, 1.7, 2.5, 1.5],
    )


def build_section_5_index(doc: Document) -> None:
    _heading(doc, "5. Index Construction")

    doc.add_paragraph(
        "Two indexes are built inline in app.py during document upload. "
        "Chroma.from_documents() constructs the dense vector store, storing all "
        "chunk embeddings with their metadata. "
        "BM25Retriever.from_documents() (langchain-community) builds the sparse "
        "TF-IDF lexical index. Both indexes are stored in Streamlit session state "
        "and rebuilt on each new document upload."
    )

    critical_note = doc.add_paragraph()
    r = critical_note.add_run(
        "IMPORTANT — Discrepancy with project brief: "
        "The system brief and canonical component list describe FAISSIndex as the "
        "vector store. The actual implementation uses ChromaDB (chromadb==1.5.0). "
        "FAISS does not appear in requirements.txt and is not imported anywhere in "
        "the codebase. All references in this document use ChromaDB, which reflects "
        "reality."
    )
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
    r.font.size = Pt(10)

    doc.add_paragraph()

    resolved_note = doc.add_paragraph()
    rn = resolved_note.add_run(
        "RESOLVED — HNSW Configuration: "
        "ChromaDB HNSW indexing parameters are now explicitly configured via "
        "HNSW_COLLECTION_METADATA (src/config.py) and passed to every "
        "Chroma.from_documents() call in app.py and all eval/ modules. "
        "Parameters: hnsw:space=cosine, hnsw:M=16, hnsw:construction_ef=200, "
        "hnsw:search_ef=100."
    )
    rn.font.bold = True
    rn.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)  # green
    rn.font.size = Pt(10)

    doc.add_paragraph()

    mmd = """flowchart LR
  Q["User Query"] --> E1["Embed\\ncreate_embedder()"]
  Q --> T1["Tokenise\\n(BM25 internal)"]
  E1 --> D["ChromaDB\\nsimilarity_search(k=10)\\nhnsw:space=cosine\\nM=16 ef_search=100"]
  T1 --> B["BM25Retriever\\ninvoke(k=10)"]
  D --> R["ReciprocalRankFusion\\nSum 1/(60+rank_i)"]
  B --> R
  R --> ST{"score_threshold\\noptional gate"}
  ST --> K["Top-5 Chunks"]"""

    ascii_lines = [
        "  User Query",
        "      │",
        "      ├──────────────────────────────┐",
        "      │                              │",
        "      ▼                              ▼",
        "  Query Embedding               BM25 Tokenise",
        "  create_embedder()             (internal to BM25Retriever)",
        "      │                              │",
        "      ▼                              ▼",
        "  ChromaDB                      BM25Retriever",
        "  similarity_search(k=10)       invoke(k=10)",
        "  hnsw:space=cosine             TF-IDF lexical match",
        "  hnsw:M=16  ef_construction=200",
        "  hnsw:search_ef=100",
        "      │                              │",
        "      └──────────────┬───────────────┘",
        "                     ▼",
        "          ReciprocalRankFusion.fuse()",
        "          score(d) = Σ 1/(60 + rank_i(d))",
        "          (src/retrieval.py:90, k=60)",
        "                     │",
        "                     ▼",
        "          [score_threshold gate]  optional",
        "          discard results where rrf_score < threshold",
        "          HybridRetriever.score_threshold (default: None)",
        "                     │",
        "                     ▼",
        "              Top-5 Chunks (deduplicated)",
    ]

    add_diagram(
        doc,
        "Index Construction and Retrieval Flow",
        mmd,
        ascii_lines,
        "fig3_index_retrieval.png",
        mmd_source_path=Path(__file__).parent / "fig3_index_retrieval.mmd",
    )

    _heading(doc, "5.1 Index Parameters", level=2)
    add_teal_table(
        doc,
        ["Parameter", "Default", "Description", "Configured In"],
        [
            ["vector_store", "ChromaDB", "Dense index backend (NOT FAISS)", "app.py"],
            ["persist_directory", "None (in-memory)", "ChromaDB on-disk path", "app.py"],
            ["bm25_backend", "rank-bm25", "Sparse lexical index library", "requirements.txt"],
            ["metadata_fields", "page, doc_name, classification, ocr_provider, chunk_id", "Per-chunk metadata", "src/ingestion.py:199"],
            ["hnsw:space", "cosine", "Distance metric matching sentence-transformer training", "src/config.py HNSW_COLLECTION_METADATA"],
            ["hnsw:M", "16", "Bidirectional links per node (higher = more accurate)", "src/config.py HNSW_COLLECTION_METADATA"],
            ["hnsw:construction_ef", "200", "Beam width at index-build time (higher = better quality)", "src/config.py HNSW_COLLECTION_METADATA"],
            ["hnsw:search_ef", "100", "Beam width at query time (higher = better recall)", "src/config.py HNSW_COLLECTION_METADATA"],
        ],
        col_widths=[1.6, 1.6, 2.8, 1.7],
    )


def build_section_6_retrieval(doc: Document) -> None:
    _heading(doc, "6. Hybrid Retrieval and Re-ranking")

    doc.add_paragraph(
        "HybridRetriever (src/retrieval.py) orchestrates both index queries and "
        "combines their results using Reciprocal Rank Fusion (RRF). "
        "The dense retriever queries ChromaDB for the top-10 most similar embeddings; "
        "the sparse retriever queries BM25Retriever for the top-10 lexically "
        "matching documents. The two ranked lists are merged by RRF with k=60, "
        "and the final top-5 documents are returned. Retrieval timing is captured "
        "as metadata on every query."
    )

    add_callout(
        doc,
        "RRF Formula (src/retrieval.py:12–18)",
        """RRF_score(d) = Σ  1 / (k + rank_i(d))

where:
  d        = candidate document
  k        = 60  (Cormack et al. 2009; src/config.py:24)
  rank_i   = position of d in ranked list i (1-based)
  i        = {semantic, lexical}

Documents found in both lists receive contributions from both terms.""",
    )

    _heading(doc, "6.1 Retrieval Parameters", level=2)
    add_teal_table(
        doc,
        ["Parameter", "Default", "Description", "Configured In"],
        [
            ["dense_top_k", "10 (= 2 × final_top_k)", "Candidates from ChromaDB", "src/retrieval.py:258"],
            ["sparse_top_k", "10 (= 2 × final_top_k)", "Candidates from BM25Retriever", "src/retrieval.py:258"],
            ["rrf_k", "60", "RRF smoothing constant", "src/config.py:24"],
            ["score_threshold", "None (disabled)", "Min RRF score gate; enable in sidebar. Typical range 0.007–0.016.", "src/retrieval.py HybridRetriever"],
            ["final_top_k", "5", "Documents passed to prompt builder", "src/config.py:23"],
        ],
        col_widths=[1.7, 1.8, 2.6, 1.6],
    )

    _heading(doc, "6.2 Benchmark Results", level=2)
    doc.add_paragraph(
        "Evaluated on the clinical trial document corpus "
        "(eval/hybrid_comparison.py):"
    )
    add_teal_table(
        doc,
        ["Metric", "Hybrid Retrieval", "Semantic-Only", "Delta"],
        [
            ["Average retrieval time", "52.92 ms", "49.91 ms", "+3.01 ms overhead"],
            ["Diversity score", "(higher)", "(lower)", "+0.243 improvement"],
        ],
        col_widths=[2.2, 1.7, 1.7, 2.1],
    )


def build_section_7_prompt(doc: Document) -> None:
    _heading(doc, "7. Prompt Augmentation")

    doc.add_paragraph(
        "Prompt construction is handled by build_adaptive_prompt() "
        "(src/prompts.py:258) and the ResponseStyler class (src/prompts.py:37). "
        "The function assembles a user-role message from five ordered components: "
        "a user-type instruction (5 templates keyed on UserType), a query-type "
        "instruction (9 templates keyed on QueryType), the retrieved document "
        "context with source citation tags, a grounding requirement, and the user "
        "question. The SYSTEM_PROMPT is passed separately as the system message "
        "to every LLM call."
    )

    note = doc.add_paragraph(
        "Note on canonical names: The project brief calls this component "
        "PromptBuilder. The actual implementation is the ResponseStyler class "
        "(src/prompts.py:37) plus the top-level function build_adaptive_prompt() "
        "(src/prompts.py:258)."
    )
    note.runs[0].font.italic = True
    note.runs[0].font.size = Pt(9.5)
    note.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    _prompt_mmd = """flowchart TD
  subgraph SYS["SYSTEM MESSAGE (every LLM call)"]
    SP["SYSTEM_PROMPT\\n'You are a clinical trials document assistant...'\\nsrc/prompts.py:21-27"]
  end
  subgraph USR["USER ROLE MESSAGE"]
    UT["UserType Instruction\\n5 templates: NOVICE / INTERMEDIATE / EXPERT / REGULATORY / EXECUTIVE\\nsrc/prompts.py:47-107"]
    QT["QueryType Instruction\\n9 templates: FACTUAL / PROCEDURE / COMPARATIVE / ...\\nsrc/prompts.py:110-171"]
    CTX["Context from Document\\n[Source 1: Page X, Chunk Y] ... top-5 chunks"]
    GR["Grounding Requirement\\n'Answer using ONLY the information in CONTEXT...'\\nsrc/prompts.py:30-34"]
    QN["QUESTION: user query"]
    AN["ANSWER:"]
  end
  SP --> UT --> QT --> CTX --> GR --> QN --> AN"""

    _prompt_ascii = [
        "+---------------------------------------------------------------------+",
        "|  SYSTEM MESSAGE  (passed separately to every LLM call)             |",
        "|  SYSTEM_PROMPT  .  src/prompts.py:21-27                            |",
        "|  \"You are a clinical trials document assistant.                    |",
        "|   Never fabricate information, invent citations...\"                |",
        "+---------------------------------------------------------------------+",
        "|  USER ROLE MESSAGE                                                  |",
        "+---------------------------------------------------------------------+",
        "|  USER TYPE INSTRUCTION  .  src/prompts.py:47-107                   |",
        "|  (5 templates keyed on UserType enum)                              |",
        "|  e.g. NOVICE: \"Use simple language, provide definitions...\"        |",
        "+---------------------------------------------------------------------+",
        "|  QUERY TYPE INSTRUCTION  .  src/prompts.py:110-171                 |",
        "|  (9 templates keyed on QueryType enum)                             |",
        "|  e.g. PROCEDURE: \"Provide numbered step-by-step instructions...\"  |",
        "+---------------------------------------------------------------------+",
        "|  CONTEXT FROM DOCUMENT                                             |",
        "|  [Source 1: Page 3, Chunk 12]  <chunk text>                        |",
        "|  [Source 2: Page 7, Chunk 31]  <chunk text>                        |",
        "|  ... (top-5 chunks from HybridRetriever)                           |",
        "+---------------------------------------------------------------------+",
        "|  GROUNDING REQUIREMENT  .  src/prompts.py:30-34                    |",
        "|  \"Answer using ONLY the information in the CONTEXT...\"            |",
        "+---------------------------------------------------------------------+",
        "|  QUESTION:  <user query>                                           |",
        "+---------------------------------------------------------------------+",
        "|  ANSWER:                                                           |",
        "+---------------------------------------------------------------------+",
    ]

    add_diagram(
        doc,
        "Prompt Layout (assembled per query)",
        _prompt_mmd,
        _prompt_ascii,
        "fig4_prompt_assembly.png",
        mmd_source_path=Path(__file__).parent / "fig4_prompt_assembly.mmd",
    )

    add_callout(
        doc,
        "SYSTEM_PROMPT — verbatim (src/prompts.py:21–27)",
        """You are a clinical trials document assistant. You answer questions strictly
based on the document context provided by the user. Never fabricate information,
invent citations, or draw on knowledge outside the provided context. If the
document context does not contain enough information to answer a question fully,
say so explicitly and describe what is missing rather than guessing.""",
    )

    add_callout(
        doc,
        "_GROUNDING_INSTRUCTION — verbatim (src/prompts.py:30–34)",
        """GROUNDING REQUIREMENT:
Answer using ONLY the information in the CONTEXT FROM DOCUMENT section above.
Do not use prior knowledge, external sources, or infer facts not explicitly
stated in the context. If the context does not contain sufficient information
to fully answer the question, explicitly state what is not covered rather than
speculating.""",
    )

    add_callout(
        doc,
        "Source Citation Tag — format_source_reference() (src/utils.py:115–129)",
        """format_source_reference(doc, index=1) -> \"[Source 1: Page 3, Chunk 12]\"""",
    )

    _heading(doc, "7.1 Persona Token Budget", level=2)
    doc.add_paragraph(
        "Note on canonical names: The project brief describes 3 persona levels "
        "(Beginner/Intermediate/Expert). The actual UserType enum has 5 values: "
        "NOVICE, INTERMEDIATE, EXPERT, REGULATORY, EXECUTIVE (src/personas.py:22)."
    ).runs[0].font.italic = True

    add_teal_table(
        doc,
        ["Persona", "max_length", "detail_level", "use_tables", "include_definitions", "key_takeaway"],
        [
            ["NOVICE", "300", "low", "No", "Yes", "Yes"],
            ["INTERMEDIATE", "500", "medium", "Yes", "No", "No"],
            ["EXPERT", "1000", "high", "Yes", "No", "No"],
            ["REGULATORY", "800", "high", "Yes", "No", "No"],
            ["EXECUTIVE", "250", "low", "Yes", "No", "Yes (exec summary)"],
        ],
        col_widths=[1.3, 0.9, 1.0, 0.95, 1.5, 1.55],
    )


def build_section_8_generation(doc: Document) -> None:
    _heading(doc, "8. Response Generation")

    doc.add_paragraph(
        "LLMFactory.create() (src/llm/factory.py:46) instantiates either "
        "OpenAIProvider or OllamaProvider based on the configured provider setting. "
        "Both providers implement the abstract LLMProvider interface "
        "(src/llm/base.py), exposing a generate(prompt, system_prompt, temperature) "
        "method and returning an LLMResponse dataclass. "
        "Temperature is fixed at 0.0 for deterministic, reproducible outputs — "
        "essential for clinical document Q&A where consistent answers build user trust."
    )

    note = doc.add_paragraph(
        "Note on canonical names: The project brief calls this component "
        "LLMGenerator. The actual implementation is the LLMProvider abstract base "
        "(src/llm/base.py) with concrete OpenAIProvider and OllamaProvider classes, "
        "instantiated via LLMFactory (src/llm/factory.py:46)."
    )
    note.runs[0].font.italic = True
    note.runs[0].font.size = Pt(9.5)
    note.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    _heading(doc, "8.1 Supported Models", level=2)
    add_teal_table(
        doc,
        ["Provider", "Model ID", "Context Window", "Vision"],
        [
            ["OpenAI", "gpt-4o", "128k", "Yes"],
            ["OpenAI", "gpt-4o-mini", "128k", "Yes"],
            ["OpenAI", "gpt-4", "8192", "No"],
            ["OpenAI", "gpt-3.5-turbo", "16.3k", "No"],
            ["Ollama", "llama3.1:8b", "4096", "No"],
            ["Ollama", "biomistral-7b", "4096", "No"],
            ["Ollama", "medgemma:4b", "4096", "No"],
            ["Ollama", "mistral:7b", "4096", "No"],
        ],
        col_widths=[1.2, 1.8, 1.5, 0.8],
    )

    _heading(doc, "8.2 LLMResponse Fields (src/llm/base.py:12–20)", level=2)
    add_teal_table(
        doc,
        ["Field", "Type", "Description"],
        [
            ["content", "str", "Generated response text"],
            ["model", "str", "Model identifier used for generation"],
            ["provider", "str", "Provider name (\"openai\" or \"ollama\")"],
            ["tokens_used", "int", "Total tokens consumed (prompt + completion)"],
            ["latency_ms", "float", "End-to-end generation latency in milliseconds"],
        ],
        col_widths=[1.5, 0.9, 5.0],
    )

    _heading(doc, "8.3 Generation Parameters", level=2)
    add_teal_table(
        doc,
        ["Parameter", "Default", "Description", "Configured In"],
        [
            ["provider", "openai", 'LLM backend ("openai" or "ollama")', "app.py"],
            ["openai_model", "gpt-4o", "OpenAI model identifier", "src/config.py:102"],
            ["ollama_model", "llama3.1:8b", "Local Ollama model", "src/config.py:122"],
            ["temperature", "0.0", "Sampling temperature (deterministic)", "app.py:940"],
            ["max_tokens", "2048", "Maximum generation tokens", "src/llm/base.py"],
            ["top_p", "None (disabled)", "Nucleus sampling; enable via sidebar toggle. Overrides temperature when set.", "src/llm/base.py, openai_provider.py, ollama_provider.py"],
            ["streaming", "generate_stream()", "Token-by-token streaming via Iterator[str]", "src/llm/openai_provider.py, ollama_provider.py"],
        ],
        col_widths=[1.5, 1.5, 2.8, 1.9],
    )


def build_section_9_config(doc: Document) -> None:
    _heading(doc, "9. Consolidated Configuration Reference")

    doc.add_paragraph(
        "All configurable parameters across the pipeline, with their defaults and "
        "source locations."
    )

    add_teal_table(
        doc,
        ["Parameter", "Default", "Stage", "Configured In"],
        [
            ["ocr_text_threshold", "50 chars", "Ingestion", "src/ingestion.py:37"],
            ["raster_dpi", "200", "Ingestion", "src/ingestion.py:40"],
            ["chunk_size", "800", "Chunking", "src/config.py:15"],
            ["chunk_overlap", "150", "Chunking", "src/config.py:16"],
            ["embedding_model", "S-PubMedBert-MS-MARCO", "Embedding", "src/config.py:35"],
            ["embedding_dimensions", "768", "Embedding", "src/config.py"],
            ["vector_store", "ChromaDB", "Indexing", "app.py"],
            ["dense_candidates", "10 (2×top_k)", "Retrieval", "src/retrieval.py:258"],
            ["sparse_candidates", "10 (2×top_k)", "Retrieval", "src/retrieval.py:258"],
            ["rrf_k", "60", "Retrieval", "src/config.py:24"],
            ["final_top_k", "5", "Retrieval", "src/config.py:23"],
            ["novice_max_length", "300", "Prompting", "src/config.py:165"],
            ["intermediate_max_length", "500", "Prompting", "src/config.py:165"],
            ["expert_max_length", "1000", "Prompting", "src/config.py:165"],
            ["regulatory_max_length", "800", "Prompting", "src/config.py:165"],
            ["executive_max_length", "250", "Prompting", "src/config.py:165"],
            ["llm_temperature", "0.0", "Generation", "app.py:940"],
            ["llm_max_tokens", "2048", "Generation", "src/llm/base.py"],
            ["faithfulness_warning", "0.45", "Post-generation", "src/faithfulness.py:17"],
            ["sentence_threshold", "0.35", "Post-generation", "src/faithfulness.py:31"],
        ],
        col_widths=[2.0, 1.6, 1.3, 2.0],
    )


def build_section_10_discrepancies(doc: Document) -> None:
    _heading(doc, "10. Brief-to-Code Discrepancy Reference")

    doc.add_paragraph(
        "The system brief uses canonical component names that differ from actual "
        "code identifiers. The table below maps brief names to code reality."
    )

    add_teal_table(
        doc,
        ["Brief Canonical", "Actual Code Name", "File", "Notes"],
        [
            ["DocumentLoader", "DocumentIngester", "src/ingestion.py:88", "Class name differs"],
            ["RecursiveSemanticSplitter", "RecursiveCharacterTextSplitter", "src/ingestion.py:115", "LangChain class; character-based, not semantic"],
            ["EmbeddingService", "create_embedder()", "src/embeddings.py:57", "Function, not class"],
            ["FAISSIndex", "ChromaDB (Chroma.from_documents())", "app.py inline", "FAISS not in requirements.txt; not used"],
            ["BM25Index", "BM25Retriever.from_documents()", "app.py inline", "LangChain community; no wrapper class"],
            ["PromptBuilder", "ResponseStyler + build_adaptive_prompt()", "src/prompts.py:37, 258", "Class + top-level function"],
            ["LLMGenerator", "LLMProvider / LLMFactory", "src/llm/base.py, factory.py", "Abstract base + factory pattern"],
            ["3 levels (Beginner/Inter/Expert)", "5 levels: NOVICE/INTERMEDIATE/EXPERT/REGULATORY/EXECUTIVE", "src/personas.py:22", "UserType enum"],
        ],
        col_widths=[1.8, 2.2, 1.8, 2.0],
    )


def build_section_11_gaps(doc: Document) -> None:
    _heading(doc, "11. Gap Resolution Status")

    doc.add_paragraph(
        "The following pipeline gaps were identified during codebase review and have "
        "since been addressed. The table below maps each original finding to the "
        "resolution implemented."
    )

    add_teal_table(
        doc,
        ["Gap", "Status", "Resolution", "Source Location"],
        [
            [
                "Score threshold filtering",
                "RESOLVED",
                "score_threshold parameter added to HybridRetriever. "
                "Low-quality RRF results below the threshold are filtered in "
                "retrieve() and retrieve_with_metadata(). Exposed as a sidebar "
                "slider (0.001–0.020) in app.py. Default: None (disabled).",
                "src/retrieval.py, src/config.py (DEFAULT_SCORE_THRESHOLD), app.py",
            ],
            [
                "FAISSIndex — brief vs. code discrepancy",
                "DOCUMENTED",
                "FAISS was never implemented. ChromaDB is the actual vector store. "
                "All document references corrected to ChromaDB. No code change "
                "required — discrepancy is a naming artefact in the original brief.",
                "Sections 5, 10 of this document",
            ],
            [
                "HNSW configuration",
                "RESOLVED",
                "Explicit HNSW parameters (hnsw:space=cosine, hnsw:M=16, "
                "hnsw:construction_ef=200, hnsw:search_ef=100) added as "
                "HNSW_COLLECTION_METADATA constant and passed to every "
                "Chroma.from_documents() call in app.py and all five eval/ modules.",
                "src/config.py, app.py, eval/hybrid_comparison.py, "
                "eval/model_comparison.py, eval/latency_measurement.py, "
                "eval/persona_evaluation.py, eval/adaptive_vs_generic.py",
            ],
            [
                "Streaming generation",
                "RESOLVED",
                "generate_stream() abstract method added to LLMProvider base class "
                "(returns Iterator[str]). Implemented in OpenAIProvider using "
                "stream=True in chat.completions.create(), and in OllamaProvider "
                "using stream=True in client.chat().",
                "src/llm/base.py, src/llm/openai_provider.py, "
                "src/llm/ollama_provider.py",
            ],
            [
                "top_p / nucleus sampling",
                "RESOLVED",
                "top_p: float | None = None added to generate(), "
                "generate_with_metadata(), and generate_stream() in the base class "
                "and both providers. Passed through to the OpenAI API and Ollama "
                "options dict. Exposed as a sidebar toggle + slider (0.1–1.0) in "
                "app.py.",
                "src/llm/base.py, src/llm/openai_provider.py, "
                "src/llm/ollama_provider.py, app.py",
            ],
            [
                "OCR evaluation integration",
                "RESOLVED",
                "eval/ocr_retrieval_quality.py created. Ingests a document via "
                "DocumentIngester (applying OCR to scanned pages), then evaluates "
                "hybrid retrieval quality separately on OCR-extracted chunks and "
                "native-text chunks. Reports per-source-type counts, RRF scores, "
                "diversity, and latency. Outputs CSV and summary TXT.",
                "eval/ocr_retrieval_quality.py",
            ],
        ],
        col_widths=[1.6, 1.0, 3.2, 1.9],
    )


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def main() -> None:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # Default paragraph font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    _add_page_numbers(doc)

    add_cover(doc)
    add_toc(doc)

    build_section_1_overview(doc)
    build_section_2_ingestion(doc)
    build_section_3_chunking(doc)
    build_section_4_embedding(doc)
    build_section_5_index(doc)
    build_section_6_retrieval(doc)
    build_section_7_prompt(doc)
    build_section_8_generation(doc)
    build_section_9_config(doc)
    build_section_10_discrepancies(doc)
    build_section_11_gaps(doc)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
