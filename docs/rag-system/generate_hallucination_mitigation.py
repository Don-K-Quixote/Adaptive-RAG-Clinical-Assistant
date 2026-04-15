"""
generate_hallucination_mitigation.py

Generates docs/rag-system/hallucination_mitigation.docx — a clinical governance
reference documenting every hallucination mitigation mechanism in the
Adaptive RAG Clinical Assistant.

Run:
    python docs/rag-system/generate_hallucination_mitigation.py

Requirements:
    pip install python-docx
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ---------------------------------------------------------------------------
# Colour constants (matching app UI theme)
# ---------------------------------------------------------------------------
NAVY = RGBColor(0x11, 0x22, 0x40)       # #112240  section headers
TEAL = RGBColor(0x0E, 0xA5, 0xC9)       # #0EA5C9  subtitle / table headers
WHITE = RGBColor(0xFF, 0xFF, 0xFF)       # #FFFFFF
RED = RGBColor(0xDC, 0x26, 0x26)        # #DC2626  gaps border

TEAL_HEX = "0EA5C9"
CALLOUT_BG_HEX = "EBF8FF"              # light blue — prompt quote boxes
DIAGRAM_BG_HEX = "F3F4F6"             # light grey — ASCII diagram boxes
RED_BORDER_HEX = "DC2626"
HEADER_TEXT_HEX = "FFFFFF"

OUTPUT_PATH = Path(__file__).parent / "hallucination_mitigation.docx"

TODAY = date.today().strftime("%Y-%m-%d")


# ===========================================================================
# Low-level XML helpers
# ===========================================================================

def _shade_cell(cell, hex_color: str) -> None:
    """Fill a table cell with a solid background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _shade_para(para, hex_color: str) -> None:
    """Apply solid paragraph background shading."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _add_left_border(para, hex_color: str, width_pt: int = 24) -> None:
    """Add a thick left border to a paragraph (gap/warning box style)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(width_pt))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)


def _page_break(doc: Document) -> None:
    """Insert an explicit page break."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _add_page_numbers(doc: Document) -> None:
    """Insert 'Page X of Y' into the document footer."""
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _field(instr: str) -> None:
        run = para.add_run()
        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar_begin)

        run2 = para.add_run()
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = f" {instr} "
        run2._r.append(instrText)

        run3 = para.add_run()
        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")
        run3._r.append(fldChar_end)

    para.add_run("Page ")
    _field("PAGE")
    para.add_run(" of ")
    _field("NUMPAGES")


# ===========================================================================
# Content helpers
# ===========================================================================

def _heading1(doc: Document, text: str) -> None:
    """Add a Heading 1 styled paragraph with navy colour."""
    para = doc.add_heading(text, level=1)
    for run in para.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(16)


def _heading2(doc: Document, text: str) -> None:
    """Add a Heading 2 styled paragraph."""
    para = doc.add_heading(text, level=2)
    for run in para.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(13)


def _body(doc: Document, text: str, bold_prefix: str = "") -> None:
    """Add a normal body paragraph, with optional bold prefix."""
    para = doc.add_paragraph()
    if bold_prefix:
        r = para.add_run(bold_prefix)
        r.bold = True
    para.add_run(text)
    return para


def _mono_ref(doc: Document, text: str, indent_cm: float = 0.0) -> None:
    """Add a monospace code-reference line (grey tint)."""
    para = doc.add_paragraph()
    if indent_cm:
        para.paragraph_format.left_indent = Cm(indent_cm)
    _shade_para(para, "F1F5F9")
    run = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)


def _callout_prompt(doc: Document, label: str, text: str) -> None:
    """
    Render a verbatim prompt quote in a light-blue styled callout box.
    Used for SYSTEM_PROMPT and _GROUNDING_INSTRUCTION.
    """
    # Label line
    label_para = doc.add_paragraph()
    _shade_para(label_para, CALLOUT_BG_HEX)
    label_run = label_para.add_run(label)
    label_run.font.bold = True
    label_run.font.name = "Courier New"
    label_run.font.size = Pt(9)
    label_run.font.color.rgb = RGBColor(0x0E, 0xA5, 0xC9)

    # Content lines
    for line in text.split("\n"):
        cp = doc.add_paragraph()
        _shade_para(cp, CALLOUT_BG_HEX)
        cp.paragraph_format.left_indent = Cm(0.5)
        cr = cp.add_run(line if line else " ")
        cr.font.name = "Courier New"
        cr.font.size = Pt(9)
        cr.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)


def _ascii_diagram(doc: Document, caption: str, lines: list[str]) -> None:
    """
    Render an ASCII art diagram in a grey-shaded monospace block,
    followed by a numbered caption.
    """
    for line in lines:
        p = doc.add_paragraph()
        _shade_para(p, DIAGRAM_BG_HEX)
        p.paragraph_format.left_indent = Cm(0.3)
        r = p.add_run(line if line else " ")
        r.font.name = "Courier New"
        r.font.size = Pt(8)

    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_para.add_run(caption)
    cap_run.font.italic = True
    cap_run.font.size = Pt(9)
    cap_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


def _table_with_header(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    """
    Add a table with a teal shaded header row and alternating body rows.
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        _shade_cell(cell, TEAL_HEX)
        para = cell.paragraphs[0]
        run = para.add_run(h)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)

    # Data rows
    for r_idx, row in enumerate(rows):
        bg = "FFFFFF" if r_idx % 2 == 0 else "F8FAFC"
        tr = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row):
            cell = tr.cells[c_idx]
            _shade_cell(cell, bg)
            para = cell.paragraphs[0]
            run = para.add_run(cell_text)
            run.font.size = Pt(9)

    doc.add_paragraph()  # spacing after table


def _gap_block(doc: Document, title: str, risk: str, remedy: str) -> None:
    """
    Add a red-left-bordered gap entry for Section 7.
    """
    # Warning symbol + title
    p_title = doc.add_paragraph()
    _add_left_border(p_title, RED_BORDER_HEX, width_pt=24)
    p_title.paragraph_format.left_indent = Cm(0.4)
    r_warn = p_title.add_run("⚠️  " + title)
    r_warn.font.bold = True
    r_warn.font.size = Pt(11)
    r_warn.font.color.rgb = RED

    # Risk
    p_risk = doc.add_paragraph()
    _add_left_border(p_risk, RED_BORDER_HEX, width_pt=24)
    p_risk.paragraph_format.left_indent = Cm(0.4)
    r1 = p_risk.add_run("Risk:  ")
    r1.font.bold = True
    r1.font.size = Pt(10)
    p_risk.add_run(risk).font.size = Pt(10)

    # Remedy
    p_rem = doc.add_paragraph()
    _add_left_border(p_rem, RED_BORDER_HEX, width_pt=24)
    p_rem.paragraph_format.left_indent = Cm(0.4)
    r2 = p_rem.add_run("What would close it:  ")
    r2.font.bold = True
    r2.font.size = Pt(10)
    p_rem.add_run(remedy).font.size = Pt(10)

    doc.add_paragraph()  # spacer


# ===========================================================================
# Cover page
# ===========================================================================

def _add_cover(doc: Document) -> None:
    doc.add_paragraph()
    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Hallucination Mitigation")
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = NAVY

    doc.add_paragraph()

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_para.add_run("Clinical RAG Assistant — Technical Documentation")
    run.font.size = Pt(14)
    run.font.color.rgb = TEAL

    doc.add_paragraph()

    meta_lines = [
        ("System:", "Adaptive RAG Clinical Assistant"),
        ("Version:", "1.0"),
        ("Date:", TODAY),
        ("Classification:", "Internal / Clinical Governance Reference"),
    ]
    for label, value in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label + "  ")
        r1.font.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        r2 = p.add_run(value)
        r2.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()

    purpose_para = doc.add_paragraph()
    purpose_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    purpose_run = purpose_para.add_run(
        "Purpose:\n"
        "This document describes every mechanism implemented to prevent the\n"
        "Clinical RAG Assistant from generating factually unsupported or fabricated\n"
        "clinical content, and the limits of those mechanisms."
    )
    purpose_run.font.size = Pt(11)
    purpose_run.font.italic = True
    purpose_run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    _page_break(doc)


# ===========================================================================
# Table of Contents placeholder
# ===========================================================================

def _add_toc(doc: Document) -> None:
    _heading1(doc, "Table of Contents")
    p = doc.add_paragraph()
    run = p.add_run(
        "[This table of contents is auto-generated by Word. "
        "To update: right-click → Update Field → Update entire table.]"
    )
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    # Word TOC field
    para = doc.add_paragraph()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run = para.add_run()
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

    _page_break(doc)


# ===========================================================================
# Section 1 — Defense-in-Depth Overview
# ===========================================================================

def _section1(doc: Document) -> None:
    _heading1(doc, "1  Defense-in-Depth Overview")

    doc.add_paragraph(
        "A single hallucination-prevention control is insufficient for a clinical RAG system. "
        "Any individual mechanism can be bypassed: a system-level instruction may be diluted "
        "by a long user turn; an in-prompt constraint may be ignored under adversarial or "
        "ambiguous phrasing; a post-generation scorer catches only what it can measure. "
        "A layered architecture ensures that a hallucination that escapes one layer must "
        "survive at least two additional independent checks before reaching the user. "
        "This is not merely best practice — it is a safety requirement. In the clinical "
        "trial domain, a fabricated eligibility criterion may enrol an ineligible patient, "
        "an incorrect p-value may inform a regulatory decision, and a misattributed adverse "
        "event may delay or distort safety reporting. Each of these outcomes can have direct "
        "patient safety consequences. The three-layer architecture implemented in this system "
        "addresses pre-generation constraints (Layers 1 and 2) and post-generation detection "
        "(Layer 3)."
    )

    doc.add_paragraph()
    doc.add_paragraph("Three-Layer Architecture Summary:").runs[0].font.bold = True

    _table_with_header(
        doc,
        ["Layer", "When Applied", "Mechanism", "File", "Key Identifier"],
        [
            ["1", "Pre-generation", "LLM role instruction forbidding fabrication",
             "src/prompts.py:21–27", "SYSTEM_PROMPT"],
            ["2", "Pre-generation", "In-prompt constraint injected immediately before the question",
             "src/prompts.py:30–34", "_GROUNDING_INSTRUCTION"],
            ["3", "Post-generation", "Cosine similarity scoring of response sentences vs. retrieved chunks",
             "src/faithfulness.py:17, 34–138", "FaithfulnessChecker.check()"],
        ],
    )

    _heading2(doc, "1.1  Figure 1 — Layered Defense Architecture")

    _ascii_diagram(doc, "Figure 1 — Layered defense-in-depth architecture", [
        "  ┌─────────────────────────────────────────────────────────────────────┐",
        "  │  PRE-GENERATION DEFENSES                                            │",
        "  │                                                                     │",
        "  │  ┌─────────────────────────────────────────────────────────────┐   │",
        "  │  │  Layer 1: SYSTEM_PROMPT  (src/prompts.py:21–27)            │   │",
        "  │  │  Passed as system-role message to every LLM call            │   │",
        "  │  │  Forbids: fabrication, invented citations, out-of-context   │   │",
        "  │  │  knowledge. Requires: explicit disclosure of gaps.          │   │",
        "  │  └─────────────────────────────────────────────────────────────┘   │",
        "  │                           │                                         │",
        "  │  ┌─────────────────────────────────────────────────────────────┐   │",
        "  │  │  Layer 2: _GROUNDING_INSTRUCTION  (src/prompts.py:30–34)   │   │",
        "  │  │  Injected in user message after context, before question     │   │",
        "  │  │  Reinforces: ONLY context, no prior knowledge, no inference  │   │",
        "  │  └─────────────────────────────────────────────────────────────┘   │",
        "  └──────────────────────────────────┬──────────────────────────────────┘",
        "                                      │",
        "  ┌───────────────────────────────────▼─────────────────────────────────┐",
        "  │  GENERATION                                                          │",
        "  │  LLM call with source-tagged context chunks [Source N: Page X, Chunk Y] │",
        "  └───────────────────────────────────┬──────────────────────────────────┘",
        "                                      │",
        "  ┌───────────────────────────────────▼─────────────────────────────────┐",
        "  │  POST-GENERATION DETECTION                                           │",
        "  │  Layer 3: FaithfulnessChecker.check()  (src/faithfulness.py:66–126) │",
        "  │  Cosine similarity per sentence vs. retrieved chunks                 │",
        "  └──────────┬──────────────────────────┬─────────────────────────────┬─┘",
        "             │                           │                             │",
        "  ┌──────────▼──────────┐  ┌────────────▼────────────┐  ┌────────────▼──────────┐",
        "  │  score < 0.25       │  │  0.25 ≤ score < 0.45    │  │  score ≥ 0.45          │",
        "  │  ⛔ Response blocked │  │  Warning banner shown    │  │  Score displayed only  │",
        "  │  st.error() refusal │  │  (orange st.warning)     │  │  No warning shown      │",
        "  └─────────────────────┘  └─────────────────────────┘  └────────────────────────┘",
    ])


# ===========================================================================
# Section 2 — Layer 1: System Prompt
# ===========================================================================

def _section2(doc: Document) -> None:
    _page_break(doc)
    _heading1(doc, "2  Layer 1: System Prompt")

    _heading2(doc, "2.1  Purpose")
    doc.add_paragraph(
        "The system prompt is the outermost defence layer. It is supplied as the "
        "system-role message in every LLM call, establishing the model's behavioural "
        "contract before any user input is processed. By setting the assistant's identity "
        "as a 'clinical trials document assistant' that operates strictly within the "
        "provided document context, the prompt creates a standing prohibition on "
        "fabrication, invented citations, and use of parametric (pre-trained) knowledge. "
        "It also establishes a positive obligation: when the context is insufficient, "
        "the model must say so explicitly rather than speculating."
    )
    doc.add_paragraph(
        "This layer is necessary but not sufficient as a standalone control. LLM "
        "instruction-following from a system prompt can be diluted when user turns are long "
        "or contain contradictory framing. Layer 2 provides a second, proximate constraint "
        "positioned at the point of generation — immediately before the question — to "
        "reinforce the system-level prohibition."
    )

    _heading2(doc, "2.2  Implementation")
    _mono_ref(doc, "src/prompts.py  →  SYSTEM_PROMPT  (lines 21–27)")
    _mono_ref(doc, "app.py:938–942  →  llm.generate(prompt=prompt, system_prompt=SYSTEM_PROMPT, temperature=0)")

    doc.add_paragraph()
    doc.add_paragraph("Verbatim text of SYSTEM_PROMPT (quoted from src/prompts.py:21–27):").runs[0].font.bold = True
    doc.add_paragraph()

    _callout_prompt(doc, "SYSTEM_PROMPT  —  src/prompts.py:21–27", (
        'SYSTEM_PROMPT = ('
        '\n    "You are a clinical trials document assistant. "'
        '\n    "You answer questions strictly based on the document context provided by the user. "'
        '\n    "Never fabricate information, invent citations, or draw on knowledge outside the provided context. "'
        '\n    "If the document context does not contain enough information to answer a question fully, "'
        '\n    "say so explicitly and describe what is missing rather than guessing."'
        '\n)'
    ))

    doc.add_paragraph()

    _heading2(doc, "2.3  Behaviours Enforced")
    doc.add_paragraph(
        "The SYSTEM_PROMPT enforces the following specific behaviours:"
    )
    behaviours = [
        ("Fabrication forbidden:", "The phrase 'Never fabricate information' creates an explicit prohibition on inventing clinical facts."),
        ("Citation invention forbidden:", "The phrase 'invent citations' explicitly prohibits the model from generating source references that do not correspond to retrieved chunks."),
        ("Out-of-context knowledge forbidden:", "'draw on knowledge outside the provided context' restricts the model to the retrieved document context exclusively."),
        ("Gap disclosure required:", "'say so explicitly and describe what is missing rather than guessing' mandates that the model surface its own limitations rather than filling gaps speculatively."),
    ]
    for label, desc in behaviours:
        p = doc.add_paragraph(style="List Bullet")
        r1 = p.add_run(label + "  ")
        r1.font.bold = True
        p.add_run(desc)

    _heading2(doc, "2.4  Scope and Limitations")
    doc.add_paragraph(
        "SYSTEM_PROMPT is passed to every LLM call in the pipeline via "
        "app.py:938–942, using the LLMProvider.generate() interface defined in "
        "src/llm/base.py. Both OpenAI (cloud) and Ollama (local) providers honour "
        "the system_prompt parameter. There is no mechanism to suppress or override "
        "this parameter at runtime."
    )
    doc.add_paragraph(
        "Limitation: system prompt influence can diminish in extended or complex user "
        "turns. This is the primary motivation for Layer 2, which places an identical "
        "grounding constraint at the proximate position immediately before the question."
    )

    _heading2(doc, "2.5  Code Reference")
    _mono_ref(doc, "src/prompts.py:21–27        SYSTEM_PROMPT (variable definition)")
    _mono_ref(doc, "src/llm/base.py:33          LLMProvider.generate(system_prompt=...)")
    _mono_ref(doc, "app.py:61                   from src.prompts import SYSTEM_PROMPT, build_adaptive_prompt")
    _mono_ref(doc, "app.py:938–942              llm.generate(prompt=prompt, system_prompt=SYSTEM_PROMPT, temperature=0)")


# ===========================================================================
# Section 3 — Layer 2: Grounding Instruction & Source Tagging
# ===========================================================================

def _section3(doc: Document) -> None:
    _page_break(doc)
    _heading1(doc, "3  Layer 2: Grounding Instruction and Source Tagging")

    _heading2(doc, "3.1  Purpose")
    doc.add_paragraph(
        "Layer 2 provides a second, in-prompt constraint that reinforces Layer 1 at the "
        "point closest to generation. Unlike the system prompt, which is processed before "
        "any user message, the grounding instruction is injected into the user message "
        "itself — positioned immediately after the retrieved context and immediately before "
        "the user question. This placement is deliberate: it is the last instruction the "
        "model reads before producing its answer, maximising its salience and minimising "
        "the risk that the constraint is 'forgotten' in the presence of a long context block."
    )
    doc.add_paragraph(
        "Layer 2 also provides source attribution infrastructure through "
        "format_source_reference() (src/utils.py:115–129), which wraps every retrieved "
        "chunk in a structured tag — [Source N: Page X, Chunk Y] — giving the model "
        "explicit, traceable anchors to cite rather than vague or invented references."
    )

    _heading2(doc, "3.2  Grounding Instruction Implementation")
    _mono_ref(doc, "src/prompts.py:30–34        _GROUNDING_INSTRUCTION (variable definition)")
    _mono_ref(doc, "src/prompts.py:249          injection point in ResponseStyler.generate_prompt()")

    doc.add_paragraph()
    doc.add_paragraph("Verbatim text of _GROUNDING_INSTRUCTION (quoted from src/prompts.py:30–34):").runs[0].font.bold = True
    doc.add_paragraph()

    _callout_prompt(doc, "_GROUNDING_INSTRUCTION  —  src/prompts.py:30–34", (
        '_GROUNDING_INSTRUCTION = """GROUNDING REQUIREMENT:'
        '\nAnswer using ONLY the information in the CONTEXT FROM DOCUMENT section above. \\'
        '\nDo not use prior knowledge, external sources, or infer facts not explicitly stated in the context. \\'
        '\nIf the context does not contain sufficient information to fully answer the question, \\'
        '\nexplicitly state what is not covered rather than speculating."""'
    ))

    doc.add_paragraph()

    _heading2(doc, "3.3  Source Tagging Implementation")
    _mono_ref(doc, "src/utils.py:115–129        format_source_reference(document, index)")
    _mono_ref(doc, "src/prompts.py:18           from .utils import format_source_reference")
    _mono_ref(doc, "src/prompts.py:276          source_ref = format_source_reference(doc, index=i)")

    doc.add_paragraph()
    doc.add_paragraph(
        "The function format_source_reference() (src/utils.py:115–129) generates a "
        "structured provenance tag for each retrieved chunk. Its return value is:"
    )
    _mono_ref(doc, '    f"[Source {index}: Page {page}, Chunk {chunk_id}]"', indent_cm=0.5)
    doc.add_paragraph(
        "where page and chunk_id are extracted from the LangChain Document metadata "
        "dictionary. This tag is prepended to the chunk text before it is concatenated "
        "into the CONTEXT FROM DOCUMENT block. The model can therefore cite sources as "
        "'[Source 2]' with a precise, verifiable location."
    )

    _heading2(doc, "3.4  Citation Directive")
    doc.add_paragraph(
        "A citation directive is conditionally injected into the formatting instructions "
        "within ResponseStyler.generate_prompt() (src/prompts.py:219–223). It is "
        "controlled by the ResponseConfig.include_references flag:"
    )
    _mono_ref(doc, "src/prompts.py:219–223      if config.include_references:")
    _mono_ref(doc, '                                 "- When referencing specific information, cite the source number"')
    _mono_ref(doc, '                                 "(e.g., \'According to [Source 2]...\' or \'As stated in [Source 1]...\')"')
    doc.add_paragraph()
    doc.add_paragraph(
        "The directive is active for Expert and Regulatory persona configurations "
        "(include_references=True) and inactive for the Novice persona "
        "(include_references=False). This is separate from _GROUNDING_INSTRUCTION, "
        "which is always injected regardless of persona."
    )

    _heading2(doc, "3.5  Figure 2 — Prompt Assembly and Injection Layout")

    _ascii_diagram(doc, "Figure 2 — Exact structure of the assembled prompt sent to the LLM", [
        "  ┌─────────────────────────────────────────────────────────────────────┐",
        "  │  SYSTEM MESSAGE  (passed as system_prompt parameter)               │",
        "  │  SYSTEM_PROMPT  (src/prompts.py:21–27)                             │",
        "  │  'You are a clinical trials document assistant. You answer          │",
        "  │   questions strictly based on the document context provided...      │",
        "  │   Never fabricate information, invent citations, or draw on         │",
        "  │   knowledge outside the provided context...'                        │",
        "  ├─────────────────────────────────────────────────────────────────────┤",
        "  │  USER MESSAGE                                                       │",
        "  │  [Persona instructions — AUDIENCE: ... REQUIREMENTS: ...]           │",
        "  │  [Query-type format instructions — FORMAT FOR <TYPE> QUERY: ...]    │",
        "  │  [Formatting directives — tables, bullet points, max_length, ...]   │",
        "  │                                                                     │",
        "  │  CONTEXT FROM DOCUMENT:                                             │",
        "  │  [Source 1: Page X, Chunk Y]     ← format_source_reference()       │",
        "  │  <retrieved chunk 1 text>         ← src/utils.py:115–129           │",
        "  │                                                                     │",
        "  │  [Source 2: Page X, Chunk Y]                                        │",
        "  │  <retrieved chunk 2 text>                                           │",
        "  │  ...                                                                │",
        "  │                                                                     │",
        "  │  _GROUNDING_INSTRUCTION  (src/prompts.py:30–34)                    │",
        "  │  'GROUNDING REQUIREMENT:                                            │",
        "  │   Answer using ONLY the information in the CONTEXT FROM DOCUMENT    │",
        "  │   section above. Do not use prior knowledge, external sources,      │",
        "  │   or infer facts not explicitly stated in the context...'           │",
        "  │                                                                     │",
        "  │  QUESTION: <user's actual query>   ← src/prompts.py:251            │",
        "  │                                                                     │",
        "  │  ANSWER:                           ← LLM begins generating here    │",
        "  └─────────────────────────────────────────────────────────────────────┘",
    ])

    _heading2(doc, "3.6  Code Reference")
    _mono_ref(doc, "src/utils.py:115–129        format_source_reference(document, index) — returns [Source N: Page X, Chunk Y]")
    _mono_ref(doc, "src/prompts.py:18           from .utils import format_source_reference")
    _mono_ref(doc, "src/prompts.py:30–34        _GROUNDING_INSTRUCTION (variable definition)")
    _mono_ref(doc, "src/prompts.py:174–255      ResponseStyler.generate_prompt() — assembles and injects both layers")
    _mono_ref(doc, "src/prompts.py:249          {_GROUNDING_INSTRUCTION} injection point in prompt template")
    _mono_ref(doc, "src/prompts.py:258–281      build_adaptive_prompt() — formats docs and calls generate_prompt()")
    _mono_ref(doc, "app.py:920                  prompt = build_adaptive_prompt(source_docs, query, response_config)")


# ===========================================================================
# Section 4 — Layer 3: Post-Generation Faithfulness Scorer
# ===========================================================================

def _section4(doc: Document) -> None:
    _page_break(doc)
    _heading1(doc, "4  Layer 3: Post-Generation Faithfulness Scorer")

    _heading2(doc, "4.1  Purpose")
    doc.add_paragraph(
        "Layer 3 is the detection layer. Unlike Layers 1 and 2, which are prophylactic "
        "(they attempt to prevent hallucinations before they occur), Layer 3 is diagnostic: "
        "it measures, after generation, how well the response is grounded in the retrieved "
        "context. It gives the user a quantified, sentence-level signal about which parts "
        "of the response may be unsupported."
    )
    doc.add_paragraph(
        "Critically, Layer 3 does not block responses. In a clinical research context, "
        "suppressing a response entirely may be worse than flagging it for review: a "
        "principal investigator who receives a 'response blocked' message has no basis "
        "for further verification. A flagged-but-visible response with a low faithfulness "
        "score and highlighted sentences gives the user the information needed to check "
        "the original document. This design treats the clinician as the ultimate safety "
        "control, supported by — not replaced by — automated detection."
    )

    _heading2(doc, "4.2  Algorithm")
    _mono_ref(doc, "src/faithfulness.py:34–138   class FaithfulnessChecker")
    _mono_ref(doc, "src/faithfulness.py:66–126   FaithfulnessChecker.check(response_text, context_documents)")
    _mono_ref(doc, "src/faithfulness.py:132–138  FaithfulnessChecker._split_sentences(text)")

    doc.add_paragraph()
    doc.add_paragraph("Algorithm (exact as implemented):").runs[0].font.bold = True
    steps = [
        ("Step 1 — Sentence splitting (src/faithfulness.py:80, 132–138):",
         "Split the generated response into sentences using the regex r'(?<=[.?!])\\s+'. "
         "Filter out sentences shorter than min_sentence_length characters (default: 15). "
         "This prevents empty lines or punctuation artefacts from inflating the score."),
        ("Step 2 — Early exit on empty input (src/faithfulness.py:82–93):",
         "If the sentence list or context_documents list is empty, return score=0.0 "
         "immediately without calling the embedder."),
        ("Step 3 — Single embedding call (src/faithfulness.py:95–98):",
         "Concatenate all sentences and all chunk texts into a single list (all_texts = "
         "sentences + chunk_texts). Pass all_texts to embedder.embed_documents() in one "
         "call. This reuses the HuggingFaceEmbeddings instance already loaded in session "
         "state — no additional model download is required."),
        ("Step 4 — Cosine similarity matrix (src/faithfulness.py:100–107):",
         "Partition the embedding matrix: first n_sentences rows → sentence_embeddings; "
         "remaining rows → chunk_embeddings. Compute cosine_similarity(sentence_embeddings, "
         "chunk_embeddings) to produce an (n_sentences × n_chunks) matrix. "
         "per_sentence_scores = max across columns (max similarity to any chunk per sentence)."),
        ("Step 5 — Aggregate score (src/faithfulness.py:108):",
         "overall_score = mean(per_sentence_scores). Range: 0.0 (no sentence has any "
         "semantic similarity to the context) to 1.0 (every sentence is maximally similar "
         "to at least one context chunk)."),
        ("Step 6 — Threshold application and result assembly (src/faithfulness.py:110–126):",
         "Identify low_confidence_indices where per_sentence_scores[i] < sentence_threshold "
         "(default 0.35). Return a FaithfulnessResult dataclass containing: score, "
         "sentence_scores, sentences, low_confidence_sentences, low_confidence_indices, "
         "context_chunks_used, latency_ms, sentence_threshold."),
    ]
    for label, desc in steps:
        p = doc.add_paragraph(style="List Number")
        r1 = p.add_run(label + "  ")
        r1.font.bold = True
        r1.font.size = Pt(10)
        p.add_run(desc).font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph("Core mathematical formulation:").runs[0].font.bold = True
    _callout_prompt(doc, "Faithfulness scoring formulae", (
        "Per-sentence score:\n"
        "  score_i = max { cosine_sim(embed(s_i), embed(c_j)) }\n"
        "             j ∈ retrieved_chunks\n"
        "\n"
        "Overall faithfulness score:\n"
        "  overall_score = (1/n) × Σ score_i   where i = 1..n sentences\n"
        "\n"
        "where embed() = embedder.embed_documents()  →  HuggingFaceEmbeddings\n"
        "      cosine_sim = sklearn.metrics.pairwise.cosine_similarity"
    ))

    _heading2(doc, "4.3  Threshold Values and UI Behaviour")
    _mono_ref(doc, "src/faithfulness.py:17      FAITHFULNESS_WARNING_THRESHOLD: float = 0.45  (hardcoded module constant)")
    _mono_ref(doc, "src/config.py               FAITHFULNESS_BLOCK_THRESHOLD: float = 0.25  (operational policy threshold)")
    _mono_ref(doc, "src/faithfulness.py:55      sentence_threshold: float = 0.35  (default constructor parameter)")
    _mono_ref(doc, "src/faithfulness.py:31      FaithfulnessResult.sentence_threshold field (default 0.35)")

    doc.add_paragraph()
    _table_with_header(
        doc,
        ["Condition", "Threshold", "UI Behaviour", "Configured In"],
        [
            ["Response blocked (hard gate)",
             "overall_score < 0.25",
             "st.error() replaces response with structured refusal message (app.py)",
             "src/config.py: FAITHFULNESS_BLOCK_THRESHOLD"],
            ["Overall score displayed",
             "Always (all scores)",
             "st.metric('Faithfulness Score', f'{score:.2f}') in Retrieval Statistics expander",
             "src/faithfulness.py:17 + app.py"],
            ["Warning banner shown",
             "overall_score < 0.45",
             "st.warning() with orange warning banner (app.py)",
             "src/faithfulness.py:17 (FAITHFULNESS_WARNING_THRESHOLD)"],
            ["Sentence-level detail shown",
             "any sentence score < 0.35",
             "Nested expander listing low-confidence sentences with individual scores in amber (#d97706)",
             "src/faithfulness.py:55 (sentence_threshold default) + app.py"],
        ],
    )

    _heading2(doc, "4.4  Figure 3 — Faithfulness Scoring Flow")

    _ascii_diagram(doc, "Figure 3 — Faithfulness scoring sequence and UI rendering", [
        "  Generated Response Text",
        "         │",
        "         ▼",
        "  FaithfulnessChecker.check()   (src/faithfulness.py:66–126)",
        "  ├── Split into sentences  (regex: r'(?<=[.?!])\\s+')",
        "  ├── Filter short sentences  (< 15 chars by default)",
        "  │",
        "  │   [For each sentence s_i and each chunk c_j:]",
        "  ├── embedder.embed_documents(sentences + chunks)  ← single call",
        "  │       └── HuggingFaceEmbeddings (reused from session state)",
        "  ├── cosine_similarity(sentence_embeddings, chunk_embeddings)",
        "  ├── score_i = max similarity across all chunks",
        "  ├── overall_score = mean(score_i)",
        "  │",
        "  └── Return FaithfulnessResult",
        "           │",
        "           ▼",
        "  app.py:1026–1054   (Retrieval Statistics expander)",
        "  ├── st.metric('Faithfulness Score', f'{score:.2f}')      ← always shown",
        "  ├── if overall_score < 0.45:                              ← FAITHFULNESS_WARNING_THRESHOLD",
        "  │       st.warning('Low faithfulness score...')           ← orange banner",
        "  └── if any sentence score < 0.35:                         ← sentence_threshold",
        "          st.expander('Low-confidence sentences (N)')",
        "              [score] sentence text  (amber #d97706)",
    ])

    _heading2(doc, "4.5  Figure 4 — Faithfulness Score Decision Flow")

    _ascii_diagram(doc, "Figure 4 — Branching logic for UI output based on threshold comparison", [
        "  FaithfulnessChecker.check() complete",
        "         │",
        "         ▼",
        "  overall_score computed",
        "         │",
        "         ├─────────────────────────────────┬───────────────────────────────┐",
        "         │  score < 0.25                   │  0.25 ≤ score < 0.45          │  score ≥ 0.45",
        "         │  (BLOCK_THRESHOLD)              │  (WARN_THRESHOLD)             │",
        "         ▼                                 ▼                               ▼",
        "  ⛔ Response blocked              Display + warning banner         Display score only",
        "  st.error() refusal              (orange st.warning)              No warning shown",
        "  User directed to source                  │",
        "                              ┌────────────┴─────────────┐",
        "                              │  all sentence scores      │  any sentence score",
        "                              │       >= 0.35             │       < 0.35",
        "                              ▼                           ▼",
        "                       Warning banner only         Warning banner +",
        "                       No sentence detail          Nested expander with",
        "                                                   low-confidence sentences",
        "                                                   (amber, with individual scores)",
    ])

    _heading2(doc, "4.6  Scope and Limitations")

    # RESOLVED callout (teal border) replacing old red wiring-gap warning
    p_resolved = doc.add_paragraph(
        "✅  RESOLVED — Session-State Wiring Gap (app.py:876 now stores embedder):"
    )
    _add_left_border(p_resolved, TEAL_HEX, width_pt=18)
    p_resolved.paragraph_format.left_indent = Cm(0.4)
    p_resolved.runs[0].font.bold = True
    p_resolved.runs[0].font.color.rgb = TEAL

    p_resolved_detail = doc.add_paragraph(
        "app.py:876 now stores st.session_state.embedder = embedder alongside vectordb, "
        "bm25_retriever, and hybrid_retriever. FaithfulnessChecker.check() is now invoked "
        "on every query. The UI rendering code is active. "
        "FAITHFULNESS_BLOCK_THRESHOLD = 0.25 is implemented in src/config.py: "
        "responses with score < 0.25 are replaced with a structured refusal message."
    )
    _add_left_border(p_resolved_detail, TEAL_HEX, width_pt=18)
    p_resolved_detail.paragraph_format.left_indent = Cm(0.4)
    p_resolved_detail.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    limitations = [
        ("Semantic similarity, not logical entailment:",
         "Cosine similarity measures vocabulary and semantic proximity. A sentence that "
         "affirms something the context denies may still score highly if it uses similar "
         "terminology. The scorer does not verify claim direction."),
        ("Automatic response blocking implemented (FAITHFULNESS_BLOCK_THRESHOLD = 0.25):",
         "Responses with score < 0.25 are blocked with a structured refusal. "
         "Responses between 0.25 and 0.45 are shown with an orange warning banner."),
        ("No NLI-based entailment checking:",
         "⚠️ A dedicated natural language inference model to verify that each response "
         "sentence is entailed by (not merely similar to) the retrieved context is not implemented."),
        ("No self-consistency sampling:",
         "⚠️ Generating N responses and comparing them to detect high-variance (potentially "
         "hallucinated) outputs is not implemented."),
    ]
    for label, desc in limitations:
        p = doc.add_paragraph(style="List Bullet")
        r1 = p.add_run(label + "  ")
        r1.font.bold = True
        p.add_run(desc)

    _heading2(doc, "4.7  Code Reference")
    _mono_ref(doc, "src/faithfulness.py:17      FAITHFULNESS_WARNING_THRESHOLD = 0.45")
    _mono_ref(doc, "src/config.py               FAITHFULNESS_BLOCK_THRESHOLD = 0.25")
    _mono_ref(doc, "src/faithfulness.py:20–32   class FaithfulnessResult (dataclass)")
    _mono_ref(doc, "src/faithfulness.py:34–138  class FaithfulnessChecker")
    _mono_ref(doc, "src/faithfulness.py:52–60   FaithfulnessChecker.__init__(embedder, sentence_threshold=0.35, min_sentence_length=15)")
    _mono_ref(doc, "src/faithfulness.py:66–126  FaithfulnessChecker.check(response_text, context_documents) → FaithfulnessResult")
    _mono_ref(doc, "src/faithfulness.py:132–138 FaithfulnessChecker._split_sentences(text)")
    _mono_ref(doc, "app.py                      from src.faithfulness import FAITHFULNESS_WARNING_THRESHOLD, FaithfulnessChecker")
    _mono_ref(doc, "app.py:876                  st.session_state.embedder = embedder  ← RESOLVED")
    _mono_ref(doc, "app.py                      if st.session_state.get('embedder') is not None:  ← now active")
    _mono_ref(doc, "app.py                      _faith_blocked gate: score < FAITHFULNESS_BLOCK_THRESHOLD → st.error() refusal")


# ===========================================================================
# Section 5 — Test Coverage
# ===========================================================================

def _section5(doc: Document) -> None:
    _page_break(doc)
    _heading1(doc, "5  Test Coverage")

    _heading2(doc, "5.1  What Is Tested")

    _mono_ref(doc, "tests/test_prompts.py  —  Layer 1 and Layer 2 tests")
    doc.add_paragraph()

    prompt_tests = [
        ("TestSystemPrompt.test_system_prompt_is_non_empty",
         "Asserts SYSTEM_PROMPT is non-empty (> 20 characters)."),
        ("TestSystemPrompt.test_system_prompt_forbids_fabrication",
         "Asserts SYSTEM_PROMPT contains 'fabricat' or 'never' (fabrication prohibition)."),
        ("TestSystemPrompt.test_system_prompt_mentions_context",
         "Asserts SYSTEM_PROMPT contains 'context' (context restriction)."),
        ("TestSystemPrompt.test_system_prompt_instructs_to_disclose_gaps",
         "Asserts SYSTEM_PROMPT contains 'missing', 'not contain', or 'explicitly' (gap disclosure)."),
        ("TestGroundingInstruction.test_grounding_present_for_novice",
         "Asserts 'GROUNDING REQUIREMENT' appears in prompts built for Novice persona."),
        ("TestGroundingInstruction.test_grounding_present_for_expert",
         "Asserts 'GROUNDING REQUIREMENT' appears in prompts built for Expert persona."),
        ("TestGroundingInstruction.test_grounding_present_for_regulatory",
         "Asserts 'GROUNDING REQUIREMENT' appears in prompts built for Regulatory persona."),
        ("TestGroundingInstruction.test_grounding_only_uses_context_phrase",
         "Asserts 'ONLY' appears in the assembled prompt."),
        ("TestGroundingInstruction.test_grounding_positioned_after_context",
         "Asserts GROUNDING REQUIREMENT appears after CONTEXT FROM DOCUMENT and before QUESTION: in the prompt."),
        ("TestGroundingInstruction.test_grounding_instructs_speculative_disclosure",
         "Asserts prompt contains 'speculating', 'speculate', or 'explicitly state'."),
        ("TestIncludeReferences.test_citation_instruction_present_when_references_enabled",
         "For Expert config (include_references=True), asserts '[Source' or 'cite the source' appears."),
        ("TestIncludeReferences.test_citation_instruction_absent_for_novice",
         "For Novice config (include_references=False), asserts 'cite the source number' does not appear."),
        ("TestIncludeReferences.test_citation_instruction_when_manually_enabled",
         "For ResponseConfig(include_references=True), asserts citation directive appears."),
        ("TestSourceReferenceFormatting.test_source_tags_present_in_context",
         "Asserts '[Source 1:' and '[Source 2:' appear in the assembled prompt."),
        ("TestSourceReferenceFormatting.test_source_tags_include_page",
         "Asserts 'Page 1' and 'Page 2' appear in the source tags."),
        ("TestSourceReferenceFormatting.test_source_tags_include_chunk",
         "Asserts 'Chunk 0' and 'Chunk 1' appear in the source tags."),
        ("TestSourceReferenceFormatting.test_empty_documents_still_builds_prompt",
         "Empty document list produces a valid prompt containing QUESTION: and GROUNDING REQUIREMENT."),
    ]

    _table_with_header(
        doc,
        ["Test Function", "What It Validates"],
        [[name, desc] for name, desc in prompt_tests],
    )

    _mono_ref(doc, "tests/test_faithfulness.py  —  Layer 3 tests")
    doc.add_paragraph()

    faith_tests = [
        ("TestFaithfulnessCheckerScores.test_perfect_score_when_response_matches_context",
         "Aligned unit vectors → score 1.0 (perfect grounding)."),
        ("TestFaithfulnessCheckerScores.test_zero_score_when_vectors_orthogonal",
         "Orthogonal sentence vector → score 0.0 (no grounding)."),
        ("TestFaithfulnessCheckerScores.test_partial_score_for_mixed_response",
         "Mixed: one grounded sentence, one not → mean = 0.5."),
        ("TestFaithfulnessCheckerScores.test_sentence_scores_length_matches_sentences",
         "sentence_scores length equals sentences length."),
        ("TestLowConfidenceFlagging.test_low_confidence_sentences_flagged",
         "Sentences with score < 0.35 appear in low_confidence_sentences and low_confidence_indices."),
        ("TestLowConfidenceFlagging.test_no_low_confidence_when_all_above_threshold",
         "All scores >= 0.35 → empty low_confidence_sentences and indices."),
        ("TestEdgeCases.test_empty_response_returns_zero_score",
         "Empty response string → score 0.0, embedder.embed_documents never called."),
        ("TestEdgeCases.test_empty_documents_returns_zero_score",
         "No context documents → score 0.0, context_chunks_used=0, embed not called."),
        ("TestEdgeCases.test_short_sentences_filtered",
         "Sentences < 15 chars are filtered out; only longer sentences are scored."),
        ("TestEdgeCases.test_whitespace_only_response_returns_zero_score",
         "Whitespace-only response treated as empty → score 0.0."),
        ("TestResultMetadata.test_latency_ms_non_negative",
         "latency_ms >= 0.0 always."),
        ("TestResultMetadata.test_context_chunks_used_equals_doc_count",
         "context_chunks_used equals the number of documents passed in."),
        ("TestResultMetadata.test_context_chunks_used_zero_for_empty_docs",
         "context_chunks_used = 0 when no documents provided."),
        ("TestConstant.test_faithfulness_warning_threshold_in_valid_range",
         "FAITHFULNESS_WARNING_THRESHOLD is between 0.0 and 1.0 exclusive."),
        ("TestConstant.test_faithfulness_warning_threshold_value",
         "FAITHFULNESS_WARNING_THRESHOLD == 0.45 (exact value check)."),
        ("TestConstant.test_threshold_comparison_triggers_warning",
         "Score 0.44 < 0.45 → condition evaluates True (warning triggered)."),
        ("TestConstant.test_threshold_comparison_does_not_trigger_warning",
         "Score 0.46 is not < 0.45 → condition evaluates False (no warning)."),
    ]

    _table_with_header(
        doc,
        ["Test Function", "What It Validates"],
        [[name, desc] for name, desc in faith_tests],
    )

    _mono_ref(doc, "tests/test_utils.py  —  format_source_reference() coverage")
    doc.add_paragraph()

    utils_tests = [
        ("TestFormatSourceReference.test_happy_path",
         "page=3, chunk_id=7, index=1 → '[Source 1: Page 3, Chunk 7]'."),
        ("TestFormatSourceReference.test_default_index_is_1",
         "Default index parameter is 1."),
        ("TestFormatSourceReference.test_missing_page_returns_na",
         "Missing page key → 'N/A' in result; chunk_id still shown."),
        ("TestFormatSourceReference.test_missing_chunk_id_returns_na",
         "Missing chunk_id key → 'N/A' in result; page still shown."),
        ("TestFormatSourceReference.test_both_missing_returns_na_na",
         "Empty metadata → '[Source 1: Page N/A, Chunk N/A]'."),
        ("TestFormatSourceReference.test_custom_index",
         "index=5 → '[Source 5: Page 10, Chunk 99]'."),
    ]

    _table_with_header(
        doc,
        ["Test Function", "What It Validates"],
        [[name, desc] for name, desc in utils_tests],
    )

    _mono_ref(doc, "tests/test_llm_system_prompt.py  —  SYSTEM_PROMPT forwarding (OpenAI + Ollama)")
    doc.add_paragraph()

    llm_sp_tests = [
        ("TestOpenAIProviderSystemPrompt.test_system_prompt_sent_as_first_message",
         "system_prompt is passed as the first system-role message to OpenAI chat.completions.create."),
        ("TestOpenAIProviderSystemPrompt.test_empty_system_prompt_omits_system_message",
         "Empty system_prompt → no system-role message in the messages list (OpenAI)."),
        ("TestOllamaProviderSystemPrompt.test_system_prompt_sent_as_first_message",
         "system_prompt is passed as the first system-role message to ollama.chat."),
        ("TestOllamaProviderSystemPrompt.test_empty_system_prompt_omits_system_message",
         "Empty system_prompt → no system-role message in the messages list (Ollama)."),
    ]

    _table_with_header(
        doc,
        ["Test Function", "What It Validates"],
        [[name, desc] for name, desc in llm_sp_tests],
    )

    _mono_ref(doc, "tests/test_integration.py  —  End-to-end pipeline (all 3 layers)")
    doc.add_paragraph()

    integ_tests = [
        ("TestEndToEndPipeline.test_layer1_system_prompt_is_non_empty",
         "Layer 1: SYSTEM_PROMPT is non-empty and contains fabrication prohibition keyword."),
        ("TestEndToEndPipeline.test_layer2_grounding_in_all_persona_prompts",
         "Layer 2: 'GROUNDING REQUIREMENT' appears in prompts for all 5 UserType personas."),
        ("TestEndToEndPipeline.test_layer3_faithfulness_checker_returns_result",
         "Layer 3: FaithfulnessChecker.check() returns a non-None result; SYSTEM_PROMPT forwarded in mock call."),
        ("TestEndToEndPipeline.test_all_three_layers_connected",
         "End-to-end: all three layers active in a single pipeline pass (mock LLM + mock embedder)."),
    ]

    _table_with_header(
        doc,
        ["Test Function", "What It Validates"],
        [[name, desc] for name, desc in integ_tests],
    )

    _heading2(doc, "5.2  What Is Not Tested")

    gaps = [
        ("Streamlit widget-level UI rendering — ⚠️ No test coverage",
         "The Streamlit UI rendering (warning banner, block gate error display, "
         "low-confidence sentence list) has no automated test. Testing Streamlit "
         "widget output requires a headless browser framework (e.g. Playwright with "
         "st.testing) and is not in scope. The threshold comparison logic underlying "
         "the gate is covered by TestConstant tests."),
    ]

    for title, desc in gaps:
        p = doc.add_paragraph(style="List Bullet")
        r1 = p.add_run(title + ":  ")
        r1.font.bold = True
        p.add_run(desc)

    _heading2(doc, "5.3  Code Reference")
    _mono_ref(doc, "tests/test_prompts.py           Layer 1 and Layer 2 unit tests (4 test classes, 17 test functions)")
    _mono_ref(doc, "tests/test_faithfulness.py      Layer 3 unit tests (5 test classes, 17 test functions)")
    _mono_ref(doc, "tests/test_utils.py             format_source_reference() unit tests (6 test functions)")
    _mono_ref(doc, "tests/test_llm_system_prompt.py SYSTEM_PROMPT forwarding tests — OpenAI + Ollama (4 test functions)")
    _mono_ref(doc, "tests/test_integration.py       End-to-end pipeline integration tests (4 test functions)")


# ===========================================================================
# Section 6 — Hallucination Risk Matrix
# ===========================================================================

def _section6(doc: Document) -> None:
    _page_break(doc)
    _heading1(doc, "6  Hallucination Risk Matrix")

    doc.add_paragraph(
        "The following table consolidates all identified hallucination risk types, "
        "where they occur in the pipeline, which mitigation is applied, and how "
        "detection is performed. Every row is traceable to actual implemented code."
    )

    _table_with_header(
        doc,
        ["Hallucination Type", "Where It Occurs", "Mitigation Applied", "Detection Method", "File(s)"],
        [
            ["Fabricated clinical fact\n(invented eligibility criterion, p-value, dosing)",
             "LLM generation phase",
             "Layer 1 + Layer 2 restrict generation to provided context",
             "Layer 3: faithfulness score < 0.45 flags; score < 0.25 blocks\n(RESOLVED: active since app.py:876)",
             "src/prompts.py, src/faithfulness.py, src/config.py"],
            ["Invented citation\n(source reference with no corresponding chunk)",
             "LLM generation phase",
             "Layer 1 explicitly forbids 'invent citations';\nformat_source_reference() provides real anchor tags",
             "Citation format test: tests/test_prompts.py →\nTestSourceReferenceFormatting",
             "src/prompts.py:21–27, src/utils.py:115–129,\ntests/test_prompts.py"],
            ["Out-of-context parametric knowledge\n(LLM drawing on pre-training, not retrieved docs)",
             "LLM generation phase",
             "Layer 1 prohibits 'knowledge outside the provided context';\nLayer 2 reinforces 'ONLY the information in the CONTEXT FROM DOCUMENT section'",
             "Layer 3: low faithfulness score (semantic similarity to context)\n(RESOLVED: active since app.py:876)",
             "src/prompts.py:21–27, src/prompts.py:30–34"],
            ["Unsupported individual sentence\n(one sentence in an otherwise-grounded response is fabricated)",
             "LLM generation phase",
             "Layer 3 flags at sentence level: any sentence with score < 0.35 is listed",
             "Sentence score < 0.35 → low_confidence_sentences list in UI\n(RESOLVED: active since app.py:876)",
             "src/faithfulness.py:55, app.py"],
            ["Missing information stated as fact\n(LLM presents a gap as a known answer)",
             "LLM generation phase",
             "Layer 1: 'say so explicitly and describe what is missing rather than guessing';\nLayer 2: 'explicitly state what is not covered rather than speculating'",
             "No automated detection — relies on prompt instruction compliance;\ntested in tests/test_prompts.py → TestGroundingInstruction",
             "src/prompts.py:21–27, src/prompts.py:30–34,\ntests/test_prompts.py"],
            ["Wrong source attribution\n(citing [Source 1] for content that came from [Source 3])",
             "LLM generation phase",
             "format_source_reference() provides precise [Source N: Page X, Chunk Y] anchors;\ncitation directive instructs model to cite by number",
             "Source tag format tested in tests/test_prompts.py;\nmissing-metadata cases tested in tests/test_utils.py",
             "src/utils.py:115–129, src/prompts.py:219–223,\ntests/test_prompts.py, tests/test_utils.py"],
            ["Persona-appropriate fabrication\n(Novice-friendly simplification that introduces inaccuracy)",
             "LLM generation phase, post-adaptation",
             "Layer 1 and Layer 2 apply regardless of persona;\nUI citation disclaimer shown when include_references=False",
             "No specific detection — faithfulness scorer does not distinguish persona contexts",
             "src/prompts.py, app.py"],
        ],
    )


# ===========================================================================
# Section 7 — Known Gaps and Future Work
# ===========================================================================

def _resolved_block(doc: Document, title: str, detail: str) -> None:
    """Add a teal-left-bordered RESOLVED callout for Section 7."""
    p_title = doc.add_paragraph()
    _add_left_border(p_title, TEAL_HEX, width_pt=24)
    p_title.paragraph_format.left_indent = Cm(0.4)
    r_title = p_title.add_run("✅  RESOLVED — " + title)
    r_title.font.bold = True
    r_title.font.size = Pt(11)
    r_title.font.color.rgb = TEAL

    p_detail = doc.add_paragraph()
    _add_left_border(p_detail, TEAL_HEX, width_pt=24)
    p_detail.paragraph_format.left_indent = Cm(0.4)
    p_detail.add_run(detail).font.size = Pt(10)

    doc.add_paragraph()  # spacer


def _section7(doc: Document) -> None:
    _page_break(doc)
    _heading1(doc, "7  Known Gaps and Future Work")

    intro = doc.add_paragraph(
        "This section is specifically formatted for clinical governance review. "
        "Resolved items are shown with a teal border. Remaining open gaps are shown "
        "with a red border. Each open gap is accompanied by a statement of the clinical "
        "risk it creates and a concrete description of what would close it."
    )
    _add_left_border(intro, RED_BORDER_HEX, width_pt=6)

    doc.add_paragraph()

    _resolved_block(
        doc,
        title="Session-State Wiring Bug (FaithfulnessChecker)",
        detail=(
            "app.py:876 now stores st.session_state.embedder = embedder alongside vectordb, "
            "bm25_retriever, and hybrid_retriever. FaithfulnessChecker.check() is now invoked "
            "on every query. The UI rendering (score metric, warning banner, low-confidence "
            "sentence list) is active code. Integration test in tests/test_integration.py "
            "verifies Layer 3 is wired end-to-end."
        ),
    )

    _resolved_block(
        doc,
        title="Automatic Response Blocking on Low Faithfulness Score",
        detail=(
            "FAITHFULNESS_BLOCK_THRESHOLD = 0.25 added to src/config.py. "
            "When faithfulness_result.score < 0.25, app.py replaces the response with "
            "a structured st.error() refusal message directing the user to consult the "
            "source document. This is distinct from the warning threshold (0.45): "
            "scores between 0.25 and 0.45 display a warning banner but show the response; "
            "scores below 0.25 block the response entirely."
        ),
    )

    _resolved_block(
        doc,
        title="No Integration Test for End-to-End Mitigation Pipeline",
        detail=(
            "tests/test_integration.py added with 4 tests: "
            "(1) SYSTEM_PROMPT is non-empty and contains fabrication prohibition, "
            "(2) GROUNDING REQUIREMENT present in all 5 persona prompts, "
            "(3) Layer 3 faithfulness check returns a non-None result with score in [0, 1], "
            "(4) all three layers active in a single end-to-end pipeline pass."
        ),
    )

    _resolved_block(
        doc,
        title="Persona-Level Grounding Relaxation Not Guarded",
        detail=(
            "When include_references is False (e.g., Novice persona), app.py now displays "
            "a UI caption: 'Citations are not shown for this expertise level. "
            "Verify answers against the source document.' "
            "This ensures users are informed when citation directives are omitted, "
            "and can verify answers independently."
        ),
    )

    _gap_block(
        doc,
        title="NLI-Based Entailment Checking Not Implemented",
        risk=(
            "Cosine similarity scoring (Layer 3) measures vocabulary and semantic proximity, "
            "not logical entailment. A response sentence that contradicts the retrieved context "
            "(e.g., 'Patients may be re-enrolled' when the protocol says 'Re-enrolment is "
            "prohibited') may still score highly if the vocabulary is similar. The scorer "
            "cannot detect directional errors of this kind."
        ),
        remedy=(
            "Integrate a dedicated Natural Language Inference model (e.g., a cross-encoder "
            "trained for NLI such as cross-encoder/nli-deberta-v3-base) to verify that each "
            "response sentence is entailed by (not merely similar to) the retrieved context "
            "chunks. This would be a fourth post-generation layer operating on the sentences "
            "already identified as low-confidence by Layer 3."
        ),
    )

    _gap_block(
        doc,
        title="Self-Consistency Sampling Not Implemented",
        risk=(
            "A single generation pass may produce a confident-sounding but unsupported "
            "answer. High confidence in the surface form of the response does not indicate "
            "factual grounding. The faithfulness scorer measures similarity to the retrieved "
            "context, not the stability of the generation across runs."
        ),
        remedy=(
            "Generate N independent responses for each query (e.g., N=3–5 with temperature > 0) "
            "and compare them. High variance across responses (measured via pairwise cosine "
            "similarity or n-gram overlap) is a signal that the LLM is uncertain and potentially "
            "hallucinating. Flag such queries with an additional warning. This approach is "
            "computationally more expensive but provides a generation-stability signal "
            "orthogonal to the faithfulness score."
            "\n\nDecision note: Self-consistency sampling was evaluated and explicitly deferred. "
            "At temperature=0 (the current setting for clinical determinism), all N generations "
            "are identical and the technique produces no variance signal. At temperature > 0, "
            "the 3–5× LLM call overhead is considered disproportionate to the marginal safety gain, "
            "given that FAITHFULNESS_BLOCK_THRESHOLD = 0.25 already blocks severely ungrounded "
            "responses. This is a documented architectural decision, not an oversight. "
            "Recommended for re-evaluation once the system is in production and per-query "
            "latency budgets are established."
        ),
    )


# ===========================================================================
# Main generator
# ===========================================================================

def generate() -> Path:
    doc = Document()

    # Page size: A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    _add_page_numbers(doc)
    _add_cover(doc)
    _add_toc(doc)

    _section1(doc)
    _section2(doc)
    _section3(doc)
    _section4(doc)
    _section5(doc)
    _section6(doc)
    _section7(doc)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = generate()
    print(f"Generated: {path}")
