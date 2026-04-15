"""
generate_adaptive_evaluation.py

Generates docs/rag-system/adaptive_evaluation.docx — a full technical reference
for the Adaptive Evaluation Suite of the Clinical RAG Assistant.

Run:
    python docs/rag-system/generate_adaptive_evaluation.py

Requirements:
    pip install python-docx
"""

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Colour constants (identical to generate_rag_pipeline.py)
# ---------------------------------------------------------------------------
NAVY = RGBColor(0x11, 0x22, 0x40)      # #112240  title
TEAL = RGBColor(0x0E, 0xA5, 0xC9)      # #0EA5C9  subtitle / table header
TEAL_HEX = "0EA5C9"
CALLOUT_BG_HEX = "EBF8FF"
DIAGRAM_BG_HEX = "F3F4F6"
RED_BORDER_HEX = "DC2626"
HEADER_TEXT_COLOR = RGBColor(0xFF, 0xFF, 0xFF)  # white header text

OUTPUT_PATH = Path(__file__).parent / "adaptive_evaluation.docx"


# ---------------------------------------------------------------------------
# Low-level XML helpers (verbatim from generate_rag_pipeline.py)
# ---------------------------------------------------------------------------

def _shade_cell(cell, hex_color: str) -> None:
    """Fill a table cell with a solid background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _shade_para(para, hex_color: str) -> None:
    """Apply a solid paragraph background (shading) to simulate a coloured box."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _add_left_border(para, hex_color: str, width_pt: int = 18) -> None:
    """Add a thick left border to a paragraph (gap / warning box style)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(width_pt))   # eighths of a point
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)


def _add_page_numbers(doc: Document) -> None:
    """Insert 'Page X of Y' into the document footer via Word field codes."""
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _field(instr: str):
        run = para.add_run()
        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar_begin)

        run2 = para.add_run()
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = f" {instr} "
        run2._r.append(instrText)

        run3 = para.add_run()
        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")
        run3._r.append(fldChar_end)

    para.add_run("Page ")
    _field("PAGE")
    para.add_run(" of ")
    _field("NUMPAGES")


# ---------------------------------------------------------------------------
# High-level content helpers (verbatim from generate_rag_pipeline.py)
# ---------------------------------------------------------------------------

def add_cover(doc: Document) -> None:
    """Add a styled cover page."""
    doc.add_paragraph()  # top margin spacer
    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Adaptive Evaluation Suite")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = NAVY

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_para.add_run("Clinical RAG Assistant — Technical Documentation")
    run.font.size = Pt(14)
    run.font.color.rgb = TEAL

    doc.add_paragraph()

    meta_lines = [
        ("System:", "Adaptive RAG Clinical Assistant"),
        ("Version:", "1.0"),
        ("Date:", "2026-02-20"),
        ("Classification:", "Internal Technical Reference"),
    ]
    for label, value in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label + "  ")
        r1.font.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(value)
        r2.font.size = Pt(11)

    doc.add_page_break()


def add_toc(doc: Document) -> None:
    """Add a Table of Contents placeholder."""
    h = doc.add_heading("Table of Contents", level=1)
    h.runs[0].font.color.rgb = NAVY

    note = doc.add_paragraph(
        "Right-click this area in Microsoft Word and select 'Update Field' "
        "to generate the Table of Contents automatically."
    )
    note.runs[0].font.italic = True
    note.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)  # grey
    note.runs[0].font.size = Pt(10)

    doc.add_page_break()


def _heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level == 1 else TEAL


def add_teal_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    col_widths: list[float] | None = None,
) -> None:
    """
    Render a table with a teal header row and alternating body rows.

    col_widths: list of Inches widths; if None, Word auto-sizes.
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr_cells[i].text = text
        _shade_cell(hdr_cells[i], TEAL_HEX)
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = HEADER_TEXT_COLOR
        run.font.size = Pt(9.5)

    # Body rows — alternate light grey
    ALT_BG = "F9FAFB"
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        bg = ALT_BG if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, text in enumerate(row_data):
            row_cells[c_idx].text = text
            _shade_cell(row_cells[c_idx], bg)
            run = row_cells[c_idx].paragraphs[0].runs[0]
            run.font.size = Pt(9.5)

    # Apply column widths if provided
    if col_widths:
        for r_idx in range(len(table.rows)):
            for c_idx, width in enumerate(col_widths):
                table.rows[r_idx].cells[c_idx].width = Inches(width)

    doc.add_paragraph()  # spacing after table


def add_callout(doc: Document, label: str, text: str) -> None:
    """
    Add a verbatim callout box (light-blue background, Courier New).

    label: displayed as bold heading above the box.
    text:  monospaced verbatim content.
    """
    lbl = doc.add_paragraph()
    r = lbl.add_run(label)
    r.font.bold = True
    r.font.size = Pt(10)

    for line in text.splitlines():
        p = doc.add_paragraph(line if line.strip() else " ")
        _shade_para(p, CALLOUT_BG_HEX)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(8.5)

    doc.add_paragraph()  # spacing


def add_gap_box(doc: Document, title: str, items: list[str]) -> None:
    """Add a red-bordered warning box listing implementation notes."""
    title_p = doc.add_paragraph()
    _add_left_border(title_p, RED_BORDER_HEX, 24)
    title_p.paragraph_format.left_indent = Inches(0.25)
    r = title_p.add_run(title)
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

    for item in items:
        p = doc.add_paragraph()
        _add_left_border(p, RED_BORDER_HEX, 18)
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(f"• {item}")
        r.font.size = Pt(10)

    doc.add_paragraph()


def add_ascii(doc: Document, caption: str, lines: list[str]) -> None:
    """Render an ASCII diagram in a grey code block with a caption."""
    cap = doc.add_paragraph()
    r = cap.add_run(f"Figure: {caption}")
    r.font.bold = True
    r.font.italic = True
    r.font.size = Pt(10)

    for line in lines:
        p = doc.add_paragraph(line if line else " ")
        _shade_para(p, DIAGRAM_BG_HEX)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.15)
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(8)

    doc.add_paragraph()


# mmdc may not be on PATH in all shells on Windows; resolve the full path via npm prefix.
_NPM_PREFIX = subprocess.run(
    ["node", "-e", "process.stdout.write(require('child_process').execSync('npm config get prefix').toString().trim())"],
    capture_output=True,
    text=True,
).stdout.strip()
_MMDC_CMD = str(Path(_NPM_PREFIX) / "mmdc.cmd") if _NPM_PREFIX else "mmdc"


def _try_mermaid(mmd: str, out_path: Path, mmd_source_path: Path | None = None) -> bool:
    """Attempt to render a Mermaid diagram to PNG. Returns True on success.

    If mmd_source_path is provided and exists, it is used as the .mmd input
    (single source of truth). Otherwise the inline mmd string is written to
    a temporary file that is cleaned up after rendering.
    """
    if mmd_source_path and mmd_source_path.exists():
        mmd_file = mmd_source_path
        temp_created = False
    else:
        mmd_file = out_path.with_suffix(".mmd")
        mmd_file.write_text(mmd, encoding="utf-8")
        temp_created = True

    try:
        result = subprocess.run(
            [
                _MMDC_CMD,
                "-i", str(mmd_file),
                "-o", str(out_path),
                "--scale", "2",
                "--backgroundColor", "white",
            ],
            capture_output=True,
            timeout=60,
        )
        return result.returncode == 0 and out_path.exists()
    except (FileNotFoundError, subprocess.TimeoutExpired, OSError):
        return False
    finally:
        if temp_created and mmd_file.exists():
            mmd_file.unlink()


def add_diagram(
    doc: Document,
    caption: str,
    mmd: str,
    ascii_lines: list[str],
    png_name: str,
    mmd_source_path: Path | None = None,
) -> None:
    """
    Attempt Mermaid PNG render; fall back to ASCII art if mmdc is unavailable.

    mmd_source_path: if supplied and the file exists, it is used as the source
    .mmd file instead of the inline mmd string.
    """
    png_path = OUTPUT_PATH.parent / png_name
    if _try_mermaid(mmd, png_path, mmd_source_path):
        cap = doc.add_paragraph()
        r = cap.add_run(f"Figure: {caption}")
        r.font.bold = True
        r.font.italic = True
        r.font.size = Pt(10)
        doc.add_picture(str(png_path), width=Inches(6.0))
        doc.add_paragraph()
    else:
        add_ascii(doc, caption, ascii_lines)


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def build_section_1_overview(doc: Document) -> None:
    _heading(doc, "1. Suite Overview")

    doc.add_paragraph(
        "The Adaptive Evaluation Suite contains four complementary evaluation modules "
        "that together prove the Clinical RAG Assistant's adaptive pipeline produces "
        "measurably better outputs than a generic RAG baseline. "
        "Adaptation matters clinically: a response calibrated for the wrong expertise "
        "level can lead to a real patient-safety failure — a novice receiving "
        "expert-level jargon may miss a critical safety signal; a PROCEDURE query "
        "misclassified as DEFINITION receives prose instead of numbered steps, "
        "which can cause procedural error. "
        "The four evals systematically measure classification accuracy, readability "
        "gradient, format instruction compliance, and head-to-head win rate."
    )

    mmd = """flowchart TD
  CLI["run_eval.py — main() argparse CLI"]
  E1["Eval 1: classification_accuracy\\nrun_classification_accuracy()"]
  E2["Eval 2: readability_analysis\\nrun_readability_analysis()"]
  E3["Eval 3: format_compliance\\nrun_format_compliance()"]
  E4["Eval 4: adaptive_vs_generic\\nrun_adaptive_vs_generic()"]
  PERSONA["Shared: persona_evaluation.py\\nrun_persona_evaluation()"]
  METRICS["eval/metrics.py\\ncalculate_all_metrics()"]
  CLI -->|"--classify"| E1
  CLI -->|"--readability"| E2
  CLI -->|"--compliance"| E3
  CLI -->|"--adaptive-vs-generic"| E4
  CLI -->|"--metrics"| METRICS
  E2 -->|"reads/generates"| PERSONA
  E3 -->|"reads/generates"| PERSONA"""

    ascii_lines = [
        "  run_eval.py  (main — argparse CLI)",
        "      │",
        "      ├──► --classify          Eval 1: run_classification_accuracy()",
        "      │                            No LLM. No document. < 1 second.",
        "      │",
        "      ├──► --readability       Eval 2: run_readability_analysis()",
        "      │                            Reads persona_responses.json cache",
        "      │                            or calls run_persona_evaluation()",
        "      │",
        "      ├──► --compliance        Eval 3: run_format_compliance()",
        "      │                            Same cache reuse as Eval 2",
        "      │",
        "      ├──► --adaptive-vs-generic  Eval 4: run_adaptive_vs_generic()",
        "      │                            Independent pipeline. Requires document.",
        "      │",
        "      └──► --metrics           eval/metrics.py: calculate_all_metrics()",
        "                                   Aggregates from existing result CSVs",
        "",
        "  Shared Infrastructure",
        "  eval/persona_evaluation.py: run_persona_evaluation()",
        "      Used by Evals 2 and 3 when persona_responses.json is absent",
    ]

    add_diagram(
        doc,
        "Evaluation Suite — Module Relationships",
        mmd,
        ascii_lines,
        "fig5_eval_master_relationship.png",
        mmd_source_path=Path(__file__).parent / "fig5_eval_master_relationship.mmd",
    )


def build_section_2_classification(doc: Document) -> None:
    doc.add_page_break()
    _heading(doc, "2. Eval 1: classification_accuracy")

    _heading(doc, "2.1 Purpose", level=2)
    doc.add_paragraph(
        "Measures how accurately the regex-based query classifier assigns one of 9 QueryType "
        "labels to incoming queries. Accurate classification is a prerequisite for adaptive "
        "prompt routing: a PROCEDURE query misclassified as DEFINITION receives prose "
        "instead of numbered steps, creating risk of procedural error. "
        "A COMPLIANCE query misclassified as SAFETY may omit regulatory citations "
        "required for audit readiness. This eval runs in under 1 second — "
        "no document, no LLM, no API calls required."
    )

    _heading(doc, "2.2 Inputs", level=2)
    add_teal_table(
        doc,
        ["Input", "Type", "Description", "Source"],
        [
            ["query", "str", "Natural language query string", "LABELED_QUERY_DATASET"],
            ["expected_type", "str", "Ground-truth QueryType label", "LABELED_QUERY_DATASET"],
            ["LABELED_QUERY_DATASET", "list[dict]", "45 queries, 5 per QueryType (hardcoded)", "eval/classification_accuracy.py:28"],
        ],
        col_widths=[1.6, 0.8, 2.8, 2.5],
    )

    _heading(doc, "2.3 Method", level=2)

    doc.add_paragraph(
        "The 9 QueryType categories are defined as an enum in src/query_classifier.py:24–35: "
        "DEFINITION, PROCEDURE, COMPLIANCE, COMPARISON, NUMERICAL, TIMELINE, "
        "SAFETY, ELIGIBILITY, COMPLEX."
    )

    doc.add_paragraph(
        "Classification algorithm (QueryClassifier.classify(), src/query_classifier.py:162–194): "
        "Priority-ordered regex pattern matching — not ML or LLM-based. "
        "Step 1: complexity check — if the query has more than 2 sentences or more than one "
        "question mark, it is classified as COMPLEX immediately. "
        "Step 2: domain-specific regex patterns are tested for each of the remaining 8 types. "
        "Step 3: if no pattern matches, DEFINITION is returned as the default fallback."
    )

    doc.add_paragraph(
        "Confidence scoring (QueryClassifier.get_confidence(), src/query_classifier.py:197–217): "
        "confidence = min(1.0, 0.5 + matches × 0.2), "
        "where matches is the count of regex patterns that fired for the predicted type."
    )

    add_callout(
        doc,
        "Accuracy Formulae",
        """Query Type Accuracy = correct_predictions / 45
Per-type Accuracy   = correct / 5   (per QueryType, 5 queries each)

Expertise Accuracy  = correct_profiles / 25
Per-type Accuracy   = correct / 5   (per UserType, 5 profiles each)

Combined Accuracy   = 0.6 × Query_Type_Acc + 0.4 × Expertise_Acc

Dataset sizes:
  LABELED_QUERY_DATASET     — 45 queries,  5 per QueryType (9 types)
  LABELED_EXPERTISE_DATASET — 25 profiles, 5 per UserType  (5 types)""",
    )

    add_callout(
        doc,
        "RESOLVED — Combined Accuracy Implemented",
        """RESOLVED: run_classification_accuracy() now calls run_expertise_accuracy()
after the query-type eval. detect_user_type() is run on each of the 25
LABELED_EXPERTISE_DATASET profiles and compared against the expected UserType.

Combined Accuracy = 0.6 × query_type_acc + 0.4 × expertise_acc

All three metrics (query type, expertise, combined) are written to:
  results/classification_accuracy_summary.txt
  results/expertise_accuracy_results.csv""",
    )

    mmd = """flowchart TD
  Q["User Query (string)"]
  CX{"Complexity Check\\n> 2 sentences OR > 1 '?'"}
  COMPLEX["COMPLEX (early return)"]
  PAT["Domain Pattern Matching\\n9 regex pattern sets"]
  PT["Predicted QueryType"]
  CONF["get_confidence()\\nmin(1.0, 0.5 + matches × 0.2)"]
  GT["Expected Type\\nLABELED_QUERY_DATASET (45 items)"]
  CMP{"predicted == expected?"}
  ACC["Overall Accuracy = correct / 45\\nPer-type: correct / 5"]
  CONFMAT["pd.crosstab()\\nConfusion Matrix"]
  Q --> CX
  CX -->|"yes"| COMPLEX
  CX -->|"no"| PAT
  PAT --> PT
  COMPLEX --> CONF
  PT --> CONF
  PT --> CMP
  GT --> CMP
  CMP -->|"all 45 queries"| ACC
  CMP -->|"all 45 queries"| CONFMAT"""

    ascii_lines = [
        "  ┌─────────────── QUERY TYPE CLASSIFICATION ────────────────────┐",
        "  │  Query (string)                                               │",
        "  │      │                                                        │",
        "  │      ▼                                                        │",
        "  │  Complexity Check  (src/query_classifier.py:162)              │",
        "  │  > 2 sentences  OR  > 1 '?'                                   │",
        "  │      │ no                │ yes                                │",
        "  │      ▼                   ▼                                    │",
        "  │  Domain Pattern Match   COMPLEX (early return)               │",
        "  │  9 regex pattern sets                                         │",
        "  │  → DEFINITION fallback if nothing matches                     │",
        "  │      │                                                        │",
        "  │      ▼                                                        │",
        "  │  Predicted QueryType                                          │",
        "  │      ├──► get_confidence()  min(1.0, 0.5 + matches×0.2)      │",
        "  │      └──► Compare to LABELED_QUERY_DATASET (45 items)        │",
        "  │               ├── Query Type Accuracy = correct / 45          │",
        "  │               └── pd.crosstab() Confusion Matrix              │",
        "  └────────────────────────────────────────────────────────────────┘",
        "",
        "  ┌─────────────── EXPERTISE CLASSIFICATION ──────────────────────┐",
        "  │  User Profile  {role, experience_years}                        │",
        "  │      │                                                         │",
        "  │      ▼                                                         │",
        "  │  detect_user_type()  (src/personas.py:111)                     │",
        "  │  Priority: REGULATORY → EXECUTIVE → EXPERT → NOVICE           │",
        "  │  Fallback: experience_years threshold                          │",
        "  │      │                                                         │",
        "  │      └──► Compare to LABELED_EXPERTISE_DATASET (25 profiles)  │",
        "  │               └── Expertise Accuracy = correct / 25            │",
        "  └────────────────────────────────────────────────────────────────┘",
        "",
        "  Combined Accuracy = 0.6 × Query_Type_Acc + 0.4 × Expertise_Acc",
    ]

    add_diagram(
        doc,
        "Classification Accuracy — Query Flow",
        mmd,
        ascii_lines,
        "fig6_classification_flow.png",
        mmd_source_path=Path(__file__).parent / "fig6_classification_flow.mmd",
    )

    _heading(doc, "2.4 Outputs", level=2)
    add_teal_table(
        doc,
        ["File", "Description"],
        [
            ["results/classification_accuracy_results.csv", "One row per query: query, expected_type, predicted_type, confidence, is_correct"],
            ["results/classification_confusion_matrix.csv", "9×9 confusion matrix via pd.crosstab(expected_type, predicted_type)"],
            ["results/expertise_accuracy_results.csv", "One row per profile: role, experience_years, expected_user_type, predicted_user_type, is_correct"],
            ["results/classification_accuracy_summary.txt", "Human-readable report: query type accuracy, expertise accuracy, combined accuracy, per-type breakdown"],
        ],
        col_widths=[3.2, 4.5],
    )

    note = doc.add_paragraph(
        "No action threshold is defined in the code. There is no automatic pass/fail "
        "gate on combined accuracy — the result is reported for human review."
    )
    note.runs[0].font.italic = True
    note.runs[0].font.size = Pt(9.5)
    note.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    _heading(doc, "2.5 Code Reference", level=2)
    add_teal_table(
        doc,
        ["Step", "File", "Function / Symbol"],
        [
            ["Query type dataset", "eval/classification_accuracy.py:28", "LABELED_QUERY_DATASET (45 items)"],
            ["Expertise dataset", "eval/classification_accuracy.py", "LABELED_EXPERTISE_DATASET (25 profiles)"],
            ["Weights", "eval/classification_accuracy.py", "COMBINED_ACCURACY_W1=0.6, W2=0.4"],
            ["Query types enum", "src/query_classifier.py:24", "QueryType (9 values)"],
            ["Query classification", "src/query_classifier.py:162", "QueryClassifier.classify()"],
            ["Confidence scoring", "src/query_classifier.py:197", "QueryClassifier.get_confidence()"],
            ["Expertise detection", "src/personas.py:111", "detect_user_type(user_profile)"],
            ["Expertise accuracy", "eval/classification_accuracy.py", "run_expertise_accuracy()"],
            ["Combined accuracy", "eval/classification_accuracy.py", "run_classification_accuracy() — calls both"],
            ["Summary report", "eval/classification_accuracy.py", "_write_summary()"],
        ],
        col_widths=[1.8, 2.8, 3.1],
    )


def build_section_3_readability(doc: Document) -> None:
    doc.add_page_break()
    _heading(doc, "3. Eval 2: readability_analysis")

    _heading(doc, "3.1 Purpose", level=2)
    doc.add_paragraph(
        "Proves that NOVICE responses are measurably simpler than EXPERT responses "
        "using established readability metrics. Clinical rationale: an overly complex "
        "response to a novice user (e.g. a patient or junior nurse) can bury a critical "
        "safety signal in jargon they cannot parse — a missed safety signal may delay "
        "reporting or cause incorrect clinical action. "
        "The evaluation reads from a cached persona_responses.json file to avoid "
        "redundant LLM calls; if the file is absent, it generates responses fresh via "
        "run_persona_evaluation()."
    )

    _heading(doc, "3.2 Inputs", level=2)
    add_teal_table(
        doc,
        ["Input", "Type", "Description", "Source"],
        [
            ["response text", "str", "Generated response for a (query, persona) pair", "persona_responses.json cache"],
            ["persona", "str", "One of: novice, intermediate, expert, regulatory, executive", "persona_responses.json"],
            ["query", "str", "The original query string", "DEFAULT_QUERIES (5 items)"],
            ["embedding_model", "str", "Embedding model identifier (default: S-PubMedBert-MS-MARCO)", "run_readability_analysis() param"],
            ["llm_model", "str", "LLM model (default: src/config.py:DEFAULT_LLM_MODEL)", "run_readability_analysis() param"],
        ],
        col_widths=[1.5, 0.8, 2.8, 2.6],
    )

    _heading(doc, "3.3 Method", level=2)

    doc.add_paragraph(
        "Seven metrics are computed per (query, persona) pair by compute_readability_metrics() "
        "(eval/readability_analysis.py:49) using the textstat library. "
        "With 5 default queries and 5 personas, this produces 25 rows in the output DataFrame."
    )

    add_teal_table(
        doc,
        ["Metric", "textstat Function", "Interpretation"],
        [
            ["flesch_reading_ease", "textstat.flesch_reading_ease()", "0–100; higher = simpler. < 30 = very difficult"],
            ["flesch_kincaid_grade", "textstat.flesch_kincaid_grade()", "US school grade level. Used for pass/fail test"],
            ["gunning_fog", "textstat.gunning_fog()", "Years of formal education needed to understand"],
            ["word_count", "len(text.split())", "Total words in response"],
            ["sentence_count", "textstat.sentence_count()", "Total sentences"],
            ["difficult_words", "textstat.difficult_words()", "Count of words with 3+ syllables"],
            ["avg_sentence_length", "word_count / max(sentence_count, 1)", "Mean words per sentence"],
        ],
        col_widths=[1.8, 2.3, 3.6],
    )

    add_callout(
        doc,
        "Pass/Fail Criteria (eval/readability_analysis.py)",
        """Primary pass condition:  novice_fk_grade  <  expert_fk_grade

This directional check confirms NOVICE responses use a lower grade level
than EXPERT responses, proving a measurable readability gradient.

Per-persona target comparison (PERSONA_GRADE_TARGETS):
  Persona        FK Grade Range   Status
  ─────────────  ───────────────  ──────
  Novice         (6.0,  9.0)      PASS / FAIL vs avg FK grade
  Intermediate   (10.0, 13.0)     PASS / FAIL vs avg FK grade
  Expert         (14.0, 18.0)     PASS / FAIL vs avg FK grade
  Regulatory     (14.0, 18.0)     PASS / FAIL vs avg FK grade
  Executive      (8.0,  12.0)     PASS / FAIL vs avg FK grade

Source: PERSONA_GRADE_TARGETS imported from eval/adaptive_vs_generic.py""",
    )

    add_callout(
        doc,
        "RESOLVED — PERSONA_GRADE_TARGETS now used in readability_analysis",
        """RESOLVED: readability_analysis.py now imports PERSONA_GRADE_TARGETS
from eval.adaptive_vs_generic and compares each persona's average FK grade
against its target range in _write_summary().

New section in readability_analysis_summary.txt:
  FK GRADE TARGET COMPARISON
  Persona   FK Grade (avg)   Target Range   Status
  -------   ─────────────    ────────────   ──────
  Novice         8.2         (6, 9)         PASS
  ...
This is in addition to — not replacing — the directional check.""",
    )

    mmd = """flowchart TD
  CACHE{"persona_responses.json\\nexists?"}
  GEN["run_persona_evaluation()\\n(LLM calls required)"]
  RESP["Response Text\\n5 queries × 5 personas = 25 rows"]
  RM["compute_readability_metrics()\\neval/readability_analysis.py:49"]
  M1["flesch_reading_ease"]
  M2["flesch_kincaid_grade"]
  M3["gunning_fog / difficult_words"]
  M4["word_count / sentence_count"]
  PASS{"novice_fk_grade < expert_fk_grade?"}
  PASSLBL["PASS"]
  FAILLBL["FAIL"]
  CSV["readability_analysis_results.csv"]
  CACHE -->|"yes — zero LLM calls"| RESP
  CACHE -->|"no"| GEN
  GEN --> RESP
  RESP --> RM
  RM --> M1
  RM --> M2
  RM --> M3
  RM --> M4
  M2 --> PASS
  PASS -->|"yes"| PASSLBL
  PASS -->|"no"| FAILLBL
  RM --> CSV"""

    ascii_lines = [
        "  persona_responses.json  exists?",
        "      │                       │",
        "      │ yes                   │ no",
        "      ▼                       ▼",
        "  Load from cache         run_persona_evaluation()",
        "  (zero LLM calls)        (calls LLM — requires document)",
        "      │                       │",
        "      └───────────┬───────────┘",
        "                  ▼",
        "  Response Text  (5 queries × 5 personas = 25 rows)",
        "                  │",
        "                  ▼",
        "  compute_readability_metrics()  (eval/readability_analysis.py:49)",
        "  textstat library",
        "      ├── flesch_reading_ease     (higher = simpler)",
        "      ├── flesch_kincaid_grade    (US grade level) ◄── used below",
        "      ├── gunning_fog             (years of education)",
        "      ├── word_count / sentence_count",
        "      └── difficult_words / avg_sentence_length",
        "                  │",
        "                  ├── readability_analysis_results.csv",
        "                  │",
        "                  ├── Pass/fail: novice_fk_grade < expert_fk_grade",
        "                  │    PASS → gradient confirmed",
        "                  │    FAIL → no measurable adaptation",
        "                  │",
        "                  └── Per-persona FK grade target comparison",
        "                       PERSONA_GRADE_TARGETS (from adaptive_vs_generic.py)",
        "                       Novice    avg_fk vs (6.0,  9.0)  → PASS / FAIL",
        "                       Interm.   avg_fk vs (10.0, 13.0) → PASS / FAIL",
        "                       Expert    avg_fk vs (14.0, 18.0) → PASS / FAIL",
        "                       Regul.    avg_fk vs (14.0, 18.0) → PASS / FAIL",
        "                       Execut.   avg_fk vs (8.0,  12.0) → PASS / FAIL",
    ]

    add_diagram(
        doc,
        "Readability Analysis — Data Flow",
        mmd,
        ascii_lines,
        "fig7_readability_flow.png",
        mmd_source_path=Path(__file__).parent / "fig7_readability_flow.mmd",
    )

    _heading(doc, "3.4 Outputs", level=2)
    add_teal_table(
        doc,
        ["File", "Description"],
        [
            ["results/readability_analysis_results.csv", "25 rows: query, persona, query_type + 7 readability metrics"],
            ["results/readability_analysis_summary.txt", "Per-persona averages, pass/fail headline, rank order by reading ease"],
        ],
        col_widths=[3.2, 4.5],
    )

    _heading(doc, "3.5 Code Reference", level=2)
    add_teal_table(
        doc,
        ["Step", "File", "Function / Symbol"],
        [
            ["Default queries", "eval/readability_analysis.py:40", "DEFAULT_QUERIES (5 items)"],
            ["Metric computation", "eval/readability_analysis.py:49", "compute_readability_metrics()"],
            ["Cache load / generate", "eval/readability_analysis.py:87", "run_readability_analysis()"],
            ["Directional pass/fail", "eval/readability_analysis.py", "_write_summary() — novice < expert FK"],
            ["Grade targets import", "eval/adaptive_vs_generic.py:56", "PERSONA_GRADE_TARGETS (imported)"],
            ["Per-persona target check", "eval/readability_analysis.py", "_write_summary() — FK GRADE TARGET COMPARISON section"],
        ],
        col_widths=[1.8, 2.8, 3.1],
    )


def build_section_4_compliance(doc: Document) -> None:
    doc.add_page_break()
    _heading(doc, "4. Eval 3: format_compliance")

    _heading(doc, "4.1 Purpose", level=2)
    doc.add_paragraph(
        "Verifies that adaptive prompt formatting instructions are obeyed in the "
        "generated response. The adaptive prompt tells the LLM to use numbered steps "
        "for PROCEDURE queries, markdown tables for COMPARISON queries, bullet points "
        "and key takeaways for NOVICE personas, and so on. "
        "This evaluation measures whether the LLM actually follows those instructions "
        "by regex-matching against 16 structural rules — the compliance score is "
        "purely mechanical and deterministic."
    )

    _heading(doc, "4.2 Inputs", level=2)
    add_teal_table(
        doc,
        ["Input", "Type", "Description", "Source"],
        [
            ["response text", "str", "Generated response text", "persona_responses.json"],
            ["query_type", "str", "Classified query type for this query", "classify_query() on each entry"],
            ["persona", "str", "Persona label for this response", "persona_responses.json"],
        ],
        col_widths=[1.5, 0.8, 2.8, 2.6],
    )

    _heading(doc, "4.3 Method", level=2)

    add_callout(
        doc,
        "Compliance Score Formula (eval/format_compliance.py:184–208)",
        """compliance_score = passed_applicable / total_applicable

where:
  passed_applicable = rules where check_rule() returns True
  total_applicable  = rules where check_rule() returns True or False
                      (None = not applicable, not counted)

check_rule() logic:
  1. Look up rule by name in FORMAT_RULES dict
  2. If condition_key == 'query_type' and query_type != condition_value: return None
  3. If condition_key == 'persona'     and persona     != condition_value: return None
  4. return bool(re.search(rule['pattern'], text, rule['flags']))""",
    )

    doc.add_paragraph(
        "FORMAT_RULES contains exactly 16 rules (eval/format_compliance.py:37–150), "
        "8 keyed on query_type and 8 keyed on persona. "
        "The full rule registry is shown below."
    )

    add_teal_table(
        doc,
        ["Rule ID", "Cond. Type", "Cond. Value", "Description", "Regex Pattern"],
        [
            ["procedure_has_numbered_steps", "query_type", "procedure", "Numbered steps present", r"^\s*\d+[\.\)]\s+\w (MULTILINE)"],
            ["comparison_has_table", "query_type", "comparison", "Markdown table present", r"\|.+\|.+\|"],
            ["eligibility_has_inclusion", "query_type", "eligibility", "Mentions inclusion", r"inclusion|include|eligible (IGNORECASE)"],
            ["eligibility_has_exclusion", "query_type", "eligibility", "Mentions exclusion", r"exclusion|exclude|ineligible (IGNORECASE)"],
            ["numerical_leads_with_number", "query_type", "numerical", "Leads with a number", r"^[\w\s]{0,60}?\d+"],
            ["timeline_contains_timeframe", "query_type", "timeline", "Contains explicit timeframe", r"\b\d+\s*(day|week|month|year)s?\b (IGNORECASE)"],
            ["safety_contains_severity", "query_type", "safety", "Mentions severity grade", r"\b(grade\s*[1-5]|mild|moderate|severe|serious)\b (IGNORECASE)"],
            ["compliance_cites_regulation", "query_type", "compliance", "Cites regulation", r"\b(FDA|EMA|ICH|GCP|21\s*CFR|ICH-GCP)\b (IGNORECASE)"],
            ["novice_has_bullet_points", "persona", "novice", "Bullet points used", r"^\s*[-•*]\s+\w (MULTILINE)"],
            ["novice_has_key_takeaway", "persona", "novice", "Key takeaway present", r"key\s+takeaway|📌 (IGNORECASE)"],
            ["novice_defines_terms", "persona", "novice", "Terms defined in parens", r"\(.{3,60}\)"],
            ["executive_has_summary", "persona", "executive", "Summary section present", r"executive\s+summary|in\s+brief|key\s+point (IGNORECASE)"],
            ["executive_has_recommendation", "persona", "executive", "Recommendation present", r"recommend|next\s+step|action\s+item|decision (IGNORECASE)"],
            ["expert_uses_technical_terms", "persona", "expert", "Technical terms used", r"\b(RECIST|ICH.GCP|21\s*CFR|SUVmax|iRECIST|CTCAE)\b (IGNORECASE)"],
            ["regulatory_cites_standard", "persona", "regulatory", "Standard/guidance cited", r"\b(FDA|EMA|ICH|GCP|CFR|guidance|compliance|audit)\b (IGNORECASE)"],
            ["intermediate_has_example", "persona", "intermediate", "Example included", r"\bfor\s+example\b|\be\.g\.\b|\bsuch\s+as\b (IGNORECASE)"],
        ],
        col_widths=[2.1, 0.9, 0.9, 1.6, 2.2],
    )

    mmd = """flowchart TD
  CACHE{"persona_responses.json\\nexists?"}
  GEN["run_persona_evaluation()"]
  RESP["Response Text + query_type + persona"]
  RULES["FORMAT_RULES\\n16 rules (37-150)"]
  CHK["check_rule()\\neval/format_compliance.py:153"]
  APPL{"Rule\\napplicable?"}
  SKIP["None — not counted"]
  RESULT["True / False"]
  SCORE["compute_compliance_score()\\npassed / total_applicable"]
  CSV["format_compliance_results.csv"]
  CACHE -->|"yes"| RESP
  CACHE -->|"no"| GEN
  GEN --> RESP
  RESP --> CHK
  RULES --> CHK
  CHK --> APPL
  APPL -->|"no"| SKIP
  APPL -->|"yes"| RESULT
  RESULT --> SCORE
  SKIP --> SCORE
  SCORE --> CSV"""

    ascii_lines = [
        "  persona_responses.json  (cache or freshly generated)",
        "      │",
        "      ▼",
        "  For each (query, persona) pair:",
        "      For each of 16 FORMAT_RULES:",
        "          │",
        "          ▼",
        "      check_rule()  (eval/format_compliance.py:153)",
        "          Applicability test:",
        "              condition_key == 'query_type' AND query_type != condition_value → None",
        "              condition_key == 'persona'     AND persona     != condition_value → None",
        "          If applicable:",
        "              bool(re.search(pattern, text, flags))",
        "              → True  (rule passed)",
        "              → False (rule failed)",
        "          │",
        "          ▼",
        "      compute_compliance_score()",
        "      score = passed_applicable / total_applicable",
        "          │",
        "          ▼",
        "  format_compliance_results.csv",
        "  format_compliance_summary.txt",
    ]

    add_diagram(
        doc,
        "Format Compliance — Rule Check Flow",
        mmd,
        ascii_lines,
        "fig8_format_compliance_flow.png",
        mmd_source_path=Path(__file__).parent / "fig8_format_compliance_flow.mmd",
    )

    _heading(doc, "4.4 Outputs", level=2)
    add_teal_table(
        doc,
        ["File", "Description"],
        [
            ["results/format_compliance_results.csv", "One row per (query, persona): 16 rule bool columns + applicable_rules_count + passed_rules_count + compliance_score"],
            ["results/format_compliance_summary.txt", "Overall score, per-persona breakdown, per-query-type breakdown, per-rule pass rate"],
        ],
        col_widths=[3.2, 4.5],
    )

    _heading(doc, "4.5 Code Reference", level=2)
    add_teal_table(
        doc,
        ["Step", "File", "Function / Symbol"],
        [
            ["Rule definitions", "eval/format_compliance.py:37", "FORMAT_RULES (16 entries)"],
            ["Rule applicability + check", "eval/format_compliance.py:153", "check_rule()"],
            ["Score computation", "eval/format_compliance.py:184", "compute_compliance_score()"],
            ["Main runner", "eval/format_compliance.py:215", "run_format_compliance()"],
            ["Summary report", "eval/format_compliance.py:305", "_write_summary()"],
        ],
        col_widths=[1.8, 2.8, 3.1],
    )


def build_section_5_adaptive_vs_generic(doc: Document) -> None:
    doc.add_page_break()
    _heading(doc, "5. Eval 4: adaptive_vs_generic")

    _heading(doc, "5.1 Purpose", level=2)
    doc.add_paragraph(
        "Head-to-head proof that the full adaptive pipeline "
        "(HybridRetriever + persona-aware prompt) outperforms a vanilla RAG baseline "
        "(semantic-only retrieval + static prompt). "
        "Both systems use the same PDF document, embedding model, and LLM — "
        "the only controlled variables are retrieval method and prompt construction. "
        "A query is an adaptive win when adaptive beats generic on at least 2 of 3 "
        "sub-metrics; the primary reported number is overall win rate."
    )

    _heading(doc, "5.2 Inputs", level=2)
    add_teal_table(
        doc,
        ["Input", "Type", "Description", "Source"],
        [
            ["COMPARISON_QUERIES", "list[dict]", "25 queries (5 per persona, covering all 9 query types)", "eval/adaptive_vs_generic.py:69"],
            ["document_path", "str", "Path to the source PDF (e.g. data/irc.pdf)", "CLI --document flag"],
            ["embedding_model", "str", "Default: S-PubMedBert-MS-MARCO", "run_adaptive_vs_generic() param"],
            ["llm_model", "str", "Default: src/config.py:DEFAULT_LLM_MODEL", "run_adaptive_vs_generic() param"],
        ],
        col_widths=[1.8, 0.8, 2.8, 2.3],
    )

    _heading(doc, "5.3 Method", level=2)

    doc.add_paragraph(
        "Generic baseline (build_generic_prompt(), eval/adaptive_vs_generic.py:195–214): "
        "retrieves top-5 chunks using ChromaDB semantic similarity search only, then "
        "builds a static, persona-free prompt: "
        "'You are a helpful assistant. Use the following context to answer the question.' "
        "No persona instructions, no query-type formatting directives."
    )

    doc.add_paragraph(
        "Adaptive system (build_adaptive_prompt() from src/prompts.py + HybridRetriever.retrieve()): "
        "retrieves top-5 chunks via RRF fusion of BM25 and dense vector results (k=60), "
        "then builds a persona-aware, query-type-aware prompt using ResponseStyler. "
        "Both systems invoke the same ChatOpenAI instance at temperature=0."
    )

    add_callout(
        doc,
        "Win Condition Logic (eval/adaptive_vs_generic.py)",
        """# Sub-metric 1: compliance
wins_compliance    = adaptive_compliance > generic_compliance

# Sub-metric 2: readability fit
wins_readability   = PERSONA_GRADE_TARGETS[persona][0] <= adaptive_fk_grade
                     <= PERSONA_GRADE_TARGETS[persona][1]

# Sub-metric 3: length adherence
adaptive_adherence = max(0.0, 1.0 - |adaptive_words - adaptive_target| / adaptive_target)
generic_adherence  = max(0.0, 1.0 - |generic_words  - 500|             / 500)
wins_length        = adaptive_adherence > generic_adherence

# Overall win (majority vote)
overall_win        = sum([wins_compliance, wins_readability, wins_length]) >= 2

# Primary metric
win_rate           = df["adaptive_overall_wins"].mean()

# Adaptive Advantage Score — composite mean of normalised deltas
compliance_delta_norm = (compliance_delta + 1.0) / 2.0       # [-1,1] → [0,1]
readability_fit_score = 1.0 if persona_appropriate else 0.0   # binary [0,1]
length_delta_norm     = ((adap_adhr - gen_adhr) + 1.0) / 2.0 # [-1,1] → [0,1]
adaptive_advantage_score = mean([compliance_delta_norm,
                                  readability_fit_score,
                                  length_delta_norm])""",
    )

    add_callout(
        doc,
        "RESOLVED — Adaptive Advantage Score Implemented",
        """RESOLVED: adaptive_advantage_score is now computed for every query row.

  Formula (all three components normalised to [0, 1]):
    compliance_delta_norm = (compliance_delta + 1.0) / 2.0
    readability_fit_score = 1.0 if persona_appropriate else 0.0
    length_delta_norm     = ((adaptive_adhr - generic_adhr) + 1.0) / 2.0
    adaptive_advantage_score = mean of the three components

  Added to:
    results/adaptive_vs_generic_results.csv   — column "adaptive_advantage_score"
    results/adaptive_vs_generic_summary.txt   — shown per-persona and overall

  Note: All scoring remains metric-based (regex compliance + textstat FK +
  word count). No LLM judge or faithfulness delta is included — these are
  documented as out-of-scope for this evaluation.""",
    )

    _heading(doc, "5.3.1 PERSONA_GRADE_TARGETS", level=2)
    doc.add_paragraph(
        "Flesch-Kincaid grade targets per persona "
        "(eval/adaptive_vs_generic.py:56–62). "
        "Used only in wins_readability sub-metric — not in readability_analysis.py."
    )
    add_teal_table(
        doc,
        ["Persona", "FK Grade Min", "FK Grade Max", "Rationale"],
        [
            ["novice", "6.0", "9.0", "6th–9th grade reading level (patient / junior nurse)"],
            ["intermediate", "10.0", "13.0", "10th–13th grade (clinical research coordinator)"],
            ["expert", "14.0", "18.0", "University/graduate level (physician / scientist)"],
            ["regulatory", "14.0", "18.0", "Graduate level — same as expert"],
            ["executive", "8.0", "12.0", "Accessible prose with brevity (C-suite)"],
        ],
        col_widths=[1.3, 1.1, 1.1, 4.2],
    )

    _heading(doc, "5.3.2 Length Targets", level=2)
    doc.add_paragraph(
        "Adaptive length targets come from response_config.max_length "
        "(get_response_config(), src/personas.py). "
        "The generic baseline always uses 500 words as its target."
    )
    add_teal_table(
        doc,
        ["Persona", "Adaptive Target (words)", "Generic Baseline Target"],
        [
            ["novice", "300", "500 (hardcoded)"],
            ["intermediate", "500", "500 (hardcoded)"],
            ["expert", "1000", "500 (hardcoded)"],
            ["regulatory", "800", "500 (hardcoded)"],
            ["executive", "250", "500 (hardcoded)"],
        ],
        col_widths=[1.5, 2.2, 4.0],
    )

    mmd = """sequenceDiagram
  participant CLI as run_eval.py
  participant DOC as PyPDFLoader
  participant EMB as create_embedder()
  participant DB as ChromaDB + BM25
  participant HYBRID as HybridRetriever
  participant ADAPT as build_adaptive_prompt()
  participant GENP as build_generic_prompt()
  participant LLM as ChatOpenAI (temp=0)
  participant SCORE as Scoring
  CLI->>DOC: load(document_path)
  DOC->>EMB: split chunks (size=800)
  EMB->>DB: Chroma.from_documents() + BM25Retriever.from_documents()
  loop 25 COMPARISON_QUERIES
    CLI->>HYBRID: retrieve(query, top_k=5) via RRF
    HYBRID->>ADAPT: build_adaptive_prompt(docs, query, config)
    ADAPT->>LLM: invoke(adaptive_prompt)
    LLM-->>SCORE: adaptive_response
    CLI->>DB: similarity_search(query, k=5)
    DB->>GENP: build_generic_prompt(context, query)
    GENP->>LLM: invoke(generic_prompt)
    LLM-->>SCORE: generic_response
    SCORE->>SCORE: wins_compliance wins_readability wins_length
    SCORE->>SCORE: overall_win = sum >= 2
  end
  SCORE-->>CLI: win_rate = mean(adaptive_overall_wins)"""

    ascii_lines = [
        "  run_eval.py",
        "      │",
        "      │  Setup (once)",
        "      ├── PyPDFLoader.load(document_path)",
        "      ├── RecursiveCharacterTextSplitter  chunk_size=800",
        "      ├── Chroma.from_documents(chunks)   → ChromaDB",
        "      └── BM25Retriever.from_documents(chunks)",
        "",
        "  For each of 25 COMPARISON_QUERIES:",
        "",
        "  ┌─────────────── ADAPTIVE ─────────────────┐",
        "  │ HybridRetriever.retrieve(query, top_k=5) │",
        "  │ (BM25 + ChromaDB fused via RRF k=60)     │",
        "  │ build_adaptive_prompt(docs, q, config)   │",
        "  │ ChatOpenAI.invoke(adaptive_prompt)        │",
        "  └──────────────────────────────────────────┘",
        "                     │",
        "  ┌─────────────── GENERIC ──────────────────┐",
        "  │ ChromaDB.similarity_search(query, k=5)   │",
        "  │ build_generic_prompt(context, query)      │",
        "  │ ChatOpenAI.invoke(generic_prompt)         │",
        "  └──────────────────────────────────────────┘",
        "                     │",
        "  ┌─────────────── SCORING ──────────────────┐",
        "  │ wins_compliance  = adap_cmpl > gen_cmpl  │",
        "  │ wins_readability = FK grade in target     │",
        "  │ wins_length      = adap_adhr > gen_adhr  │",
        "  │ overall_win      = sum(wins) >= 2         │",
        "  │                                           │",
        "  │ compliance_delta_norm = (delta+1)/2       │",
        "  │ readability_fit_score = 1.0 / 0.0         │",
        "  │ length_delta_norm     = (delta+1)/2       │",
        "  │ adaptive_advantage_score = mean of above  │",
        "  └──────────────────────────────────────────┘",
        "                     │",
        "  win_rate = mean(adaptive_overall_wins)     ← Primary binary metric",
        "  avg_adv  = mean(adaptive_advantage_score)  ← Composite score [0,1]",
    ]

    add_diagram(
        doc,
        "Adaptive vs Generic — Head-to-Head Sequence",
        mmd,
        ascii_lines,
        "fig9_adaptive_vs_generic_flow.png",
        mmd_source_path=Path(__file__).parent / "fig9_adaptive_vs_generic_flow.mmd",
    )

    _heading(doc, "5.4 Outputs", level=2)
    add_teal_table(
        doc,
        ["File", "Description"],
        [
            ["results/adaptive_vs_generic_results.csv", "One row per query: all metric columns + adaptive_overall_wins"],
            ["results/adaptive_vs_generic_detailed.json", "Full JSON with response text, generation latency per query"],
            ["results/adaptive_vs_generic_summary.txt", "Win rate, compliance delta, readability fit — by persona and query type"],
        ],
        col_widths=[3.2, 4.5],
    )

    _heading(doc, "5.5 Code Reference", level=2)
    add_teal_table(
        doc,
        ["Step", "File", "Function / Symbol"],
        [
            ["Query set", "eval/adaptive_vs_generic.py:69", "COMPARISON_QUERIES (25 items)"],
            ["Grade targets", "eval/adaptive_vs_generic.py:56", "PERSONA_GRADE_TARGETS"],
            ["Generic prompt", "eval/adaptive_vs_generic.py:195", "build_generic_prompt()"],
            ["Adaptive prompt", "src/prompts.py:258", "build_adaptive_prompt()"],
            ["Hybrid retrieval", "src/retrieval.py", "HybridRetriever.retrieve()"],
            ["Win scoring", "eval/adaptive_vs_generic.py", "wins_compliance / wins_readability / wins_length"],
            ["Advantage score", "eval/adaptive_vs_generic.py", "adaptive_advantage_score = mean(3 normalised deltas)"],
            ["Main runner", "eval/adaptive_vs_generic.py", "run_adaptive_vs_generic()"],
            ["Summary report", "eval/adaptive_vs_generic.py", "_write_summary() — includes per-persona advantage score"],
        ],
        col_widths=[1.8, 2.8, 3.1],
    )


def build_section_6_scoring_summary(doc: Document) -> None:
    doc.add_page_break()
    _heading(doc, "6. Scoring Summary")

    doc.add_paragraph(
        "The table below summarises all four evaluation modules — their primary inputs, "
        "method, output metric, and the pass condition reported in each summary file."
    )

    add_teal_table(
        doc,
        ["Eval", "Primary Input", "Method", "Output Metric", "Pass Condition"],
        [
            [
                "1. classification_accuracy",
                "45 labeled queries + 25 expertise profiles",
                "Regex classifier + detect_user_type() vs ground-truth",
                "Combined accuracy = 0.6×QT_acc + 0.4×expertise_acc",
                "No auto-threshold — human review",
            ],
            [
                "2. readability_analysis",
                "25 persona responses (5q × 5p)",
                "textstat 7-metric suite + PERSONA_GRADE_TARGETS comparison",
                "FK grade per persona + per-persona PASS/FAIL vs target",
                "novice_fk_grade < expert_fk_grade AND per-persona target check",
            ],
            [
                "3. format_compliance",
                "25 persona responses + query_type",
                "16 regex rules; passed_applicable / total",
                "compliance_score ∈ [0, 1]",
                "No auto-threshold — per-rule pass rate in report",
            ],
            [
                "4. adaptive_vs_generic",
                "25 comparison queries + PDF",
                "3 sub-metric majority vote + composite advantage score",
                "win_rate + adaptive_advantage_score ∈ [0, 1]",
                "win_rate > 0.5 + advantage_score > 0.5",
            ],
        ],
        col_widths=[1.7, 1.5, 1.9, 1.5, 1.1],
    )

    doc.add_paragraph(
        "If any metric fails: "
        "(1) For classification accuracy, review the confusion matrix to identify "
        "which query types are most confused and refine their regex patterns in "
        "src/query_classifier.py. "
        "(2) For readability, inspect the per-persona FK averages — if NOVICE responses "
        "have high FK grades, the persona prompt instructions in src/prompts.py need "
        "stronger simplification directives. "
        "(3) For compliance, review per-rule pass rates — rules near 0% indicate the LLM "
        "is ignoring specific formatting directives; strengthen those instructions in "
        "src/prompts.py. "
        "(4) For adaptive win rate below 50%, the retrieval or prompt advantage is "
        "insufficient; consider strengthening the hybrid retrieval weights or "
        "the persona-specific prompt instructions."
    )


def build_section_7_running(doc: Document) -> None:
    doc.add_page_break()
    _heading(doc, "7. Running the Suite")

    _heading(doc, "7.1 CLI Commands (run_eval.py)", level=2)

    add_callout(
        doc,
        "Command Reference",
        """# Run all evaluations (adaptive suite + benchmarks)
python run_eval.py --document data/irc.pdf --all

# Run individual adaptive evaluations
python run_eval.py --classify                                # no document required
python run_eval.py --document data/irc.pdf --readability
python run_eval.py --document data/irc.pdf --compliance
python run_eval.py --document data/irc.pdf --adaptive-vs-generic

# Aggregate metrics from existing result CSVs
python run_eval.py --metrics""",
    )

    _heading(doc, "7.2 Direct Python Imports", level=2)

    add_callout(
        doc,
        "Module-Level Imports",
        """from eval.classification_accuracy import run_classification_accuracy
from eval.readability_analysis    import run_readability_analysis
from eval.format_compliance       import run_format_compliance
from eval.adaptive_vs_generic     import run_adaptive_vs_generic
from eval.metrics                 import calculate_all_metrics

# Eval 1 — no arguments needed
df_cls = run_classification_accuracy()

# Evals 2, 3, 4 — document path required for fresh generation
df_read = run_readability_analysis(document_path="data/irc.pdf")
df_comp = run_format_compliance(document_path="data/irc.pdf")
df_avsg = run_adaptive_vs_generic(document_path="data/irc.pdf")""",
    )

    _heading(doc, "7.3 Streamlit UI Trigger", level=2)
    doc.add_paragraph(
        "The evaluation suite can also be triggered from the Streamlit web interface. "
        "Navigate to Tab 4 — 'Benchmark Evals' — and click the 'Run Selected Evaluations' "
        "button. The relevant code is in app.py around line 1165. "
        "Results are written to the results/ directory and displayed inline in the UI."
    )

    _heading(doc, "7.4 Output Directory", level=2)
    add_teal_table(
        doc,
        ["Eval", "Output Files"],
        [
            ["classification_accuracy", "results/classification_accuracy_results.csv, classification_confusion_matrix.csv, classification_accuracy_summary.txt"],
            ["readability_analysis", "results/readability_analysis_results.csv, readability_analysis_summary.txt"],
            ["format_compliance", "results/format_compliance_results.csv, format_compliance_summary.txt"],
            ["adaptive_vs_generic", "results/adaptive_vs_generic_results.csv, adaptive_vs_generic_detailed.json, adaptive_vs_generic_summary.txt"],
            ["metrics aggregator", "Reads from all of the above CSVs — no new files written"],
        ],
        col_widths=[1.8, 5.9],
    )


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def main() -> None:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # Default paragraph font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    _add_page_numbers(doc)

    add_cover(doc)
    add_toc(doc)

    build_section_1_overview(doc)
    build_section_2_classification(doc)
    build_section_3_readability(doc)
    build_section_4_compliance(doc)
    build_section_5_adaptive_vs_generic(doc)
    build_section_6_scoring_summary(doc)
    build_section_7_running(doc)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
