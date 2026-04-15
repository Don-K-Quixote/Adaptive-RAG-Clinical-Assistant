"""
Generate docs/rag-system/adaptive_evaluation.docx for the Clinical RAG Assistant.

Produces a professionally formatted technical document covering all four
adaptive evaluation modules: classification_accuracy, readability_analysis,
format_compliance, and adaptive_vs_generic.

All figures are embedded as ASCII art (monospace callout blocks), making the
document fully self-contained without requiring external PNG files or mmdc.

Usage:
    python scripts/gen_adaptive_docx.py
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Colour palette
# ---------------------------------------------------------------------------
NAVY = RGBColor(0x11, 0x22, 0x40)
TEAL = RGBColor(0x0E, 0xA5, 0xC9)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEX_TEAL = "0EA5C9"
HEX_NAVY = "112240"
HEX_LIGHT_BLUE = "E0F2FE"
HEX_AMBER = "FEF3C7"
HEX_RED = "DC2626"

OUTPUT_PATH = Path("docs/rag-system/adaptive_evaluation.docx")

# ---------------------------------------------------------------------------
# ASCII figures
# ---------------------------------------------------------------------------

ASCII_SUITE_OVERVIEW = """\
  run_eval.py  (main -- argparse CLI)
      |
      +---> --classify              Eval 1: run_classification_accuracy()
      |                             No LLM. No document required. < 1 sec.
      |
      +---> --readability           Eval 2: run_readability_analysis()
      |                             Reads persona_responses.json cache or
      |                             calls run_persona_evaluation()
      |
      +---> --compliance            Eval 3: run_format_compliance()
      |                             Same cache reuse as Eval 2
      |
      +---> --adaptive-vs-generic   Eval 4: run_adaptive_vs_generic()
      |                             Independent pipeline. Requires document.
      |
      +---> --metrics               eval/metrics.py: calculate_all_metrics()
                                    Aggregates from existing result CSVs

  Shared Infrastructure
  eval/persona_evaluation.py: run_persona_evaluation()
      Used by Evals 2 and 3 when persona_responses.json is absent"""

ASCII_CLASSIFICATION = """\
  [LABELED_QUERY_DATASET]
  45 queries -- 9 QueryType labels x 5 queries each
  (hardcoded in eval/classification_accuracy.py:28)
      |
      v
  [QueryClassifier.classify()]       src/query_classifier.py:162
  Priority-ordered regex pattern matching
      |
      +-------> [predicted_type]     result label for each query
      |
      +-------> [confidence_score]   min(1.0, 0.5 + matches x 0.2)
      |
      v
  [run_classification_accuracy()]    eval/classification_accuracy.py:266
  overall_accuracy = correct / 45
  confusion_matrix via pd.crosstab(expected, predicted)
      |
      +-------> [classification_accuracy_results.csv]
      |         one row per query: query, expected, predicted, correct
      |
      +-------> [classification_confusion_matrix.csv / .txt]
      |         9x9 confusion matrix
      |
      +-------> [classification_accuracy_summary.txt]
                overall accuracy, per-type accuracy, confusion matrix"""

ASCII_READABILITY = """\
  [persona_responses.json]
  25 responses: 5 queries x 5 personas
  (NOVICE / INTERMEDIATE / EXPERT / REGULATORY / EXECUTIVE)
      |
      v
  [compute_readability_metrics()]    eval/readability_analysis.py:49
  using the textstat library
      |
      +-------> flesch_reading_ease       0-100; higher = simpler
      |
      +-------> flesch_kincaid_grade      US school grade level
      |         ^--- Pass/Fail pivot: novice_fk_grade < expert_fk_grade
      |
      +-------> gunning_fog               years of education needed
      |
      +-------> word_count / sentence_count / difficult_words
      |
      +-------> avg_sentence_length
      |
      v
  25 rows written (one per query-persona pair)
      |
      +-------> [readability_analysis_results.csv]
      |         query, persona, query_type, 7 metrics per row
      |
      +-------> [readability_analysis_summary.txt]
                per-persona averages, pass/fail headline"""

ASCII_COMPLIANCE = """\
  [persona_responses.json] + [query_type labels]
  25 (query, persona) pairs
      |
      v
  [check_rule(rule_id, text, query_type, persona)]
  eval/format_compliance.py:153
  returns True / False / None (None = not applicable)
      |
      +-------> QUERY-TYPE RULES (8):
      |         procedure_has_numbered_steps  comparison_has_table
      |         eligibility_has_inclusion     eligibility_has_exclusion
      |         numerical_leads_with_number   timeline_contains_timeframe
      |         safety_contains_severity      compliance_cites_regulation
      |
      +-------> PERSONA RULES (8):
      |         novice_has_bullet_points      novice_has_key_takeaway
      |         novice_defines_terms          executive_has_summary
      |         executive_has_recommendation  expert_uses_technical_terms
      |         regulatory_cites_standard     intermediate_has_example
      |
      v
  [compute_compliance_score()]       eval/format_compliance.py:184
  score = passed_applicable / total_applicable
  (None responses not counted in denominator)
      |
      +-------> [format_compliance_results.csv]
      |         one row per (query, persona), 16 rule bool columns
      |
      +-------> [format_compliance_summary.txt]
                overall score, per-persona and per-rule breakdown"""

ASCII_ADAPTIVE_VS_GENERIC = """\
  [COMPARISON_QUERIES: 25 queries, 5 personas x 5 each]
  (eval/adaptive_vs_generic.py:69)
      |
      +-------> GENERIC BASELINE
      |         build_generic_prompt()       eval/adaptive_vs_generic.py:195
      |         ChromaDB semantic-only retrieval
      |         Fixed 500-word response target
      |
      +-------> ADAPTIVE SYSTEM
      |         build_adaptive_prompt()      src/prompts.py:258
      |         HybridRetriever RRF fusion   src/retrieval.py
      |         Persona-specific word limits
      |
      v
  [Win Scoring -- majority vote]     eval/adaptive_vs_generic.py:351
      |
      +-------> wins_compliance:
      |         adaptive_compliance > generic_compliance
      |
      +-------> wins_readability:
      |         PERSONA_GRADE_TARGETS[persona][0]
      |         <= adaptive_fk_grade
      |         <= PERSONA_GRADE_TARGETS[persona][1]
      |
      +-------> wins_length:
      |         adaptive_adherence > generic_adherence
      |         adherence = max(0, 1 - abs(words - target) / target)
      |
      v
  overall_win = 2 or more sub-metrics won by adaptive
  win_rate    = mean(adaptive_overall_wins)
  pass cond.  = win_rate > 0.5
      |
      +-------> [adaptive_vs_generic_results.csv]
      |         one row per query: metrics + adaptive_overall_wins bool
      |
      +-------> [adaptive_vs_generic_detailed.json]
      |         full response text, generation times, all sub-metrics
      |
      +-------> [adaptive_vs_generic_summary.txt]
                win rate, compliance delta, readability summary"""

# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _heading1(doc: Document, text: str):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(16)
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(6)
    return h


def _heading2(doc: Document, text: str):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(13)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    return h


def _para(doc: Document, text: str = "", size_pt: float = 10.5):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size_pt)
    p.paragraph_format.space_after = Pt(6)
    return p


def _mono_para(doc: Document, text: str, size_pt: float = 9):
    """Monospace paragraph for code references."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(size_pt)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    return p


def _add_table(
    doc: Document, headers: list[str], rows: list[list[str]]
):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        _set_cell_bg(hdr_cells[i], HEX_TEAL)
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(10)

    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = str(val)
            for para in row_cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9.5)

    doc.add_paragraph()
    return table


def _callout(doc: Document, text: str, label: str = "", bg_hex: str = HEX_LIGHT_BLUE):
    """Bordered callout box — light blue for formulas, amber for warnings."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        bdr = OxmlElement(f"w:{side}")
        bdr.set(qn("w:val"), "single")
        bdr.set(qn("w:sz"), "4")
        bdr.set(qn("w:space"), "4")
        bdr.set(qn("w:color"), "0EA5C9")
        pBdr.append(bdr)
    pPr.append(pBdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), bg_hex)
    pPr.append(shd)
    if label:
        lr = p.add_run(f"{label}\n")
        lr.bold = True
        lr.font.size = Pt(9)
        lr.font.color.rgb = NAVY
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(8.5)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(8)
    return p


def _ascii_figure(doc: Document, ascii_art: str, caption: str):
    """Embed an ASCII diagram in a monospace callout, then add a caption."""
    _callout(doc, ascii_art, label="", bg_hex="F8FAFC")
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = NAVY
    doc.add_paragraph()


def _gap_warning(doc: Document, title: str, body: str):
    """Red left-bar callout for known gaps / discrepancies."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), HEX_RED)
    pBdr.append(left)
    pPr.append(pBdr)
    p.paragraph_format.left_indent = Inches(0.25)
    r1 = p.add_run(f"[!] {title}: ")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r2 = p.add_run(body)
    r2.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(8)


def _page_break(doc: Document):
    doc.add_page_break()


def _add_footer(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _add_toc(doc: Document):
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar)
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run._r.append(fldChar2)
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar3)
    note = doc.add_paragraph(
        "[ Press Ctrl+A, then F9 in Word to update this Table of Contents ]"
    )
    note.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# ---------------------------------------------------------------------------
# Document sections
# ---------------------------------------------------------------------------


def build_cover(doc: Document):
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Adaptive Evaluation Suite")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Clinical RAG Assistant -- Technical Documentation")
    r2.font.size = Pt(16)
    r2.font.color.rgb = TEAL

    doc.add_paragraph()

    meta = [
        ("System", "Adaptive RAG Clinical Assistant"),
        ("Document", "Adaptive Evaluation Suite"),
        ("Version", "1.0"),
        ("Date", str(date.today())),
        ("Classification", "Internal Technical Reference"),
    ]
    for label, val in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = p.add_run(f"{label}:  ")
        lr.bold = True
        lr.font.size = Pt(11)
        vr = p.add_run(val)
        vr.font.size = Pt(11)

    _page_break(doc)


def build_toc(doc: Document):
    _heading1(doc, "Table of Contents")
    _add_toc(doc)
    _page_break(doc)


def build_section1(doc: Document):
    _heading1(doc, "1. Suite Overview")

    _para(doc, (
        "The Adaptive Evaluation Suite contains four complementary evaluation modules that "
        "together prove the Clinical RAG Assistant's adaptive pipeline produces measurably "
        "better outputs than a non-adaptive baseline. Each module is independently runnable "
        "via the run_eval.py CLI or the Streamlit Benchmark Evals tab."
    ))

    _ascii_figure(doc, ASCII_SUITE_OVERVIEW, "Figure 1 -- Evaluation Suite: Module Relationships")

    _para(doc, (
        "Eval 1 (classification_accuracy) requires no document or LLM -- it tests the "
        "regex-based query classifier against a hardcoded labelled dataset. "
        "Evals 2 and 3 (readability_analysis and format_compliance) reuse the "
        "persona_responses.json cache produced by eval/persona_evaluation.py to avoid "
        "redundant LLM calls. "
        "Eval 4 (adaptive_vs_generic) runs a fresh head-to-head comparison against a "
        "generic baseline for each query."
    ))

    _page_break(doc)


def build_section2(doc: Document):
    _heading1(doc, "2. Eval 1: classification_accuracy")

    _heading2(doc, "2.1 Purpose")
    _para(doc, (
        "Measures how accurately the regex-based query classifier assigns one of 9 QueryType "
        "labels to incoming queries. Accurate classification is a prerequisite for all "
        "adaptive behaviour: the system uses the query type to select retrieval strategy, "
        "prompt style, and response format. A misclassified query receives the wrong "
        "adaptive treatment."
    ))

    _heading2(doc, "2.2 Inputs")
    _add_table(doc,
        ["Input", "Type", "Description", "Source"],
        [
            ["query", "str", "Natural language query string", "LABELED_QUERY_DATASET"],
            ["expected_type", "str", "Ground-truth QueryType label", "LABELED_QUERY_DATASET"],
            ["LABELED_QUERY_DATASET", "list[dict]", "45 queries, 5 per QueryType (hardcoded)", "eval/classification_accuracy.py:28"],
        ]
    )

    _heading2(doc, "2.3 Method")
    _para(doc, (
        "The 9 QueryType categories are defined as an enum in src/query_classifier.py:24-35: "
        "DEFINITION, PROCEDURE, COMPLIANCE, COMPARISON, NUMERICAL, TIMELINE, SAFETY, "
        "ELIGIBILITY, GENERAL."
    ))
    _para(doc, (
        "Classification algorithm (QueryClassifier.classify(), src/query_classifier.py:162-194): "
        "Priority-ordered regex pattern matching -- not ML or LLM-based. "
        "The first matching pattern wins; GENERAL is the default fallback."
    ))
    _para(doc, (
        "Confidence scoring (QueryClassifier.get_confidence(), src/query_classifier.py:197-217): "
        "confidence = min(1.0, 0.5 + matches x 0.2), where matches is the count of regex "
        "patterns that triggered for the query."
    ))

    _callout(doc, (
        "Accuracy Formulae\n\n"
        "Overall Accuracy  = correct_predictions / 45\n"
        "Per-type Accuracy = correct / 5   (per QueryType, 5 queries each)"
    ), label="Accuracy Formulae")

    _gap_warning(doc,
        "Brief vs Code Discrepancy",
        "Combined Accuracy (w1 x Query Type Acc + w2 x Expertise Acc) is described in "
        "the project brief but is NOT computed anywhere in classification_accuracy.py. "
        "No expertise classification is performed -- only query type accuracy is measured."
    )

    _ascii_figure(doc, ASCII_CLASSIFICATION,
                  "Figure 2 -- Classification Accuracy: Query Flow")

    _heading2(doc, "2.4 Outputs")
    _add_table(doc,
        ["File", "Description"],
        [
            ["results/classification_accuracy_results.csv",
             "One row per query: query, expected_type, predicted_type, correct (bool), confidence_score"],
            ["results/classification_confusion_matrix.csv / .txt",
             "9x9 confusion matrix via pd.crosstab(expected_type, predicted_type)"],
            ["results/classification_accuracy_summary.txt",
             "Human-readable report: overall accuracy, per-type accuracy, and confusion matrix"],
        ]
    )
    _para(doc, (
        "No action threshold is defined in the code. There is no automatic pass/fail gate "
        "on overall accuracy -- the result is reported for human review."
    ))

    _heading2(doc, "2.5 Code Reference")
    _add_table(doc,
        ["Step", "File", "Function / Symbol"],
        [
            ["Dataset", "eval/classification_accuracy.py:28", "LABELED_QUERY_DATASET"],
            ["Query types enum", "src/query_classifier.py:24", "QueryType (9 values)"],
            ["Classification", "src/query_classifier.py:162", "QueryClassifier.classify()"],
            ["Confidence scoring", "src/query_classifier.py:197", "QueryClassifier.get_confidence()"],
            ["Accuracy computation", "eval/classification_accuracy.py:266", "run_classification_accuracy()"],
            ["Confusion matrix", "eval/classification_accuracy.py:311", "pd.crosstab(expected_type, predicted_type)"],
            ["Summary report", "eval/classification_accuracy.py:335", "_write_summary()"],
        ]
    )

    _page_break(doc)


def build_section3(doc: Document):
    _heading1(doc, "3. Eval 2: readability_analysis")

    _heading2(doc, "3.1 Purpose")
    _para(doc, (
        "Proves that NOVICE responses are measurably simpler than EXPERT responses using "
        "established readability metrics. Clinical rationale: an overly complex response to "
        "a NOVICE coordinator may cause clinical protocol misinterpretation; an overly simple "
        "response to an EXPERT principal investigator wastes time and omits necessary detail."
    ))

    _heading2(doc, "3.2 Inputs")
    _add_table(doc,
        ["Input", "Type", "Description", "Source"],
        [
            ["response text", "str", "Generated response for a (query, persona) pair", "persona_responses.json cache"],
            ["persona", "str", "One of: novice, intermediate, expert, regulatory, executive", "persona_responses.json"],
            ["query", "str", "The original query string", "DEFAULT_QUERIES (5 items)"],
            ["embedding_model", "str", "Embedding model identifier (default: S-PubMedBert-MS-MARCO)", "run_readability_analysis() param"],
            ["llm_model", "str", "LLM model (default: src/config.py:DEFAULT_LLM_MODEL)", "run_readability_analysis() param"],
        ]
    )

    _heading2(doc, "3.3 Method")
    _para(doc, (
        "Seven metrics are computed per (query, persona) pair by compute_readability_metrics() "
        "(eval/readability_analysis.py:49) using the textstat library. "
        "With 5 queries and 5 personas this produces 25 data rows."
    ))
    _add_table(doc,
        ["Metric", "textstat Function", "Interpretation"],
        [
            ["flesch_reading_ease", "textstat.flesch_reading_ease()", "0-100; higher = simpler. <30 = very difficult"],
            ["flesch_kincaid_grade", "textstat.flesch_kincaid_grade()", "US school grade level. Used for pass/fail pivot"],
            ["gunning_fog", "textstat.gunning_fog()", "Years of formal education needed to understand"],
            ["word_count", "len(text.split())", "Total words in response"],
            ["sentence_count", "textstat.sentence_count()", "Total sentences"],
            ["difficult_words", "textstat.difficult_words()", "Count of words with 3+ syllables"],
            ["avg_sentence_length", "word_count / max(sentence_count, 1)", "Mean words per sentence"],
        ]
    )

    _callout(doc, (
        "Pass/Fail Criterion  (eval/readability_analysis.py:196-209)\n\n"
        "Pass condition:  novice_fk_grade  <  expert_fk_grade\n\n"
        "This is a directional ordering check -- not a threshold check.\n"
        "The evaluation confirms that NOVICE responses use a lower Flesch-Kincaid\n"
        "grade level than EXPERT responses, proving the adaptive persona system\n"
        "produces a measurable readability gradient."
    ), label="Pass/Fail Criterion")

    _gap_warning(doc,
        "Brief vs Code Discrepancy",
        "The project brief suggests readability_analysis maps FK scores to per-persona "
        "target grade ranges. PERSONA_GRADE_TARGETS lives in eval/adaptive_vs_generic.py:56-62 "
        "and is NOT used in readability_analysis.py. No target threshold comparison is "
        "performed here. Readability eval is directional only."
    )

    _ascii_figure(doc, ASCII_READABILITY,
                  "Figure 3 -- Readability Analysis: Data Flow")

    _heading2(doc, "3.4 Outputs")
    _add_table(doc,
        ["File", "Description"],
        [
            ["results/readability_analysis_results.csv",
             "25 rows: query, persona, query_type + 7 readability metrics per row"],
            ["results/readability_analysis_summary.txt",
             "Per-persona averages, pass/fail headline (novice_fk < expert_fk)"],
        ]
    )

    _heading2(doc, "3.5 Code Reference")
    _add_table(doc,
        ["Step", "File", "Function / Symbol"],
        [
            ["Default queries", "eval/readability_analysis.py:40", "DEFAULT_QUERIES (5 items)"],
            ["Metric computation", "eval/readability_analysis.py:49", "compute_readability_metrics()"],
            ["Cache load / generate", "eval/readability_analysis.py:87", "run_readability_analysis()"],
            ["Pass/fail test", "eval/readability_analysis.py:196", "_write_summary() -- novice < expert FK"],
            ["Grade targets (NOT used here)", "eval/adaptive_vs_generic.py:56", "PERSONA_GRADE_TARGETS"],
        ]
    )

    _page_break(doc)


def build_section4(doc: Document):
    _heading1(doc, "4. Eval 3: format_compliance")

    _heading2(doc, "4.1 Purpose")
    _para(doc, (
        "Verifies that adaptive prompt formatting instructions are obeyed in the generated "
        "response. The adaptive prompt tells the LLM to use numbered steps for PROCEDURE "
        "queries, markdown tables for COMPARISON queries, and bullet points for NOVICE "
        "personas. This evaluation confirms those instructions actually produce compliant "
        "output."
    ))

    _heading2(doc, "4.2 Inputs")
    _add_table(doc,
        ["Input", "Type", "Description", "Source"],
        [
            ["response text", "str", "Generated response text", "persona_responses.json"],
            ["query_type", "str", "Classified query type for this query", "classify_query() on each entry"],
            ["persona", "str", "Persona label for this response", "persona_responses.json"],
        ]
    )

    _heading2(doc, "4.3 Method")
    _callout(doc, (
        "Compliance Score Formula  (eval/format_compliance.py:184-208)\n\n"
        "compliance_score = passed_applicable / total_applicable\n\n"
        "where:\n"
        "  passed_applicable = rules where check_rule() returns True\n"
        "  total_applicable  = rules where check_rule() returns True or False\n"
        "                      (None = not applicable, not counted)\n\n"
        "check_rule() logic:\n"
        "  1. Look up rule by name in FORMAT_RULES dict\n"
        "  2. If condition_key == 'query_type' and query_type != condition_value: return None\n"
        "  3. If condition_key == 'persona'     and persona     != condition_value: return None\n"
        "  4. return bool(re.search(rule['pattern'], text, rule['flags']))"
    ), label="Compliance Score Formula")

    _para(doc, (
        "FORMAT_RULES contains exactly 16 rules (eval/format_compliance.py:37-150), "
        "8 keyed on query_type and 8 keyed on persona. The full rule registry is shown below:"
    ))
    _add_table(doc,
        ["Rule ID", "Cond. Type", "Cond. Value", "Description", "Regex Pattern"],
        [
            ["procedure_has_numbered_steps", "query_type", "procedure", "Numbered steps present", r"^\s*\d+[\.\)]\s+\w (MULTILINE)"],
            ["comparison_has_table", "query_type", "comparison", "Markdown table present", r"\|.+\|.+\|"],
            ["eligibility_has_inclusion", "query_type", "eligibility", "Mentions inclusion", "inclusion|include|eligible (IGNORECASE)"],
            ["eligibility_has_exclusion", "query_type", "eligibility", "Mentions exclusion", "exclusion|exclude|ineligible (IGNORECASE)"],
            ["numerical_leads_with_number", "query_type", "numerical", "Leads with a number", r"^[\w\s]{0,60}?\d+"],
            ["timeline_contains_timeframe", "query_type", "timeline", "Contains explicit timeframe", r"\b\d+\s*(day|week|month|year)s?\b (IGNORECASE)"],
            ["safety_contains_severity", "query_type", "safety", "Mentions severity grade", r"\b(grade\s*[1-5]|mild|moderate|severe|serious)\b"],
            ["compliance_cites_regulation", "query_type", "compliance", "Cites regulation", r"\b(FDA|EMA|ICH|GCP|21\s*CFR|ICH-GCP)\b (IGNORECASE)"],
            ["novice_has_bullet_points", "persona", "novice", "Bullet points used", r"^\s*[-*]\s+\w (MULTILINE)"],
            ["novice_has_key_takeaway", "persona", "novice", "Key takeaway present", "key takeaway (IGNORECASE)"],
            ["novice_defines_terms", "persona", "novice", "Terms defined in parens", r"\(.{3,60}\)"],
            ["executive_has_summary", "persona", "executive", "Summary section present", "executive summary|in brief|key point"],
            ["executive_has_recommendation", "persona", "executive", "Recommendation present", "recommend|next step|action item|decision"],
            ["expert_uses_technical_terms", "persona", "expert", "Technical terms used", r"\b(RECIST|ICH.GCP|21\s*CFR|SUVmax|iRECIST)\b"],
            ["regulatory_cites_standard", "persona", "regulatory", "Standard/guidance cited", r"\b(FDA|EMA|ICH|GCP|CFR|guidance|compliance)\b"],
            ["intermediate_has_example", "persona", "intermediate", "Example included", r"for example|e\.g\.|such as (IGNORECASE)"],
        ]
    )

    _ascii_figure(doc, ASCII_COMPLIANCE,
                  "Figure 4 -- Format Compliance: Rule Check Flow")

    _heading2(doc, "4.4 Outputs")
    _add_table(doc,
        ["File", "Description"],
        [
            ["results/format_compliance_results.csv",
             "One row per (query, persona): 16 rule bool columns + compliance_score"],
            ["results/format_compliance_summary.txt",
             "Overall score, per-persona breakdown, per-rule pass rate"],
        ]
    )

    _heading2(doc, "4.5 Code Reference")
    _add_table(doc,
        ["Step", "File", "Function / Symbol"],
        [
            ["Rule definitions", "eval/format_compliance.py:37", "FORMAT_RULES (16 entries)"],
            ["Rule applicability + check", "eval/format_compliance.py:153", "check_rule()"],
            ["Score computation", "eval/format_compliance.py:184", "compute_compliance_score()"],
            ["Main runner", "eval/format_compliance.py:215", "run_format_compliance()"],
            ["Summary report", "eval/format_compliance.py:305", "_write_summary()"],
        ]
    )

    _page_break(doc)


def build_section5(doc: Document):
    _heading1(doc, "5. Eval 4: adaptive_vs_generic")

    _heading2(doc, "5.1 Purpose")
    _para(doc, (
        "Head-to-head proof that the full adaptive pipeline (HybridRetriever + persona-aware "
        "prompt) outperforms a vanilla RAG baseline (semantic-only retrieval + a fixed generic "
        "prompt) across three sub-metrics: format compliance, persona-appropriate readability, "
        "and word-count length adherence."
    ))

    _heading2(doc, "5.2 Inputs")
    _add_table(doc,
        ["Input", "Type", "Description", "Source"],
        [
            ["COMPARISON_QUERIES", "list[dict]", "25 queries (5 per persona, covering all 5 QueryTypes)", "eval/adaptive_vs_generic.py:69"],
            ["document_path", "str", "Path to the source PDF (e.g. data/irc.pdf)", "CLI --document flag"],
            ["embedding_model", "str", "Default: S-PubMedBert-MS-MARCO", "run_adaptive_vs_generic() param"],
            ["llm_model", "str", "Default: src/config.py:DEFAULT_LLM_MODEL", "run_adaptive_vs_generic() param"],
        ]
    )

    _heading2(doc, "5.3 Method")
    _para(doc, (
        "Generic baseline (build_generic_prompt(), eval/adaptive_vs_generic.py:195-214): "
        "retrieves top-5 chunks using ChromaDB semantic similarity search only, "
        "constructs a prompt with no persona instructions and no query-type adaptation, "
        "targets a fixed 500-word response."
    ))
    _para(doc, (
        "Adaptive system (build_adaptive_prompt() from src/prompts.py + "
        "HybridRetriever.retrieve()): retrieves top-5 chunks via RRF fusion of BM25 and "
        "dense vector search, constructs a prompt with full persona instructions, "
        "query-type instructions, and the _GROUNDING_INSTRUCTION hallucination guard."
    ))

    _callout(doc, (
        "Win Condition Logic  (eval/adaptive_vs_generic.py:351-354)\n\n"
        "# Sub-metric 1: compliance\n"
        "wins_compliance    = adaptive_compliance > generic_compliance\n\n"
        "# Sub-metric 2: readability fit\n"
        "wins_readability   = PERSONA_GRADE_TARGETS[persona][0]\n"
        "                     <= adaptive_fk_grade\n"
        "                     <= PERSONA_GRADE_TARGETS[persona][1]\n\n"
        "# Sub-metric 3: length adherence\n"
        "adaptive_adherence = max(0.0, 1.0 - |adaptive_words - adaptive_target| / adaptive_target)\n"
        "generic_adherence  = max(0.0, 1.0 - |generic_words  - 500|             / 500)\n"
        "wins_length        = adaptive_adherence > generic_adherence\n\n"
        "# Overall win (majority vote)\n"
        "overall_win        = sum([wins_compliance, wins_readability, wins_length]) >= 2\n\n"
        "# Primary metric\n"
        "win_rate           = df[\"adaptive_overall_wins\"].mean()"
    ), label="Win Condition Logic")

    _gap_warning(doc,
        "Brief vs Code Discrepancies",
        "'Adaptive Advantage Score' (composite mean of multiple deltas) is described in "
        "the project brief but is NOT computed. The actual primary metric is win_rate. "
        "No LLM judge is used -- all scoring is metric-based (regex compliance, textstat FK "
        "grade, word-count adherence). No faithfulness delta metric is included."
    )

    _heading2(doc, "5.3.1 PERSONA_GRADE_TARGETS")
    _para(doc, (
        "Flesch-Kincaid grade targets per persona (eval/adaptive_vs_generic.py:56-62). "
        "Used only in wins_readability sub-metric -- not in readability_analysis.py."
    ))
    _add_table(doc,
        ["Persona", "FK Grade Min", "FK Grade Max", "Rationale"],
        [
            ["novice", "6.0", "9.0", "6th-9th grade reading level (patient / junior coordinator)"],
            ["intermediate", "10.0", "13.0", "10th-13th grade (clinical research coordinator)"],
            ["expert", "14.0", "18.0", "University/graduate level (physician / scientist)"],
            ["regulatory", "14.0", "18.0", "Graduate level -- same as expert"],
            ["executive", "8.0", "12.0", "Accessible prose with brevity (C-suite)"],
        ]
    )

    _heading2(doc, "5.3.2 Length Targets")
    _para(doc, (
        "Adaptive length targets come from response_config.max_length "
        "(get_response_config(), src/personas.py). "
        "The generic baseline always uses 500 words as its fixed target."
    ))
    _add_table(doc,
        ["Persona", "Adaptive Target (words)", "Generic Baseline Target"],
        [
            ["novice", "300", "500 (hardcoded)"],
            ["intermediate", "500", "500 (hardcoded)"],
            ["expert", "1000", "500 (hardcoded)"],
            ["regulatory", "800", "500 (hardcoded)"],
            ["executive", "250", "500 (hardcoded)"],
        ]
    )

    _ascii_figure(doc, ASCII_ADAPTIVE_VS_GENERIC,
                  "Figure 5 -- Adaptive vs Generic: Head-to-Head Sequence")

    _heading2(doc, "5.4 Outputs")
    _add_table(doc,
        ["File", "Description"],
        [
            ["results/adaptive_vs_generic_results.csv",
             "One row per query: all metric columns + adaptive_overall_wins bool"],
            ["results/adaptive_vs_generic_detailed.json",
             "Full JSON with response text, generation times, all sub-metric scores"],
            ["results/adaptive_vs_generic_summary.txt",
             "Win rate, compliance delta, readability and length adherence summaries"],
        ]
    )

    _heading2(doc, "5.5 Code Reference")
    _add_table(doc,
        ["Step", "File", "Function / Symbol"],
        [
            ["Query set", "eval/adaptive_vs_generic.py:69", "COMPARISON_QUERIES (25 items)"],
            ["Grade targets", "eval/adaptive_vs_generic.py:56", "PERSONA_GRADE_TARGETS"],
            ["Generic prompt", "eval/adaptive_vs_generic.py:195", "build_generic_prompt()"],
            ["Adaptive prompt", "src/prompts.py:258", "build_adaptive_prompt()"],
            ["Hybrid retrieval", "src/retrieval.py", "HybridRetriever.retrieve()"],
            ["Win scoring", "eval/adaptive_vs_generic.py:351", "wins_compliance / wins_readability / wins_length"],
            ["Main runner", "eval/adaptive_vs_generic.py:222", "run_adaptive_vs_generic()"],
            ["Summary report", "eval/adaptive_vs_generic.py:420", "_write_summary()"],
        ]
    )

    _page_break(doc)


def build_section6(doc: Document):
    _heading1(doc, "6. Scoring Summary")

    _para(doc, (
        "The table below summarises all four evaluation modules -- their primary inputs, "
        "method, output metric, and the pass condition reported in each summary file."
    ))

    _add_table(doc,
        ["Eval", "Primary Input", "Method", "Output Metric", "Pass Condition"],
        [
            [
                "1. classification_accuracy",
                "45 labeled queries",
                "Regex classifier vs ground-truth labels",
                "Overall accuracy (correct / 45)",
                "No auto-threshold -- human review",
            ],
            [
                "2. readability_analysis",
                "25 persona responses (5q x 5p)",
                "textstat 7-metric suite per response",
                "FK grade per persona",
                "novice_fk_grade < expert_fk_grade",
            ],
            [
                "3. format_compliance",
                "25 persona responses + query_type",
                "16 regex rules; passed_applicable / total_applicable",
                "compliance_score in [0, 1]",
                "No auto-threshold -- per-rule pass rate inspected manually",
            ],
            [
                "4. adaptive_vs_generic",
                "25 comparison queries + PDF",
                "3 sub-metric majority vote per query",
                "win_rate = mean(adaptive_overall_wins)",
                "win_rate > 0.5 (adaptive wins majority)",
            ],
        ]
    )

    _para(doc, (
        "If any metric fails: (1) For classification accuracy, review the confusion matrix "
        "to identify which query types are most confused and refine their regex patterns "
        "in src/query_classifier.py. "
        "(2) For readability, inspect per-persona FK averages and adjust persona prompt "
        "instructions in src/prompts.py. "
        "(3) For format compliance, identify which rules fail most often and strengthen "
        "the corresponding prompt instruction. "
        "(4) For adaptive_vs_generic, a win_rate below 0.5 indicates the generic baseline "
        "is competitive -- investigate which sub-metric is dragging the result down."
    ))

    _page_break(doc)


def build_section7(doc: Document):
    _heading1(doc, "7. Running the Suite")

    _heading2(doc, "7.1 CLI Commands (run_eval.py)")
    _callout(doc, (
        "# Run all evaluations (adaptive suite + benchmarks)\n"
        "python run_eval.py --document data/irc.pdf --all\n\n"
        "# Run individual adaptive evaluations\n"
        "python run_eval.py --classify                                # no document required\n"
        "python run_eval.py --document data/irc.pdf --readability\n"
        "python run_eval.py --document data/irc.pdf --compliance\n"
        "python run_eval.py --document data/irc.pdf --adaptive-vs-generic\n\n"
        "# Aggregate metrics from existing result CSVs\n"
        "python run_eval.py --metrics"
    ), label="Command Reference")

    _heading2(doc, "7.2 Direct Python Imports")
    _callout(doc, (
        "from eval.classification_accuracy import run_classification_accuracy\n"
        "from eval.readability_analysis    import run_readability_analysis\n"
        "from eval.format_compliance       import run_format_compliance\n"
        "from eval.adaptive_vs_generic     import run_adaptive_vs_generic\n"
        "from eval.metrics                 import calculate_all_metrics\n\n"
        "# Eval 1 -- no arguments needed\n"
        "df_cls = run_classification_accuracy()\n\n"
        "# Evals 2, 3, 4 -- document path required for fresh generation\n"
        "df_read = run_readability_analysis(document_path=\"data/irc.pdf\")\n"
        "df_comp = run_format_compliance(document_path=\"data/irc.pdf\")\n"
        "df_avsg = run_adaptive_vs_generic(document_path=\"data/irc.pdf\")"
    ), label="Module-Level Imports")

    _heading2(doc, "7.3 Streamlit UI Trigger")
    _para(doc, (
        "The evaluation suite can also be triggered from the Streamlit web interface. "
        "Navigate to Tab 4 -- 'Benchmark Evals' -- and click the 'Run Selected Evaluations' "
        "button. Results are rendered inline and downloadable as CSV/TXT files."
    ))

    _heading2(doc, "7.4 Output Directory")
    _add_table(doc,
        ["Eval", "Output Files"],
        [
            ["classification_accuracy",
             "results/classification_accuracy_results.csv, "
             "classification_confusion_matrix.csv/.txt, "
             "classification_accuracy_summary.txt"],
            ["readability_analysis",
             "results/readability_analysis_results.csv, readability_analysis_summary.txt"],
            ["format_compliance",
             "results/format_compliance_results.csv, format_compliance_summary.txt"],
            ["adaptive_vs_generic",
             "results/adaptive_vs_generic_results.csv, adaptive_vs_generic_detailed.json, "
             "adaptive_vs_generic_summary.txt"],
            ["metrics aggregator",
             "Reads from all of the above CSVs -- no new files written"],
        ]
    )


# ---------------------------------------------------------------------------
# Main builder
# ---------------------------------------------------------------------------


def build_document() -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    _add_footer(doc)
    build_cover(doc)
    build_toc(doc)
    build_section1(doc)
    build_section2(doc)
    build_section3(doc)
    build_section4(doc)
    build_section5(doc)
    build_section6(doc)
    build_section7(doc)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
