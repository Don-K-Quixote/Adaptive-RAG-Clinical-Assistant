"""
Generate benchmark_evaluations.docx for the Clinical RAG Assistant.

Produces a professionally formatted technical document covering all four
benchmark evaluation modules: model_comparison, hybrid_comparison,
persona_evaluation, and latency_measurement.

Usage:
    python scripts/gen_benchmark_docx.py
"""

import os
import subprocess
import tempfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Colour palette — hex strings for XML, RGBColor for python-docx run/font
# ---------------------------------------------------------------------------
NAVY = RGBColor(0x11, 0x22, 0x40)       # #112240 dark navy
TEAL = RGBColor(0x0E, 0xA5, 0xC9)       # #0EA5C9 teal accent
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE_BG = RGBColor(0xE0, 0xF2, 0xFE)  # callout background
RED_BORDER = RGBColor(0xDC, 0x26, 0x26)

# Hex strings for direct XML shading (RGBColor is not subscriptable everywhere)
HEX_TEAL = "0EA5C9"
HEX_NAVY = "112240"
HEX_LIGHT_BLUE = "E0F2FE"
HEX_RED = "DC2626"

OUTPUT_PATH = Path("docs/rag-system/benchmark_evaluations_v2.docx")

# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str):
    """Set cell background using a plain hex string e.g. '0EA5C9'."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_run_font(run, size_pt=10, bold=False, italic=False,
                  color: RGBColor = None, mono=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    if mono:
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")


def _heading1(doc: Document, text: str):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(16)
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(6)
    return h


def _heading2(doc: Document, text: str):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(13)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    return h


def _para(doc: Document, text: str = "", size_pt=10.5):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size_pt)
    p.paragraph_format.space_after = Pt(6)
    return p


def _mono_para(doc: Document, text: str, size_pt=9):
    """Monospace paragraph for code references."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(size_pt)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    return p


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    """Add a table with teal shaded header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        _set_cell_bg(hdr_cells[i], HEX_TEAL)
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(10)

    # Data rows
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = str(val)
            for para in row_cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9.5)

    doc.add_paragraph()  # spacing after table
    return table


def _callout(doc: Document, text: str, label: str = ""):
    """Light-blue callout box with monospace content (for prompts)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        bdr = OxmlElement(f"w:{side}")
        bdr.set(qn("w:val"), "single")
        bdr.set(qn("w:sz"), "4")
        bdr.set(qn("w:space"), "4")
        bdr.set(qn("w:color"), "0EA5C9")
        pBdr.append(bdr)
    pPr.append(pBdr)
    # Background shading
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "E0F2FE")
    pPr.append(shd)
    if label:
        lr = p.add_run(f"{label}\n")
        lr.bold = True
        lr.font.size = Pt(9)
        lr.font.color.rgb = NAVY
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(8)
    return p


def _page_break(doc: Document):
    doc.add_page_break()


def _add_footer(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _add_toc(doc: Document):
    """Insert a Word TOC field."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar)
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run._r.append(fldChar2)
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar3)
    doc.add_paragraph(
        "[ Press Ctrl+A, then F9 in Word to update this Table of Contents ]"
    ).runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# ---------------------------------------------------------------------------
# Mermaid diagram helpers
# ---------------------------------------------------------------------------

def _try_render_mermaid(mmd_text: str, png_path: str) -> bool:
    """Attempt to render Mermaid diagram to PNG via mmdc CLI."""
    try:
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".mmd", delete=False, encoding="utf-8"
        ) as f:
            f.write(mmd_text)
            tmp = f.name
        result = subprocess.run(
            ["mmdc", "-i", tmp, "-o", png_path, "-b", "white", "-w", "900"],
            capture_output=True, timeout=30
        )
        os.unlink(tmp)
        return result.returncode == 0 and Path(png_path).exists()
    except Exception:
        return False


def _insert_diagram(doc: Document, mmd_text: str, caption: str,
                    png_filename: str, ascii_fallback: str):
    """Embed pre-rendered PNG if it exists; otherwise fall back to ASCII art."""
    png_path = Path("docs/rag-system") / png_filename
    if png_path.exists():
        doc.add_picture(str(png_path), width=Inches(6.0))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        # ASCII fallback
        _mono_para(doc, ascii_fallback, size_pt=8)
    cap_p = doc.add_paragraph(caption)
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap_p.runs:
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = NAVY
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Document sections
# ---------------------------------------------------------------------------

def build_cover(doc: Document):
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Benchmark Evaluations")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Clinical RAG Assistant — Technical Documentation")
    r2.font.size = Pt(16)
    r2.font.color.rgb = TEAL

    doc.add_paragraph()

    meta = [
        ("System", "Clinical RAG Assistant"),
        ("Document", "Benchmark Evaluations"),
        ("Version", "2.0"),
        ("Date", str(date.today())),
        ("Status", "Active Development"),
    ]
    for label, val in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = p.add_run(f"{label}: ")
        lr.bold = True
        lr.font.size = Pt(11)
        vr = p.add_run(val)
        vr.font.size = Pt(11)

    _page_break(doc)


def build_toc(doc: Document):
    _heading1(doc, "Table of Contents")
    _add_toc(doc)
    _page_break(doc)


def build_section1(doc: Document):
    _heading1(doc, "1  Benchmark Suite Overview")

    _para(doc, (
        "The Benchmark Evaluation framework provides systematic, reproducible measurements of "
        "the Clinical RAG Assistant's core capabilities — independently of any single user query. "
        "Unlike the Adaptive Evaluation suite, which assesses system behaviour on live queries, "
        "benchmarks run against a fixed test corpus and a controlled query set, making them "
        "suitable for regression testing, model selection, and infrastructure planning."
    ))
    _para(doc, (
        "Benchmarks are triggered on demand through the Streamlit Evaluate tab "
        "(app.py, lines 1115–1535) or programmatically via each module's entry-point function. "
        "They are not scheduled automatically; a clinical engineer or evaluator must initiate "
        "a run by uploading a reference document and clicking the relevant benchmark button. "
        "Results are persisted to the results/ directory as CSV and JSON files and displayed "
        "inline in the UI upon completion. The aggregate metrics module (eval/metrics.py) "
        "combines all four benchmark outputs into a single summary report."
    ))

    _heading2(doc, "1.1  Benchmark Summary Table")
    _add_table(doc,
        ["Benchmark", "Module File", "What It Measures", "When Run"],
        [
            [
                "Model Comparison",
                "eval/model_comparison.py",
                "Retrieval speed, result diversity, and per-query-type performance across six embedding model configurations",
                "On demand (30–60 min); requires reference document upload",
            ],
            [
                "Hybrid vs Semantic",
                "eval/hybrid_comparison.py",
                "Head-to-head retrieval quality (speed + diversity) of Hybrid RRF vs semantic-only vs lexical-only retrieval",
                "On demand (5–10 min); requires reference document upload",
            ],
            [
                "Persona Evaluation",
                "eval/persona_evaluation.py",
                "Response quality across all five user personas: word count adherence, generation latency, readability grades (Flesch, FK, Gunning Fog), and faithfulness score vs retrieved context",
                "On demand (10–15 min); requires document upload and LLM API access",
            ],
            [
                "Latency Measurement",
                "eval/latency_measurement.py",
                "End-to-end pipeline latency across setup and per-query stages (retrieval + generation), SLA pass/fail reporting (LATENCY_SLA), sequential throughput (queries/sec), and timestamped run snapshots",
                "On demand (~5 min); requires document upload and LLM API access",
            ],
        ]
    )

    _heading2(doc, "1.2  Benchmarking Pipeline Flow")

    mmd = """graph TD
    A[Test Corpus / Clinical Trial PDF] --> B[Document Loader & Chunker]
    B --> C1[eval/model_comparison.py\nrun_model_comparison]
    B --> C2[eval/hybrid_comparison.py\nrun_hybrid_comparison]
    B --> C3[eval/persona_evaluation.py\nrun_persona_evaluation]
    B --> C4[eval/latency_measurement.py\nrun_latency_measurement]
    C1 --> D[Scoring & Metrics\neval/metrics.py]
    C2 --> D
    C3 --> D
    C4 --> D
    D --> E[results/ CSV + JSON Files]
    E --> F[Streamlit Benchmark Results Panel\napp.py:1115-1535]
    F --> G[calculate_all_metrics\nfinal_metrics_summary.txt]"""

    ascii_art = """
  [Test Corpus / Clinical Trial PDF]
              |
    [DocumentLoader + RecursiveSemanticSplitter]
    /          |           |              \\
[model_    [hybrid_    [persona_    [latency_
comparison] comparison] evaluation] measurement]
    \\          |           |              /
        [eval/metrics.py: calculate_all_metrics()]
                      |
              [results/ CSV + JSON]
                      |
        [Streamlit Evaluate Tab — app.py:1115-1535]
                      |
          [final_metrics_summary.txt]
"""
    _insert_diagram(doc, mmd, "Figure 1 — Benchmarking Pipeline Flow",
                    "fig1_pipeline.png", ascii_art)

    _page_break(doc)


def build_section2(doc: Document):
    _heading1(doc, "2  Benchmark 1: Model Comparison")

    _heading2(doc, "2.1  Purpose")
    _para(doc, (
        "The Model Comparison benchmark measures how different embedding models perform "
        "when indexing and retrieving clinical trial document chunks. Because embedding "
        "quality directly determines which text passages are surfaced in response to a query, "
        "this benchmark informs the production model selection decision: given a fixed clinical "
        "document corpus, which embedding model yields the fastest retrieval and the most "
        "diverse, non-redundant result sets?"
    ))
    _para(doc, (
        "Six model configurations are evaluated (two general-purpose, three biomedical, "
        "one lightweight) against a ten-query test set drawn from RECIST 1.1 and clinical "
        "imaging workflows. The benchmark does not require LLM access — it measures retrieval "
        "quality only, making it fast and cost-free beyond compute time."
    ))

    _heading2(doc, "2.2  Inputs")
    _add_table(doc,
        ["Input", "Type", "Source", "Description"],
        [
            ["Reference document", "PDF file", "User-uploaded via Streamlit UI", "Clinical trial document (e.g. RECIST protocol); chunked using DEFAULT_CHUNK_SIZE=800, overlap=150"],
            ["Embedding model list", "list[str]", "src/config.py: EMBEDDING_MODELS dict", "Six model keys: all-mpnet-base-v2, all-MiniLM-L6-v2, S-PubMedBert-MS-MARCO (default), BioSimCSE-BioLinkBERT, BioBERT, bert-tiny-mnli"],
            ["Test query set", "list[dict]", "eval/model_comparison.py: DEFAULT_QUERIES", "10 queries with type labels: medical_technical (7), procedural (1), general (2)"],
            ["output_dir", "str", "Caller / Streamlit UI", "Directory for CSV/TXT output; default 'results'"],
        ]
    )

    _heading2(doc, "2.3  Method")
    _para(doc, (
        "For each embedding model in the configured list, the benchmark executes the "
        "following steps via evaluate_model() (eval/model_comparison.py):"
    ))
    steps = [
        "Load and chunk the reference document using load_and_chunk_document(), which calls RecursiveCharacterTextSplitter with DEFAULT_CHUNK_SIZE=800 and DEFAULT_CHUNK_OVERLAP=150.",
        "Initialise the EmbeddingService for the candidate model and build a FAISS vector index. Record index_time (seconds).",
        "For each of the 10 test queries, perform a TOP_K=5 vector similarity search and record retrieval_time_ms.",
        "Compute diversity_score for the 5 retrieved chunks using calculate_diversity_score() from src/utils.py.",
        "Aggregate per-model results into a DataFrame and write to results/model_comparison_results.csv.",
    ]
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(s).font.size = Pt(10.5)

    _para(doc, "")
    _heading2(doc, "2.3.1  Diversity Score Formula")
    _para(doc, (
        "Result diversity is the primary quality metric (beyond speed). "
        "It is computed by calculate_diversity_score() in src/utils.py:"
    ))
    _callout(doc, (
        "diversity_score = 0.4 × page_diversity + 0.6 × content_diversity\n\n"
        "page_diversity   = |unique_pages| / |total_retrieved_chunks|\n\n"
        "content_diversity = 1 - mean_pairwise_jaccard_similarity\n\n"
        "jaccard(A, B) = |words(A) ∩ words(B)| / |words(A) ∪ words(B)|\n\n"
        "Range: 0.0 (all chunks identical) → 1.0 (all chunks unique)"
    ), label="Diversity Score")

    _para(doc, (
        "No statistical significance testing is applied between models in the current implementation. "
        "Rankings are determined by raw metric values (mean retrieval time, mean diversity score) "
        "across the 10-query test set."
    ))

    _heading2(doc, "2.4  Diagram")
    mmd2 = """graph TD
    Q[10-Query Test Set\nmedical_technical + procedural + general] --> M1[all-mpnet-base-v2\nGeneral 768-dim]
    Q --> M2[all-MiniLM-L6-v2\nGeneral 384-dim]
    Q --> M3[S-PubMedBert-MS-MARCO\nMedical 768-dim DEFAULT]
    Q --> M4[BioSimCSE-BioLinkBERT\nMedical 768-dim]
    Q --> M5[BioBERT\nMedical 768-dim]
    Q --> M6[bert-tiny-mnli\nLightweight 128-dim]
    M1 --> R[Per-Model Scores\nindex_time · retrieval_time_ms · diversity_score]
    M2 --> R
    M3 --> R
    M4 --> R
    M5 --> R
    M6 --> R
    R --> CSV[model_comparison_results.csv\nmodel_comparison_summary.txt]"""

    ascii2 = """
  [10-Query Test Set]
  /    |    |    |    |    \\
[mpnet][MiniLM][PubMed][BioSim][BioBERT][tiny]
  \\    |    |    |    |    /
  [Per-Model: index_time, retrieval_ms, diversity_score]
              |
  [model_comparison_results.csv + _summary.txt]
"""
    _insert_diagram(doc, mmd2, "Figure 2 — Model Comparison Flow",
                    "fig2_model_comparison.png", ascii2)

    _heading2(doc, "2.5  Output")
    _add_table(doc,
        ["Metric", "Scale / Unit", "Interpretation", "File"],
        [
            ["index_time", "seconds", "Time to build FAISS index for all chunks; lower is better", "model_comparison_results.csv"],
            ["retrieval_time_ms", "milliseconds / query", "Mean query execution time across 10 queries; lower is better", "model_comparison_results.csv"],
            ["diversity_score", "0.0 – 1.0", "Mean result diversity; higher = less redundant chunks surfaced", "model_comparison_results.csv"],
            ["num_results", "integer (always 5)", "Number of chunks retrieved per query (DEFAULT_TOP_K=5)", "model_comparison_results.csv"],
        ]
    )
    _para(doc, (
        "The 'winning' model is not automatically flagged in the benchmark output. "
        "A clinical engineer reviews the summary file (results/model_comparison_summary.txt) "
        "and selects the model offering the best diversity_score at acceptable retrieval_time_ms "
        "for the clinical deployment context. The default production model is "
        "S-PubMedBert-MS-MARCO (medical domain, 768 dimensions) as configured in "
        "src/config.py: EMBEDDING_MODELS."
    ))

    _heading2(doc, "2.6  Code Reference")
    _mono_para(doc, "eval/model_comparison.py")
    _mono_para(doc, "  load_and_chunk_document(file_path, chunk_size, chunk_overlap) -> list[Document]")
    _mono_para(doc, "  evaluate_model(model_key, chunks, queries, output_dir) -> dict")
    _mono_para(doc, "  run_model_comparison(document_path, models, queries, output_dir) -> DataFrame")
    _mono_para(doc, "")
    _mono_para(doc, "src/utils.py :: calculate_diversity_score(documents) -> float")
    _mono_para(doc, "src/config.py :: EMBEDDING_MODELS (dict), DEFAULT_CHUNK_SIZE=800, DEFAULT_CHUNK_OVERLAP=150, DEFAULT_TOP_K=5")

    _add_table(doc,
        ["Parameter", "Default", "Description", "Configured In"],
        [
            ["models", "All 6 keys in EMBEDDING_MODELS", "List of embedding model keys to evaluate", "src/config.py: EMBEDDING_MODELS"],
            ["queries", "DEFAULT_QUERIES (10 queries)", "Test query set with type labels", "eval/model_comparison.py: DEFAULT_QUERIES"],
            ["chunk_size", "800 tokens", "Chunk size for document splitting", "src/config.py: DEFAULT_CHUNK_SIZE"],
            ["chunk_overlap", "150 tokens", "Overlap between consecutive chunks", "src/config.py: DEFAULT_CHUNK_OVERLAP"],
            ["top_k", "5", "Number of chunks retrieved per query", "src/config.py: DEFAULT_TOP_K"],
            ["output_dir", "'results'", "Directory for CSV and summary output", "Caller / Streamlit UI"],
        ]
    )
    _page_break(doc)


def build_section3(doc: Document):
    _heading1(doc, "3  Benchmark 2: Hybrid vs Semantic Retrieval")

    _heading2(doc, "3.1  Purpose")
    _para(doc, (
        "This benchmark quantifies the retrieval benefit of combining dense vector search "
        "with BM25 lexical search via Reciprocal Rank Fusion (RRF), relative to using "
        "either retrieval method in isolation. In a clinical trial context, this distinction "
        "is consequential: a query for 'SUVmax calculation' requires exact-term matching "
        "(lexical strength), while 'tumour assessment timing' requires semantic understanding. "
        "The benchmark provides empirical evidence for the hybrid strategy's superiority "
        "across a mixed medical and procedural query set."
    ))

    _heading2(doc, "3.2  Inputs")
    _add_table(doc,
        ["Input", "Type", "Source", "Description"],
        [
            ["Reference document", "PDF file", "User-uploaded via Streamlit UI", "Clinical trial PDF; chunked with same defaults as Model Comparison"],
            ["embedding_model", "str", "Caller / Streamlit UI", "Default: 'S-PubMedBert-MS-MARCO'; used for both dense and RRF retrieval"],
            ["Test query set", "list[dict]", "eval/hybrid_comparison.py: DEFAULT_QUERIES", "10 queries: medical (5), procedural (5)"],
            ["output_dir", "str", "Caller / Streamlit UI", "Default: 'results'"],
        ]
    )

    _heading2(doc, "3.3  Method")
    _para(doc, (
        "For each test query, compare_retrieval_methods() (eval/hybrid_comparison.py) "
        "executes three retrieval strategies in sequence and records timing and diversity "
        "for each:"
    ))
    strategies = [
        ("Hybrid (RRF)", "HybridRetriever.retrieve() — fuses FAISS dense results and BM25 lexical results using Reciprocal Rank Fusion with RRF_K_CONSTANT=60."),
        ("Semantic-only", "HybridRetriever.semantic_only() — FAISS vector similarity search only; no lexical component."),
        ("Lexical-only", "HybridRetriever.lexical_only() — BM25 keyword search only; no dense vector component."),
    ]
    for name, desc in strategies:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{name}: ").bold = True
        p.add_run(desc).font.size = Pt(10.5)

    _para(doc, "")
    _heading2(doc, "3.3.1  RRF Fusion Formula")
    _callout(doc, (
        "RRF score(d) = Σ_i  1 / (k + rank_i(d))\n\n"
        "where:\n"
        "  k        = RRF_K_CONSTANT = 60   (src/config.py)\n"
        "  rank_i   = position of document d in retrieval list i (1-indexed)\n"
        "  Σ_i      = sum over semantic retrieval list (i=1) and lexical list (i=2)\n\n"
        "Higher RRF score → document ranks higher in the fused list.\n"
        "k=60 is the default from Cormack et al. (2009) and dampens rank-order sensitivity."
    ), label="Reciprocal Rank Fusion (RRF)")

    _heading2(doc, "3.3.2  Computed Metrics")
    _para(doc, (
        "The benchmark computes the following metrics per query "
        "(implemented in compare_retrieval_methods(), eval/hybrid_comparison.py):"
    ))
    _add_table(doc,
        ["Metric", "Formula / Definition", "Scale"],
        [
            ["hybrid_time_ms / semantic_time_ms / lexical_time_ms", "Wall-clock time for each retrieval call (time.perf_counter)", "milliseconds"],
            ["hybrid_diversity / semantic_diversity / lexical_diversity", "calculate_diversity_score() applied to each result set", "0.0 – 1.0"],
            ["hybrid_in_both", "Count of documents appearing in both dense and lexical top-k", "integer"],
            ["hybrid_semantic_overlap", "|hybrid_results ∩ semantic_results| / top_k", "0.0 – 1.0"],
            ["hybrid_lexical_overlap", "|hybrid_results ∩ lexical_results| / top_k", "0.0 – 1.0"],
            ["semantic_lexical_overlap", "|semantic_results ∩ lexical_results| / top_k", "0.0 – 1.0"],
            ["diversity_improvement", "hybrid_diversity − semantic_diversity", "signed float"],
            ["time_overhead_ms", "hybrid_time_ms − semantic_time_ms", "milliseconds (overhead cost of fusion)"],
        ]
    )
    _para(doc, (
        "Note: Precision@k, Recall@k, MRR, and NDCG@k are NOT implemented in this benchmark. "
        "The benchmark relies on diversity_score as a proxy for retrieval quality because "
        "no ground-truth relevance labels exist for the test query set. "
        "If a labelled evaluation set is added in a future version, standard IR metrics "
        "could be incorporated."
    ))

    _heading2(doc, "3.3.3  Clinical Rationale for High-Recall Operation")
    _para(doc, (
        "In clinical trial document retrieval, failing to surface a relevant passage "
        "(a missed eligibility criterion, an unretrieved safety threshold) carries "
        "a higher risk than surfacing a marginally less relevant passage. "
        "The hybrid strategy is preferred because RRF fusion increases result diversity, "
        "reducing the probability of systematically missing relevant content. "
        "The time_overhead_ms cost is accepted as justified by this safety-oriented design goal."
    ))

    _heading2(doc, "3.4  Diagram")
    mmd3 = """graph LR
    Q[Test Query\n10 queries: medical + procedural] --> H[Hybrid Retriever\nDense FAISS + BM25 + RRF k=60]
    Q --> S[Semantic Retrieval\nDense FAISS Only]
    Q --> L[Lexical Retrieval\nBM25 Only]
    H --> M[compare_retrieval_methods\neval/hybrid_comparison.py]
    S --> M
    L --> M
    M --> R[Per-Query Metrics\nhybrid_time_ms · semantic_time_ms · lexical_time_ms\nhybrid_diversity · diversity_improvement · time_overhead_ms]
    R --> CSV2[hybrid_vs_semantic_comparison.csv\nhybrid_vs_semantic_summary.txt]"""

    ascii3 = """
[Test Query] ──── [Hybrid: FAISS + BM25 + RRF k=60] ───┐
     │                                                    │
     ├──── [Semantic: FAISS Dense Only] ─────────────────┤
     │                                                    │
     └──── [Lexical: BM25 Only] ──────────────────────── ┤
                                                          │
                              [compare_retrieval_methods()]
                                         │
              [hybrid_time_ms, semantic_time_ms, lexical_time_ms]
              [hybrid_diversity, diversity_improvement, time_overhead_ms]
                                         │
            [hybrid_vs_semantic_comparison.csv + _summary.txt]
"""
    _insert_diagram(doc, mmd3, "Figure 3 — Hybrid vs Semantic Comparison Flow",
                    "fig3_hybrid_semantic.png", ascii3)

    _heading2(doc, "3.5  Output")
    _add_table(doc,
        ["Metric", "Scale / Unit", "Interpretation", "File"],
        [
            ["hybrid/semantic/lexical_time_ms", "milliseconds", "Per-query retrieval latency for each strategy", "hybrid_vs_semantic_comparison.csv"],
            ["hybrid/semantic/lexical_diversity", "0.0 – 1.0", "Result diversity per strategy; higher = better coverage", "hybrid_vs_semantic_comparison.csv"],
            ["diversity_improvement", "signed float", "Positive = hybrid outperforms semantic on diversity", "hybrid_vs_semantic_comparison.csv"],
            ["time_overhead_ms", "milliseconds", "Extra latency cost of hybrid over semantic-only", "hybrid_vs_semantic_comparison.csv"],
            ["pct_improved (summary)", "percentage", "Fraction of queries where hybrid_diversity > semantic_diversity", "hybrid_vs_semantic_summary.txt"],
        ]
    )

    _heading2(doc, "3.6  Code Reference")
    _mono_para(doc, "eval/hybrid_comparison.py")
    _mono_para(doc, "  compare_retrieval_methods(query, hybrid_retriever, top_k=5) -> dict")
    _mono_para(doc, "  run_hybrid_comparison(document_path, queries, embedding_model, output_dir) -> DataFrame")
    _mono_para(doc, "")
    _mono_para(doc, "src/retrieval.py :: HybridRetriever.retrieve()")
    _mono_para(doc, "src/retrieval.py :: HybridRetriever.semantic_only()")
    _mono_para(doc, "src/retrieval.py :: HybridRetriever.lexical_only()")
    _mono_para(doc, "src/retrieval.py :: ReciprocalRankFusion.fuse()")
    _mono_para(doc, "src/config.py    :: RRF_K_CONSTANT=60, DEFAULT_SEMANTIC_WEIGHT=0.6, DEFAULT_BM25_WEIGHT=0.4")

    _add_table(doc,
        ["Parameter", "Default", "Description", "Configured In"],
        [
            ["embedding_model", "S-PubMedBert-MS-MARCO", "Dense embedding model used for vector retrieval", "src/config.py: EMBEDDING_MODELS"],
            ["queries", "DEFAULT_QUERIES (10 queries)", "10 medical + procedural test queries", "eval/hybrid_comparison.py: DEFAULT_QUERIES"],
            ["top_k", "5", "Number of results per retrieval call", "src/config.py: DEFAULT_TOP_K"],
            ["RRF_K_CONSTANT", "60", "RRF rank damping constant (Cormack et al. 2009)", "src/config.py: RRF_K_CONSTANT"],
            ["DEFAULT_SEMANTIC_WEIGHT", "0.6", "Weight for dense results in score blending (informational; RRF rank-based fusion used)", "src/config.py"],
            ["DEFAULT_BM25_WEIGHT", "0.4", "Weight for lexical results in score blending (informational)", "src/config.py"],
            ["output_dir", "'results'", "Output directory for CSV and summary files", "Caller / Streamlit UI"],
        ]
    )
    _page_break(doc)


def build_section4(doc: Document):
    _heading1(doc, "4  Benchmark 3: Persona Evaluation")

    _heading2(doc, "4.1  Purpose")
    _para(doc, (
        "The Persona Evaluation benchmark assesses whether the system produces responses "
        "that are genuinely adapted to the expertise level and communication needs of each "
        "user persona. It generates responses to a fixed five-query test set across all "
        "five configured user types and records response characteristics (word count, "
        "generation latency, format flags) per persona. "
        "This benchmark informs decisions about prompt engineering for each persona and "
        "validates that the adaptive response system delivers meaningfully different outputs "
        "for a NOVICE coordinator compared to an EXPERT principal investigator."
    ))

    _heading2(doc, "4.2  Inputs")
    _add_table(doc,
        ["Input", "Type", "Source", "Description"],
        [
            ["Reference document", "PDF file", "User-uploaded via Streamlit UI", "Clinical trial PDF"],
            ["embedding_model", "str", "Caller / Streamlit UI", "Default: 'S-PubMedBert-MS-MARCO'"],
            ["llm_model", "str", "src/config.py: DEFAULT_LLM_MODEL", "Default: 'gpt-4o-mini'"],
            ["Test query set", "list[str]", "eval/persona_evaluation.py: DEFAULT_QUERIES", "5 clinical trial questions covering RECIST, lesion measurement, compliance, imaging comparison, schedule"],
            ["User persona list", "list[UserType]", "eval/persona_evaluation.py: USER_TYPES", "All 5 UserType enum values: NOVICE, INTERMEDIATE, EXPERT, REGULATORY, EXECUTIVE"],
            ["output_dir", "str", "Caller / Streamlit UI", "Default: 'results'"],
        ]
    )

    _heading2(doc, "4.3  Method")
    _para(doc, (
        "For each query and each persona, evaluate_persona_responses() "
        "(eval/persona_evaluation.py) executes the following pipeline:"
    ))
    steps = [
        "Retrieve top-k=5 chunks using HybridRetriever.retrieve() with the configured embedding model.",
        "Classify the query using QueryClassifier.classify() to determine the QueryType.",
        "Build a ResponseConfig for the (UserType, QueryType) combination via get_response_config() (src/personas.py).",
        "Construct an adaptive prompt via build_adaptive_prompt() (src/prompts.py), incorporating the retrieved context, persona instructions, query-type instructions, and _GROUNDING_INSTRUCTION.",
        "Send the prompt to the configured LLM (gpt-4o-mini by default) and record the response, word count, character length, and generation_time_ms.",
        "Aggregate all 25 responses (5 queries × 5 personas) into a JSON structure and an HTML report.",
    ]
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(s).font.size = Pt(10.5)

    _para(doc, "")
    _heading2(doc, "4.3.1  Persona Configurations")
    _add_table(doc,
        ["Persona (UserType)", "detail_level", "max_length (words)", "Key Format Flags"],
        [
            ["NOVICE", "low", "300", "include_definitions=True, use_bullet_points=True, include_key_takeaway=True"],
            ["INTERMEDIATE", "medium", "500", "use_tables=True, include_examples=True"],
            ["EXPERT", "high", "1000", "use_tables=True, include_references=True"],
            ["REGULATORY", "high", "800", "use_tables=True, include_references=True, color_coding=True"],
            ["EXECUTIVE", "low", "250", "include_executive_summary=True, include_recommendations=True, use_tables=True"],
        ]
    )
    _para(doc, (
        "Note: The expertise level labels used in this documentation suite — "
        "Beginner, Intermediate, Expert — map to the codebase UserType values "
        "NOVICE, INTERMEDIATE, and EXPERT respectively. "
        "REGULATORY and EXECUTIVE are additional personas specific to the clinical "
        "trial governance context and have no direct equivalent in the three-label scheme."
    ))

    _heading2(doc, "4.3.2  Metrics Captured Per Response")
    _add_table(doc,
        ["Metric", "Type", "Description"],
        [
            ["word_count", "integer", "Number of words in the generated response"],
            ["response_length", "integer", "Character count of the generated response"],
            ["generation_time_ms", "float", "LLM generation wall-clock time in milliseconds"],
            ["detail_level", "str ('low'/'medium'/'high')", "Detail level from ResponseConfig for this persona"],
            ["use_tables / include_definitions / include_key_takeaway / include_executive_summary", "bool", "Format flags from ResponseConfig; used for format compliance cross-reference"],
            ["num_sources", "integer", "Number of document chunks retrieved (always 5 = DEFAULT_TOP_K)"],
            ["query_type", "str", "QueryType.value from QueryClassifier.classify()"],
            ["flesch_reading_ease", "float | None", "Flesch Reading Ease score (textstat); higher = simpler text. None if textstat not installed."],
            ["flesch_kincaid_grade", "float | None", "Flesch-Kincaid Grade Level (textstat). None if textstat not installed."],
            ["gunning_fog", "float | None", "Gunning Fog Index (textstat). None if textstat not installed."],
            ["faithfulness_score", "float | None", "Mean sentence-level cosine similarity to retrieved context (FaithfulnessChecker). None if embedder not passed."],
            ["faithfulness_warning", "bool | None", "True if faithfulness_score < FAITHFULNESS_WARNING_THRESHOLD (0.45). None if not scored."],
            ["faithfulness_low_confidence_count", "int | None", "Number of response sentences below sentence_threshold=0.35. None if not scored."],
            ["faithfulness_latency_ms", "float | None", "FaithfulnessChecker wall-clock time. None if not scored."],
        ]
    )
    _callout(doc, (
        "[RESOLVED] Readability and faithfulness quality scoring added to evaluate_persona_responses().\n\n"
        "Readability: textstat.flesch_reading_ease / flesch_kincaid_grade / gunning_fog computed per response.\n"
        "  Requires: conda install -c conda-forge textstat\n\n"
        "Faithfulness: FaithfulnessChecker(embedder).check(response, documents) run per (query, persona).\n"
        "  Score = mean cosine similarity of response sentences to retrieved context chunks.\n"
        "  Warning threshold: FAITHFULNESS_WARNING_THRESHOLD = 0.45 (src/faithfulness.py)\n\n"
        "Both metrics appear in persona_responses.json under each persona's metrics dict.\n"
        "Embedder is passed from run_persona_evaluation() so no additional model download is needed."
    ), label="Quality Scoring — RESOLVED")

    _heading2(doc, "4.4  Diagram")
    mmd4 = """graph TD
    Q[5 Test Queries] --> CL[QueryClassifier.classify\nsrc/query_classifier.py]
    Q --> HR[HybridRetriever.retrieve\nsrc/retrieval.py]
    CL --> RC[get_response_config\nsrc/personas.py]
    HR --> PB[build_adaptive_prompt\nsrc/prompts.py]
    RC --> PB
    PB --> N[UserType.NOVICE\nmax 300 words]
    PB --> I[UserType.INTERMEDIATE\nmax 500 words]
    PB --> E[UserType.EXPERT\nmax 1000 words]
    PB --> R[UserType.REGULATORY\nmax 800 words]
    PB --> X[UserType.EXECUTIVE\nmax 250 words]
    N --> SC[evaluate_persona_responses\neval/persona_evaluation.py]
    I --> SC
    E --> SC
    R --> SC
    X --> SC
    SC --> OUT1[persona_responses.json]
    SC --> OUT2[persona_responses_formatted.html]"""

    ascii4 = """
  [5 Test Queries]
      |        \\
[QueryClassifier] [HybridRetriever]
      |               |
[get_response_config] [Retrieved Chunks]
      |               |
      └── [build_adaptive_prompt()] ──┐
                                      │
        ┌─────────────────────────────┤
        │                             │
   [NOVICE]  [INTER.]  [EXPERT]  [REG.]  [EXEC.]
   300w       500w      1000w     800w    250w
        │
   [evaluate_persona_responses(embedder=embedder)]
        │
        ├── [textstat: flesch_reading_ease, flesch_kincaid_grade, gunning_fog]
        │
        ├── [FaithfulnessChecker(embedder).check(response, docs)]
        │     → faithfulness_score, faithfulness_warning,
        │       faithfulness_low_confidence_count
        │
   [persona_responses.json + _formatted.html]
   [save_run_snapshot() → results/history/<timestamp>/]
"""
    _insert_diagram(doc, mmd4, "Figure 4 — Persona Evaluation Flow",
                    "fig4_persona.png", ascii4)

    _heading2(doc, "4.5  Output")
    _add_table(doc,
        ["Output", "Format", "Content", "Path"],
        [
            ["persona_responses.json", "JSON", "List of dicts: {query, user_type, response, config, metrics} for all 25 (query, persona) combinations. Each metrics dict now includes readability grades and faithfulness_score.", "results/persona_responses.json"],
            ["persona_responses_formatted.html", "HTML", "Human-readable formatted report grouping responses by query and persona; includes word count and timing per response", "results/persona_responses_formatted.html"],
            ["results/history/<timestamp>/", "directory", "Timestamped snapshot of persona_responses.json + HTML from this run (save_run_snapshot())", "results/history/"],
        ]
    )
    _para(doc, (
        "Readability and faithfulness scores are written directly into each persona's metrics dict "
        "in persona_responses.json. A clinical engineer can inspect faithfulness_warning=True entries "
        "to identify responses that are not well-grounded in the retrieved context. "
        "Automated word-count adherence can be cross-referenced against "
        "RESPONSE_LENGTH_LIMITS in src/config.py."
    ))

    _heading2(doc, "4.6  Code Reference")
    _mono_para(doc, "eval/persona_evaluation.py")
    _mono_para(doc, "  evaluate_persona_responses(query, hybrid_retriever, llm, top_k=5, embedder=None) -> dict")
    _mono_para(doc, "  generate_html_report(results, output_path) -> None")
    _mono_para(doc, "  run_persona_evaluation(document_path, queries, embedding_model, llm_model, output_dir) -> list[dict]")
    _mono_para(doc, "")
    _mono_para(doc, "src/faithfulness.py :: FaithfulnessChecker(embedder).check(response, docs) -> FaithfulnessResult")
    _mono_para(doc, "src/faithfulness.py :: FAITHFULNESS_WARNING_THRESHOLD = 0.45")
    _mono_para(doc, "textstat            :: flesch_reading_ease(), flesch_kincaid_grade(), gunning_fog()")
    _mono_para(doc, "src/utils.py        :: save_run_snapshot(source_files, snapshot_dir) -> Path")
    _mono_para(doc, "")
    _mono_para(doc, "src/personas.py    :: UserType (enum), ResponseConfig (dataclass), get_response_config()")
    _mono_para(doc, "src/prompts.py     :: build_adaptive_prompt(), ResponseStyler.USER_TYPE_INSTRUCTIONS")
    _mono_para(doc, "src/retrieval.py   :: HybridRetriever.retrieve()")
    _mono_para(doc, "src/query_classifier.py :: QueryClassifier.classify(), QueryType (enum)")
    _mono_para(doc, "src/config.py      :: DEFAULT_LLM_MODEL='gpt-4o-mini', RESPONSE_LENGTH_LIMITS")

    _add_table(doc,
        ["Parameter", "Default", "Description", "Configured In"],
        [
            ["embedding_model", "S-PubMedBert-MS-MARCO", "Embedding model for hybrid retrieval", "src/config.py"],
            ["llm_model", "gpt-4o-mini", "LLM model for response generation", "src/config.py: DEFAULT_LLM_MODEL"],
            ["queries", "DEFAULT_QUERIES (5 queries)", "Fixed test query set", "eval/persona_evaluation.py"],
            ["USER_TYPES", "[NOVICE, INTERMEDIATE, EXPERT, REGULATORY, EXECUTIVE]", "All five personas evaluated", "eval/persona_evaluation.py"],
            ["top_k", "5", "Retrieved chunks per query", "src/config.py: DEFAULT_TOP_K"],
            ["RESPONSE_LENGTH_LIMITS", "{novice:300, intermediate:500, expert:1000, regulatory:800, executive:250}", "Max word count per persona (informational; enforced via prompt instruction)", "src/config.py"],
            ["output_dir", "'results'", "Output directory for JSON and HTML reports", "Caller / Streamlit UI"],
        ]
    )
    _page_break(doc)


def build_section5(doc: Document):
    _heading1(doc, "5  Benchmark 4: Latency Measurement")

    _heading2(doc, "5.1  Purpose")
    _para(doc, (
        "The Latency Measurement benchmark instruments the end-to-end query pipeline "
        "to identify throughput bottlenecks and verify that the system meets response-time "
        "expectations for interactive clinical use. It measures two distinct phases: "
        "(1) one-time setup latency (document loading, chunking, embedding model loading, "
        "and index construction) and (2) per-query latency (retrieval + LLM generation), "
        "averaged over multiple runs to reduce noise. "
        "The benchmark informs infrastructure decisions such as caching strategies, "
        "model selection trade-offs (speed vs quality), and acceptable query rates."
    ))

    _heading2(doc, "5.2  Inputs")
    _add_table(doc,
        ["Input", "Type", "Source", "Description"],
        [
            ["Reference document", "PDF file", "User-uploaded via Streamlit UI", "Clinical trial PDF"],
            ["embedding_model", "str", "Caller / Streamlit UI", "Default: 'S-PubMedBert-MS-MARCO'"],
            ["llm_model", "str", "src/config.py: DEFAULT_LLM_MODEL", "Default: 'gpt-4o-mini'"],
            ["Test query set", "list[str]", "eval/latency_measurement.py: DEFAULT_QUERIES", "10 representative clinical trial questions (expanded from 5 for more robust P50/P95 estimates)"],
            ["num_runs", "int", "Caller / Streamlit UI", "Default: 3 — number of times each query is executed for averaging"],
            ["output_dir", "str", "Caller / Streamlit UI", "Default: 'results'"],
        ]
    )

    _heading2(doc, "5.3  Method")
    _heading2(doc, "5.3.1  Setup Phase (One-Time Timing)")
    _para(doc, (
        "The following setup operations are individually timed using time.perf_counter() "
        "within run_latency_measurement() (eval/latency_measurement.py). "
        "Each measurement is a single wall-clock reading (not averaged):"
    ))
    _add_table(doc,
        ["Stage", "Operation", "Measured Variable"],
        [
            ["Document Loading", "PDF file read and text extraction", "doc_loading_time (seconds)"],
            ["Chunking", "RecursiveCharacterTextSplitter with chunk_size=800, overlap=150", "chunking_time (seconds)"],
            ["Embedder Loading", "HuggingFace model load and initialisation", "embedder_loading_time (seconds)"],
            ["Vector Indexing", "FAISS index construction from all chunk embeddings", "vector_indexing_time (seconds)"],
            ["BM25 Indexing", "BM25Retriever index construction from all chunks", "bm25_indexing_time (seconds)"],
        ]
    )

    _heading2(doc, "5.3.2  Per-Query Phase (Averaged Over num_runs)")
    _para(doc, (
        "For each query in the test set, measure_pipeline_latency() "
        "(eval/latency_measurement.py) executes the query num_runs=3 times and "
        "aggregates timing statistics:"
    ))
    _add_table(doc,
        ["Stage", "Operation", "Instrumentation"],
        [
            ["Retrieval", "HybridRetriever.retrieve(query, top_k=5)", "time.perf_counter() before and after retrieve() call"],
            ["LLM Generation", "LLM inference call with adaptive prompt", "time.perf_counter() before and after LLM call"],
            ["Total", "retrieval_time_ms + generation_time_ms", "Derived sum per run"],
        ]
    )

    _heading2(doc, "5.3.3  Statistical Aggregation")
    _callout(doc, (
        "For each query, over num_runs independent executions:\n\n"
        "retrieval_mean_ms  = mean(retrieval_times)\n"
        "retrieval_std_ms   = std(retrieval_times)\n"
        "retrieval_min_ms   = min(retrieval_times)\n"
        "retrieval_max_ms   = max(retrieval_times)\n\n"
        "generation_mean_ms = mean(generation_times)\n"
        "generation_std_ms  = std(generation_times)\n"
        "generation_min_ms  = min(generation_times)\n"
        "generation_max_ms  = max(generation_times)\n\n"
        "total_mean_ms      = retrieval_mean_ms + generation_mean_ms\n"
        "total_std_ms       = std(total_times)\n\n"
        "Across all queries (from eval/metrics.py: calculate_latency_metrics()):\n"
        "  p50_total_ms = 50th percentile of total_mean_ms values\n"
        "  p95_total_ms = 95th percentile of total_mean_ms values\n"
        "  max_total_ms = maximum total_mean_ms across all queries\n\n"
        "Throughput (sequential):\n"
        "  wall_elapsed_s = total wall-clock seconds for all query iterations\n"
        "  throughput_qps = len(queries) / wall_elapsed_s\n\n"
        "SLA pass/fail (src/config.py: LATENCY_SLA):\n"
        "  retrieval PASS: retrieval_mean_ms <= 500 ms\n"
        "  generation PASS: generation_mean_ms <= 8000 ms\n"
        "  total PASS: total_mean_ms <= 8000 ms\n"
        "  → Annotated per query in console output and in latency_summary.txt"
    ), label="Latency Aggregation Formulas + SLA + Throughput")

    _callout(doc, (
        "[RESOLVED] Test set expanded from 5 to 10 queries for more robust P50/P95 estimates.\n\n"
        "[RESOLVED] SLA pass/fail thresholds added to src/config.py: LATENCY_SLA.\n"
        "  retrieval_ms=500, generation_ms=8000, total_ms=8000 (all in milliseconds).\n"
        "  Each query is annotated [PASS] or [FAIL] in latency_summary.txt.\n\n"
        "[RESOLVED] Sequential throughput (queries/second) measured and written to\n"
        "  latency_summary.txt: THROUGHPUT section.\n\n"
        "Note: Concurrent load testing (multi-threaded throughput) remains out of scope.\n"
        "  The benchmark measures sequential single-threaded execution only."
    ), label="Latency Limitations — RESOLVED")

    _heading2(doc, "5.4  Diagram")
    mmd5 = """graph LR
    A[Query In] -->|t0| B[HybridRetriever.retrieve\nFAISS + BM25 + RRF]
    B -->|t1 - t0 = retrieval_ms| C[Retrieved Chunks\ntop_k=5]
    C -->|t1| D[build_adaptive_prompt\nUserType.INTERMEDIATE]
    D -->|t2| E[LLM Generation\ngpt-4o-mini]
    E -->|t3 - t2 = generation_ms| F[Response Out]
    F --> G[measure_pipeline_latency\nnum_runs=3 mean + std + min + max]
    G --> H[latency_results.csv\nlatency_summary.txt]"""

    ascii5 = """
[Query In] ──t0──► [HybridRetriever.retrieve()]
                           │
                    retrieval_ms = t1 - t0
                    SLA check: <= 500 ms [PASS/FAIL]
                           │
                   [build_adaptive_prompt()]
                           │
                   [LLM: gpt-4o-mini] ──t2──►
                           │
                    generation_ms = t3 - t2
                    SLA check: <= 8000 ms [PASS/FAIL]
                           │
                   [measure_pipeline_latency()
                    num_runs=3 → mean, std, min, max]
                    total SLA check: <= 8000 ms [PASS/FAIL]
                           │
                   [10-query wall time → throughput_qps]
                           │
             [latency_results.csv + latency_summary.txt]
             [save_run_snapshot() → results/history/<ts>/]
"""
    _insert_diagram(doc, mmd5, "Figure 5 — Latency Measurement Flow",
                    "fig5_latency.png", ascii5)

    _heading2(doc, "5.5  Latency Budget Table")
    _para(doc, (
        "The following table documents the latency stages instrumented in the benchmark. "
        "SLA thresholds for per-query stages are now defined in src/config.py: LATENCY_SLA "
        "and applied automatically in latency_summary.txt. "
        "Setup-phase stages (one-time costs) do not have SLA thresholds — they are "
        "reported for informational purposes only."
    ))
    _add_table(doc,
        ["Stage", "Metric Variables", "SLA Threshold", "Configured In"],
        [
            ["Document Loading (setup)", "load_time (ms)", "— (setup, not SLA-gated)", "N/A"],
            ["Chunking (setup)", "chunk_time (ms)", "— (setup, not SLA-gated)", "N/A"],
            ["Embedder Loading (setup)", "embed_load_time (ms)", "— (setup, not SLA-gated)", "N/A"],
            ["Vector Indexing (setup)", "index_time (ms)", "— (setup, not SLA-gated)", "N/A"],
            ["BM25 Indexing (setup)", "bm25_time (ms)", "— (setup, not SLA-gated)", "N/A"],
            ["Retrieval (per query)", "retrieval_mean_ms, std, min, max", "≤ 500 ms [PASS/FAIL]", "src/config.py: LATENCY_SLA['retrieval_ms']"],
            ["LLM Generation (per query)", "generation_mean_ms, std, min, max", "≤ 8000 ms [PASS/FAIL]", "src/config.py: LATENCY_SLA['generation_ms']"],
            ["Total E2E (per query)", "total_mean_ms, std", "≤ 8000 ms [PASS/FAIL]", "src/config.py: LATENCY_SLA['total_ms']"],
            ["P50 Total (aggregate)", "p50_total_ms across 10 queries", "— (informational)", "eval/metrics.py: calculate_latency_metrics()"],
            ["P95 Total (aggregate)", "p95_total_ms across 10 queries", "— (informational)", "eval/metrics.py: calculate_latency_metrics()"],
            ["Sequential Throughput", "throughput_qps = len(queries) / wall_elapsed_s", "— (informational)", "eval/latency_measurement.py"],
        ]
    )

    _heading2(doc, "5.6  Code Reference")
    _mono_para(doc, "eval/latency_measurement.py")
    _mono_para(doc, "  measure_pipeline_latency(query, hybrid_retriever, llm, user_type=INTERMEDIATE, num_runs=3) -> dict")
    _mono_para(doc, "  run_latency_measurement(document_path, queries, embedding_model, llm_model, num_runs, output_dir) -> DataFrame")
    _mono_para(doc, "")
    _mono_para(doc, "eval/metrics.py :: calculate_latency_metrics(df) -> dict")
    _mono_para(doc, "  → avg_retrieval_ms, avg_generation_ms, avg_total_ms, p50_total_ms, p95_total_ms, max_total_ms")
    _mono_para(doc, "")
    _mono_para(doc, "src/retrieval.py :: HybridRetriever.retrieve()")
    _mono_para(doc, "src/config.py    :: DEFAULT_LLM_MODEL='gpt-4o-mini', DEFAULT_TOP_K=5")
    _mono_para(doc, "src/config.py    :: LATENCY_SLA = {'retrieval_ms': 500, 'generation_ms': 8000, 'total_ms': 8000}")
    _mono_para(doc, "src/utils.py     :: save_run_snapshot(source_files, snapshot_dir) -> Path")

    _add_table(doc,
        ["Parameter", "Default", "Description", "Configured In"],
        [
            ["embedding_model", "S-PubMedBert-MS-MARCO", "Embedding model for retrieval", "src/config.py"],
            ["llm_model", "gpt-4o-mini", "LLM model for generation", "src/config.py: DEFAULT_LLM_MODEL"],
            ["queries", "DEFAULT_QUERIES (10 queries)", "Fixed test query set for latency profiling (expanded from 5 for robust percentile estimates)", "eval/latency_measurement.py"],
            ["num_runs", "3", "Executions per query for averaging", "eval/latency_measurement.py"],
            ["user_type", "UserType.INTERMEDIATE", "Persona used for prompt construction during latency test", "eval/latency_measurement.py: measure_pipeline_latency()"],
            ["top_k", "5", "Retrieved chunks per query", "src/config.py: DEFAULT_TOP_K"],
            ["output_dir", "'results'", "Output directory for CSV and summary", "Caller / Streamlit UI"],
        ]
    )
    _page_break(doc)


def build_section6(doc: Document):
    _heading1(doc, "6  Comparative Analysis & Reporting")

    _heading2(doc, "6.1  Results Storage")
    _para(doc, (
        "All benchmark outputs are written to the results/ directory at the end of each run. "
        "The primary output files (CSV, JSON, HTML, TXT) are overwritten on each run — "
        "this is intentional so that the Streamlit UI always displays the most recent results. "
        "To preserve run history, save_run_snapshot() (src/utils.py) automatically copies "
        "the output files to results/history/<YYYYMMDDTHHMMSS>/ at the end of each run. "
        "This is called by run_latency_measurement() and run_persona_evaluation()."
    ))
    _callout(doc, (
        "[RESOLVED] Run versioning implemented via save_run_snapshot().\n\n"
        "  src/utils.py :: save_run_snapshot(source_files, snapshot_dir) -> Path\n\n"
        "  Called at end of:\n"
        "    eval/latency_measurement.py: run_latency_measurement()\n"
        "      → snapshots: latency_results.csv, latency_summary.txt\n"
        "    eval/persona_evaluation.py: run_persona_evaluation()\n"
        "      → snapshots: persona_responses.json, persona_responses_formatted.html\n\n"
        "  Snapshots land in: results/history/<YYYYMMDDTHHMMSS>/\n"
        "  Each run directory is a complete, immutable copy of that run's outputs.\n"
        "  Trend analysis: compare CSV files across snapshot directories."
    ), label="Run Versioning — RESOLVED")
    _add_table(doc,
        ["File", "Format", "Written By", "Content"],
        [
            ["model_comparison_results.csv", "CSV", "eval/model_comparison.py: run_model_comparison()", "Per-query results: model, query, query_type, retrieval_time_ms, diversity_score, num_results"],
            ["model_comparison_summary.txt", "Text", "eval/model_comparison.py: run_model_comparison()", "Per-model aggregate: mean retrieval_time_ms, mean diversity_score"],
            ["hybrid_vs_semantic_comparison.csv", "CSV", "eval/hybrid_comparison.py: run_hybrid_comparison()", "Per-query: all timing and diversity metrics for 3 retrieval strategies"],
            ["hybrid_vs_semantic_summary.txt", "Text", "eval/hybrid_comparison.py: run_hybrid_comparison()", "Aggregate: mean times, mean diversities, diversity_improvement, pct_improved by query type"],
            ["latency_results.csv", "CSV", "eval/latency_measurement.py: run_latency_measurement()", "Per-query aggregated latency statistics (mean, std, min, max for retrieval + generation)"],
            ["latency_summary.txt", "Text", "eval/latency_measurement.py: run_latency_measurement()", "Setup times + overall latency statistics across all queries"],
            ["persona_responses.json", "JSON", "eval/persona_evaluation.py: run_persona_evaluation()", "All 25 (query, persona) responses with ResponseConfig and metrics"],
            ["persona_responses_formatted.html", "HTML", "eval/persona_evaluation.py: generate_html_report()", "Human-readable formatted output grouped by query"],
            ["final_metrics_summary.txt", "Text", "eval/metrics.py: generate_summary_report()", "Unified summary across all 8 evaluation modules (4 adaptive + 4 benchmark)"],
            ["results/history/<ts>/", "directory", "src/utils.py: save_run_snapshot()", "Timestamped copy of latency_results + persona_responses from each run"],
        ]
    )

    _heading2(doc, "6.2  Aggregate Metrics Module")
    _para(doc, (
        "eval/metrics.py provides calculate_all_metrics(results_dir='results') which "
        "loads all CSV and JSON result files and computes aggregate statistics across "
        "all evaluation modules. It is triggered via the 'Aggregate Metrics' button "
        "in the Streamlit Evaluate tab and writes results/final_metrics_summary.txt."
    ))
    _mono_para(doc, "eval/metrics.py :: calculate_all_metrics(results_dir) -> dict")
    _mono_para(doc, "  calculate_model_metrics(df)    -> {models: {...}, best_speed, best_diversity}")
    _mono_para(doc, "  calculate_hybrid_metrics(df)   -> {avg_times, avg_diversities, pct_improved}")
    _mono_para(doc, "  calculate_latency_metrics(df)  -> {avg_retrieval_ms, p50_total_ms, p95_total_ms}")
    _mono_para(doc, "  calculate_persona_metrics(data) -> {per-persona avg_word_count, avg_generation_ms}")
    _mono_para(doc, "  generate_summary_report(metrics, output_path) -> None")

    _heading2(doc, "6.3  Streamlit Results Panel")
    _para(doc, (
        "The Streamlit Evaluate tab (app.py, lines 1115–1535) surfaces benchmark results "
        "through the following pattern for each benchmark:"
    ))
    steps = [
        "User uploads a reference PDF document.",
        "User selects the benchmark from the Benchmark Evaluations section.",
        "User clicks the Run button; the entry-point function executes synchronously.",
        "Results are stored in st.session_state.benchmark_results[benchmark_name].",
        "Streamlit re-renders the results section showing metrics in cards and download buttons.",
        "CSV/TXT download buttons allow export of raw results files from results/.",
        "The 'Aggregate Metrics' button calls calculate_all_metrics() and displays the unified summary.",
    ]
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(s).font.size = Pt(10.5)

    _para(doc, (
        "Cross-run trend analysis is supported by comparing CSV files across snapshot directories "
        "in results/history/. The Streamlit UI does not yet render cross-run delta charts — "
        "a clinical engineer exports the CSVs from multiple snapshot directories for manual comparison."
    ))

    _heading2(doc, "6.4  Results Storage & Comparison Flow")
    mmd6 = """graph TD
    B1[model_comparison_results.csv] --> M[eval/metrics.py\ncalculate_all_metrics]
    B2[hybrid_vs_semantic_comparison.csv] --> M
    B3[latency_results.csv] --> M
    B4[persona_responses.json] --> M
    A1[classification_accuracy_results.csv] --> M
    A2[readability_analysis_results.csv] --> M
    A3[format_compliance_results.csv] --> M
    A4[adaptive_vs_generic_results.csv] --> M
    M --> S[final_metrics_summary.txt]
    S --> UI[Streamlit Evaluate Tab\nst.session_state.benchmark_results\napp.py:1115-1535]"""

    ascii6 = """
[model_comparison_results.csv]    ──┐
[hybrid_vs_semantic_comparison.csv] ┤
[latency_results.csv]               ┤── [eval/metrics.py]
[persona_responses.json]            ┤    calculate_all_metrics()
[classification_accuracy.csv]       ┤
[readability_analysis.csv]          ┤
[format_compliance.csv]             ┤
[adaptive_vs_generic.csv]          ──┘
                                         │
                              [final_metrics_summary.txt]
                                         │
                          [Streamlit Evaluate Tab — app.py:1115-1535]
"""
    _insert_diagram(doc, mmd6, "Figure 6 — Results Storage & Aggregation Flow",
                    "fig6_results_storage.png", ascii6)

    _heading2(doc, "6.5  Sample Benchmark Report Layout")
    _callout(doc, (
        "╔══════════════════════════════════════════════════════════════╗\n"
        "║  CLINICAL RAG BENCHMARK REPORT                               ║\n"
        "║  Run: [timestamp — manual]  |  Config: [manual versioning]   ║\n"
        "╠══════════════════════════════════════════════════════════════╣\n"
        "║  MODEL COMPARISON                                            ║\n"
        "║    Best diversity model:  [model_key]                        ║\n"
        "║    Mean retrieval time:   [X] ms                             ║\n"
        "║    Mean diversity score:  [0.0-1.0]                          ║\n"
        "╠══════════════════════════════════════════════════════════════╣\n"
        "║  HYBRID VS SEMANTIC                                          ║\n"
        "║    Hybrid avg diversity:  [X]                                ║\n"
        "║    Semantic avg diversity:[Y]                                 ║\n"
        "║    Diversity improvement: [+D]  |  % improved: [XX%]         ║\n"
        "║    Time overhead:         [ms]                               ║\n"
        "╠══════════════════════════════════════════════════════════════╣\n"
        "║  PERSONA EVALUATION                                          ║\n"
        "║    NOVICE   avg words: [W] | gen_time: [T] ms                ║\n"
        "║    INTERMED avg words: [W] | gen_time: [T] ms                ║\n"
        "║    EXPERT   avg words: [W] | gen_time: [T] ms                ║\n"
        "║    REGULAT. avg words: [W] | gen_time: [T] ms                ║\n"
        "║    EXEC.    avg words: [W] | gen_time: [T] ms                ║\n"
        "╠══════════════════════════════════════════════════════════════╣\n"
        "║  LATENCY MEASUREMENT                                         ║\n"
        "║    Avg retrieval:     [X] ms ± [std]                         ║\n"
        "║    Avg generation:    [Y] ms ± [std]                         ║\n"
        "║    Avg total E2E:     [Z] ms                                  ║\n"
        "║    P50 total:         [p50] ms  |  P95 total: [p95] ms        ║\n"
        "╠══════════════════════════════════════════════════════════════╣\n"
        "║  SOURCE: results/final_metrics_summary.txt                   ║\n"
        "╚══════════════════════════════════════════════════════════════╝"
    ), label="Sample Benchmark Report (populated from final_metrics_summary.txt)")
    _page_break(doc)


def build_section7(doc: Document):
    _heading1(doc, "7  Running the Benchmarks")

    _heading2(doc, "7.1  Via the Streamlit UI")
    _para(doc, (
        "All four benchmarks are accessible from the Evaluate tab in the Streamlit UI "
        "(app.py, Benchmark Evaluations section, lines ~1300–1535):"
    ))
    steps = [
        "Launch the application: streamlit run app.py",
        "Navigate to the Evaluate tab.",
        "Upload a reference clinical trial PDF document.",
        "Scroll to the Benchmark Evaluations subsection.",
        "Click the Run button for the desired benchmark.",
        "Wait for completion (estimated times: Model Comparison 30–60 min; Hybrid vs Semantic 5–10 min; Persona Evaluation 10–15 min; Latency Measurement ~3 min).",
        "View results in the rendered metrics cards and click Download to export CSV/TXT files.",
        "Optionally click Aggregate Metrics to generate final_metrics_summary.txt.",
    ]
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(s).font.size = Pt(10.5)

    _heading2(doc, "7.2  Programmatic API")
    _para(doc, "Each benchmark exposes a standalone entry-point function:")
    _callout(doc, (
        "from eval.model_comparison import run_model_comparison\n"
        "df = run_model_comparison(\n"
        "    document_path='path/to/protocol.pdf',\n"
        "    models=None,      # None = all 6 models from config\n"
        "    queries=None,     # None = DEFAULT_QUERIES (10 queries)\n"
        "    output_dir='results'\n"
        ")\n\n"
        "from eval.hybrid_comparison import run_hybrid_comparison\n"
        "df = run_hybrid_comparison(\n"
        "    document_path='path/to/protocol.pdf',\n"
        "    embedding_model='S-PubMedBert-MS-MARCO',\n"
        "    output_dir='results'\n"
        ")\n\n"
        "from eval.persona_evaluation import run_persona_evaluation\n"
        "results = run_persona_evaluation(\n"
        "    document_path='path/to/protocol.pdf',\n"
        "    embedding_model='S-PubMedBert-MS-MARCO',\n"
        "    llm_model='gpt-4o-mini',\n"
        "    output_dir='results'\n"
        ")\n\n"
        "from eval.latency_measurement import run_latency_measurement\n"
        "df = run_latency_measurement(\n"
        "    document_path='path/to/protocol.pdf',\n"
        "    num_runs=3,\n"
        "    output_dir='results'\n"
        ")\n\n"
        "# Aggregate all results\n"
        "from eval.metrics import calculate_all_metrics\n"
        "metrics = calculate_all_metrics(results_dir='results')"
    ), label="Programmatic Entry Points")

    _heading2(doc, "7.3  Prerequisite: OPENAI_API_KEY")
    _para(doc, (
        "Persona Evaluation and Latency Measurement require a valid OPENAI_API_KEY "
        "environment variable (or equivalent local Ollama configuration) because they "
        "invoke LLM generation. Model Comparison and Hybrid vs Semantic do not require "
        "LLM access — they measure retrieval only."
    ))

    _heading2(doc, "7.4  Gaps & Known Limitations")

    # ---- RESOLVED items — teal callout boxes --------------------------------
    resolved_items = [
        (
            "Persona evaluation: readability + faithfulness scoring — RESOLVED",
            "evaluate_persona_responses() now computes Flesch, FK Grade, and Gunning Fog "
            "(textstat) and FaithfulnessChecker cosine-similarity score per response. "
            "Metrics are stored in persona_responses.json per (query, persona) entry.\n"
            "Code: eval/persona_evaluation.py, src/faithfulness.py, textstat"
        ),
        (
            "Latency SLA thresholds — RESOLVED",
            "LATENCY_SLA = {retrieval_ms: 500, generation_ms: 8000, total_ms: 8000} added to "
            "src/config.py. Each query is annotated [PASS] or [FAIL] in latency_summary.txt.\n"
            "Code: src/config.py: LATENCY_SLA, eval/latency_measurement.py"
        ),
        (
            "Latency test set expanded + throughput measured — RESOLVED",
            "DEFAULT_QUERIES expanded from 5 to 10 queries. Sequential throughput "
            "(queries/second) measured and written to the THROUGHPUT section of latency_summary.txt.\n"
            "Code: eval/latency_measurement.py"
        ),
        (
            "Run versioning — RESOLVED",
            "save_run_snapshot() (src/utils.py) copies result files to results/history/<timestamp>/ "
            "at the end of run_latency_measurement() and run_persona_evaluation().\n"
            "Code: src/utils.py: save_run_snapshot(), eval/latency_measurement.py, eval/persona_evaluation.py"
        ),
    ]
    for title, body in resolved_items:
        _callout(doc, f"{body}", label=f"[RESOLVED] {title}")

    # ---- Remaining limitations — red-border paragraphs ----------------------
    remaining_gaps = [
        ("No ground-truth relevance labels",
         "Precision@k, Recall@k, MRR, and NDCG@k cannot be computed without a labelled "
         "relevance set for the test query corpus. "
         "Diversity score continues to serve as a retrieval quality proxy. "
         "Adding a labelled evaluation set would enable standard IR metrics."),
        ("No concurrent load testing",
         "Sequential throughput (queries/second) is now measured, but concurrent multi-user "
         "load testing is out of scope. "
         "The benchmark measures single-threaded sequential query execution only. "
         "A separate load-testing tool (e.g. Locust, k6) would be required for concurrency profiling."),
    ]
    for title, body in remaining_gaps:
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "12")
        left.set(qn("w:space"), "4")
        left.set(qn("w:color"), "DC2626")
        pBdr.append(left)
        pPr.append(pBdr)
        p.paragraph_format.left_indent = Inches(0.25)
        r1 = p.add_run(f"⚠️  {title}: ")
        r1.bold = True
        r1.font.size = Pt(10.5)
        r2 = p.add_run(body)
        r2.font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(8)


# ---------------------------------------------------------------------------
# Main builder
# ---------------------------------------------------------------------------

def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    _add_footer(doc)
    build_cover(doc)
    build_toc(doc)
    build_section1(doc)
    build_section2(doc)
    build_section3(doc)
    build_section4(doc)
    build_section5(doc)
    build_section6(doc)
    build_section7(doc)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
